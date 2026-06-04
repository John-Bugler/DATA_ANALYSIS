# =============================================================================
# ANALÝZA TRHU POZEMKŮ – POROVNÁVACÍ METODA (IVS / Market Approach)
# =============================================================================

import os
import warnings
import numpy as np
import pandas as pd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
import matplotlib.colors as mcolors
from matplotlib.colors import Normalize
from matplotlib.lines import Line2D
import seaborn as sns
from scipy import stats
import scipy.interpolate as interp
from shapely.geometry import Polygon
import folium
import branca.colormap as cm
import geopandas as gpd
import contextily as ctx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from docx.opc.constants import RELATIONSHIP_TYPE as RT

warnings.filterwarnings("ignore")

# =============================================================================
# 0. KONFIGURACE A PARAMETRY
# =============================================================================

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
INPUT_FILE = os.path.join(SCRIPT_DIR, "AI.xlsx")
OUT_XLSX   = os.path.join(SCRIPT_DIR, "Vystup_Analyza.xlsx")
OUT_MAP    = os.path.join(SCRIPT_DIR, "mapa_pozemky.html")
OUT_DOCX   = os.path.join(SCRIPT_DIR, "Znalecka_Zprava.docx")

# --- PARAMETRY OCEŇOVANÉ NEMOVITOSTI ---
OCE_DATUM        = pd.Timestamp("2026-06-02")
OCE_KU           = "Smíchov"          
OCE_OKRES        = "Hlavní město Praha" 
#OCE_UP_VSTUP     = ["LR", "ZMK", "PS", "PZO"]  # Zadejte libovolný počet zkratek do seznamu
OCE_UP_VSTUP     = ["OB", "OB-C", "OB-D", "OB-E", "SV", "SV-C", "SV-G", "SV-I", "OV", "OV-E", "OV-D"]  # Zadejte libovolný počet zkratek do seznamu

if isinstance(OCE_UP_VSTUP, str):
    OCE_UP_LIST = [OCE_UP_VSTUP.strip()]
else:
    OCE_UP_LIST = [str(x).strip() for x in OCE_UP_VSTUP]
OCE_UP = ", ".join(OCE_UP_LIST)

OCE_VYMERA       = 2000#138332.0           
OCE_PP_INDEX     = 0.70               
OCE_LAT          = 50.07132229220769  
OCE_LON          = 14.364680438635116

# =============================================================================
# !!! EXPERTNÍ NASTAVENÍ FILTRACE EXTRÉMŮ (OŘEZ VZORKU) !!!
# =============================================================================
# Tyto parametry můžete měnit, pokud jsou data nekvalitní, příliš oříznutá 
# nebo naopak pokud model propouští nesmyslné hodnoty.

# 1. Přísnost statistického filtru (Tukeyho násobek IQR pro Cenu i Výměru)
OUTLIER_IQR_MULT = 3.0 #1.5
# MATEMATIKA: Definuje šířku "kleští". Běžná hodnota je 1.5 (odřízne vše nad 75. percentil + 1.5x rozptyl trhu).
# LADĚNÍ: Pokud máte velmi nekonzistentní trh a mizí vám příliš mnoho dat, zvyšte na 2.0 nebo 3.0 (tzv. Far Outliers). 
# Tím zajistíte, že se vyřadí jen ty naprosté a neoddiskutovatelné extrémy. Pokud chcete naopak čistší střed, snižte.

# 2. Ochranné toleranční pásmo pro Výměru (Násobek vůči oceňované ploše)
VYMERA_TOLERANCE_NASOBEK = 5.0 #10.0
# MATEMATIKA: Pokud má oceňovaný pozemek extrémní plochu (např. 14 hektarů), běžná statistika by všechny větší obchody zahodila.
# LADĚNÍ: Hodnota 10.0 znamená, že model exaktně GARANTUJE ponechání pozemků, které jsou až 10x menší nebo 10x větší než oceňovaný.
# Pokud oceňujete běžný pozemek (např. 1000 m²) a chcete striktnější srovnatelnost, snižte na 3.0 nebo 5.0.

# 3. Spodní percentilový ořez (Ochrana proti pravostranné asymetrii nemovitostí)
PERCENTIL_DNO = 0.05
# MATEMATIKA: Křivka cen nemovitostí je asymetrická (dlouhý ocas doprava k luxusu, dole naráží na nulu). 
# Pokud matematický IQR výpočet "probije" podlahu (vrátí např. zápornou mez), model místo toho odřízne spodních X % prodejů.
# LADĚNÍ: 0.05 znamená, že vždy zahodí 5 % nejlevnějšího "odpadu" (např. prodeje symbolických podílů).

# 4. Absolutní minimální přípustná cena (Záchranná brzda zespodu)
MIN_POVOLENA_JC = 1000.0 #20.0
# MATEMATIKA: Tvrdý absolutní limit v Kč/m². Cokoliv pod tuto hodnotu je bez milosti smazáno ještě před koeficienty.
# LADĚNÍ: Pokud v datech vidíte příliš mnoho nesmyslných "korunových" převodů v rodině, zvyšte např. na 50.0.

# --- MANTINELY KOREKČNÍCH KOEFICIENTŮ (Rozptyl modelu) ---
K_CLIP_MIN = 0.20 #0.01  # Maximální povolená sleva (0.01 = srážka ceny na 1 %)
K_CLIP_MAX = 15.0 #150.0 # Maximální povolená přirážka (150.0 = obrovský skok u přechodu zeleň -> stavební)
# =============================================================================

SENSITIVITY_DELTA = 0.10         # Odchylka pro tornádový graf (zde 10 %)
DATE_REF = OCE_DATUM                 
CMAP_JC = "RdYlGn_r" # Globální paleta: Zelená = Levné, Žlutá = Střed, Červená = Drahé

# =============================================================================
# 1. NAČTENÍ DAT A PŘÍPRAVA
# =============================================================================

print("=" * 70)
print(f"ANALÝZA TRHU POZEMKŮ – {OCE_OKRES} / {OCE_KU}")
print("=" * 70)

df_raw = pd.read_excel(INPUT_FILE, sheet_name="data", header=0)
print(f"[1] Načtena data: {len(df_raw)} řádků")

df = df_raw.copy()
df.columns = df.columns.str.strip()
df["datum_podani"] = pd.to_datetime(df["datum_podani"], errors="coerce")
df["datum_den"] = df["datum_podani"].dt.date
if "JC_calc" in df.columns: df["JC"] = pd.to_numeric(df["JC_calc"], errors="coerce")
elif "#JC [Kč/m2]" in df.columns: df["JC"] = pd.to_numeric(df["#JC [Kč/m2]"], errors="coerce")
else: df["JC"] = pd.to_numeric(df["cenovy_udaj"], errors="coerce") / pd.to_numeric(df["#CELKOVA_VYMERA"], errors="coerce")
df["vymera"] = pd.to_numeric(df["#CELKOVA_VYMERA"], errors="coerce")
df["cas_v_dnech"] = (DATE_REF - df["datum_podani"]).dt.days
if "UP" in df.columns: df["UP"] = df["UP"].astype(str).str.strip()

def dominant_up(sub):
    if sub["UP"].nunique() == 1: return sub["UP"].iloc[0]
    return sub.groupby("UP")["vymera"].sum().idxmax()

agg_dict = {"datum_podani":"first", "cas_v_dnech":"first", "JC":"first", "vymera":"first", "ku_nazev":"first", "refPoint_lat":"first", "refPoint_lon":"first", "geometry_posList":"first", "parcel_number":"first", "ruian_parcela_id":"first"}
df_dedup = df.groupby("cislo_vkladu").agg({k:v for k,v in agg_dict.items() if k in df.columns}).reset_index()
up_per_vklad = df.groupby("cislo_vkladu").apply(dominant_up).reset_index()
up_per_vklad.columns = ["cislo_vkladu", "UP"]
df_dedup = df_dedup.merge(up_per_vklad, on="cislo_vkladu")
df_dedup.dropna(subset=["JC", "vymera", "datum_podani"], inplace=True)
df_dedup = df_dedup[df_dedup["JC"] > 0].reset_index(drop=True)

def parse_geometry(geom_str):
    try:
        if pd.isna(geom_str) or str(geom_str).strip() == "": return None
        nums = list(map(float, str(geom_str).split()))
        if len(nums) < 6 or len(nums) % 2 != 0: return None
        coords = [(nums[i], nums[i+1]) for i in range(0, len(nums), 2)]
        poly = Polygon(coords)
        return poly if poly.is_valid else poly.buffer(0)
    except Exception: return None

if "geometry_posList" in df_dedup.columns:
    df_dedup["polygon"] = df_dedup["geometry_posList"].apply(parse_geometry)
    df_dedup["pp_area"] = df_dedup["polygon"].apply(lambda p: p.area if p is not None else np.nan)
    df_dedup["pp_perim"] = df_dedup["polygon"].apply(lambda p: p.length if p is not None else np.nan)
    df_dedup["pp_index"] = df_dedup.apply(lambda r: (4 * np.pi * r["pp_area"]) / (r["pp_perim"] ** 2) if (pd.notna(r["pp_perim"]) and r["pp_perim"] > 0) else np.nan, axis=1)
else: df_dedup["pp_index"] = np.nan

# =============================================================================
# 2. DETEKCE EXTRÉMŮ (Cenová homogenizace)
# =============================================================================
print("[2] Detekce outlierů a cenová homogenizace …")

mask_oce_raw = df_dedup["UP"].isin(OCE_UP_LIST)
if mask_oce_raw.sum() > 0:
    median_oce_up_raw = df_dedup.loc[mask_oce_raw, "JC"].median()
else:
    median_oce_up_raw = df_dedup["JC"].median()

median_jc_per_up_raw = df_dedup.groupby("UP")["JC"].median()
df_dedup["JC_norm"] = df_dedup.apply(lambda r: r["JC"] * (median_oce_up_raw / median_jc_per_up_raw.get(r["UP"], median_oce_up_raw)), axis=1)

# A) IQR filtry pro HOMOGENIZOVANOU CENU
Q1_jcn, Q3_jcn = df_dedup["JC_norm"].quantile(0.25), df_dedup["JC_norm"].quantile(0.75)
IQR_jcn = Q3_jcn - Q1_jcn
if IQR_jcn == 0: IQR_jcn = median_oce_up_raw * 0.20

lower_bound_jc_norm = Q1_jcn - (OUTLIER_IQR_MULT * IQR_jcn)
upper_bound_jc_norm = Q3_jcn + (OUTLIER_IQR_MULT * IQR_jcn)

if lower_bound_jc_norm <= 0: 
    lower_bound_jc_norm = df_dedup["JC_norm"].quantile(PERCENTIL_DNO)
lower_bound_jc_norm = max(lower_bound_jc_norm, MIN_POVOLENA_JC)

# B) Striktní absolutní filtr VÝMĚRY (Žádná statistika IQR, čistě znalcova tolerance)
lower_bound_vym = OCE_VYMERA / VYMERA_TOLERANCE_NASOBEK
upper_bound_vym = OCE_VYMERA * VYMERA_TOLERANCE_NASOBEK

# Aplikace obou filtrů na data
df_dedup["outlier"] = (df_dedup["JC_norm"] < lower_bound_jc_norm) | (df_dedup["JC_norm"] > upper_bound_jc_norm) | \
                      (df_dedup["vymera"] < lower_bound_vym) | (df_dedup["vymera"] > upper_bound_vym)
df_outliers = df_dedup[df_dedup["outlier"]].copy()
df_clean = df_dedup[~df_dedup["outlier"]].copy().reset_index(drop=True)

# =============================================================================
# 3. KOREKCE A ETALON
# =============================================================================
print("[3] Popisné statistiky a koeficienty …")

def weighted_stats(series, weights=None):
    s = series.dropna()
    w = weights.loc[s.index].fillna(0) if weights is not None else None
    return {"počet": len(s), "průměr": np.average(s, weights=w) if w is not None and w.sum() > 0 else s.mean(), "min": s.min(), "P5": s.quantile(0.05), "Q1": s.quantile(0.25), "medián": s.median(), "Q3": s.quantile(0.75), "P95": s.quantile(0.95), "max": s.max(), "std": s.std()}

stats_cas = weighted_stats(df_clean["cas_v_dnech"])
stats_vym = weighted_stats(df_clean["vymera"], df_clean["vymera"])
stats_jc = weighted_stats(df_clean["JC"], df_clean["vymera"])

df_sub_up = df_clean[df_clean["UP"].isin(OCE_UP_LIST)]
has_sub_up = len(df_sub_up) >= 3 
if has_sub_up: stats_sub_jc = weighted_stats(df_sub_up["JC"], df_sub_up["vymera"])

df_reg = df_clean[df_clean["JC"] > 0].copy()
if len(df_reg) > 1 and df_reg["cas_v_dnech"].nunique() > 1:
    slope_k1, intercept_k1, r_k1, _, _ = stats.linregress(-df_reg["cas_v_dnech"] / 365.25, np.log(df_reg["JC"]))
else: slope_k1, intercept_k1, r_k1 = 0.0, df_clean["JC"].median(), 0.0
df_clean["K1"] = np.exp(slope_k1 * (df_clean["cas_v_dnech"] / 365.25)).clip(K_CLIP_MIN, K_CLIP_MAX)
rocni_zmena_k1 = (np.exp(slope_k1) - 1) * 100

if len(df_reg) > 1 and df_reg["vymera"].nunique() > 1:
    slope_k4, intercept_k4, r_k4, _, _ = stats.linregress(np.log(df_reg["vymera"]), np.log(df_reg["JC"]))
else: slope_k4, intercept_k4, r_k4 = 0.0, 0.0, 0.0
df_clean["K4"] = ((OCE_VYMERA / df_clean["vymera"]) ** slope_k4).clip(K_CLIP_MIN, K_CLIP_MAX)

median_jc_per_ku = df_clean.groupby("ku_nazev")["JC"].median()
median_oce_ku = median_jc_per_ku.get(OCE_KU, df_clean["JC"].median())
df_clean["K2"] = (median_oce_ku / df_clean["ku_nazev"].map(median_jc_per_ku)).clip(K_CLIP_MIN, K_CLIP_MAX)

median_jc_per_up = df_clean.groupby("UP")["JC"].median()
mask_oce_clean = df_clean["UP"].isin(OCE_UP_LIST)
median_oce_up = df_clean.loc[mask_oce_clean, "JC"].median() if mask_oce_clean.sum() > 0 else df_clean["JC"].median()
df_clean["K3"] = (median_oce_up / df_clean["UP"].map(median_jc_per_up)).clip(K_CLIP_MIN, K_CLIP_MAX)

pp_valid = df_clean["pp_index"].notna() & (df_clean["pp_index"] > 0)
df_clean.loc[pp_valid, "K5"] = (OCE_PP_INDEX / df_clean.loc[pp_valid, "pp_index"]).clip(K_CLIP_MIN, K_CLIP_MAX)
df_clean.loc[~pp_valid, "K5"] = 1.0

df_clean["JC_upravena"] = df_clean["JC"] * df_clean["K1"] * df_clean["K2"] * df_clean["K3"] * df_clean["K4"] * df_clean["K5"]
valid_mask = df_clean["JC_upravena"].notna() & df_clean["vymera"].notna()
etalon_vazeny = np.average(df_clean.loc[valid_mask, "JC_upravena"], weights=df_clean.loc[valid_mask, "vymera"]) if valid_mask.sum() > 0 else df_clean["JC"].median()
etalon_median = df_clean["JC_upravena"].median() if pd.notna(df_clean["JC_upravena"].median()) else etalon_vazeny
etalon_std = df_clean["JC_upravena"].std() if pd.notna(df_clean["JC_upravena"].std()) else 0.0

def prepocitej_citlivost(df_base, zmeny):
    d = df_base.copy()
    s1 = slope_k1 * (1 + zmeny.get("s_k1", 0))
    mk2 = median_oce_ku * (1 + zmeny.get("m_k2", 0))
    mk3 = median_oce_up * (1 + zmeny.get("m_k3", 0))
    s4 = slope_k4 * (1 + zmeny.get("s_k4", 0))
    pk5 = OCE_PP_INDEX * (1 + zmeny.get("p_k5", 0))
    
    k1 = np.exp(s1 * (d["cas_v_dnech"]/365.25)).clip(K_CLIP_MIN, K_CLIP_MAX)
    k2 = (mk2 / d["ku_nazev"].map(d.groupby("ku_nazev")["JC"].median())).clip(K_CLIP_MIN, K_CLIP_MAX)
    k3 = (mk3 / d["UP"].map(d.groupby("UP")["JC"].median())).clip(K_CLIP_MIN, K_CLIP_MAX)
    k4 = ((OCE_VYMERA / d["vymera"]) ** s4).clip(K_CLIP_MIN, K_CLIP_MAX)
    k5 = pd.Series(1.0, index=d.index); v = d["pp_index"].notna() & (d["pp_index"] > 0)
    k5.loc[v] = (pk5 / d.loc[v, "pp_index"]).clip(K_CLIP_MIN, K_CLIP_MAX)
    j = d["JC"] * k1 * k2 * k3 * k4 * k5; vm = j.notna() & d["vymera"].notna()
    return np.average(j[vm], weights=d.loc[vm, "vymera"]) if vm.sum() > 0 else d["JC"].median()

sens = {}
for k, p in {"K1":"s_k1", "K2":"m_k2", "K3":"m_k3", "K4":"s_k4", "K5":"p_k5"}.items():
    u, dn = prepocitej_citlivost(df_clean, {p:SENSITIVITY_DELTA}), prepocitej_citlivost(df_clean, {p:-SENSITIVITY_DELTA})
    sens[k] = {"base": etalon_vazeny, "up": u, "dn": dn, "rozpeti": abs(u-dn)}
tornado_order = sorted(sens.keys(), key=lambda x: sens[x]["rozpeti"], reverse=True)

# =============================================================================
# 4. GRAFY A WORD DOKUMENT
# =============================================================================
print("[4] Generuji grafy …")

GRAPH_FILES = {}
sns.set_theme(style="whitegrid", font="Arial")

# 1. Histogram JC
fig, ax = plt.subplots(figsize=(10, 5))
bins = np.logspace(np.log10(max(df_dedup["JC"].min(), 1)), np.log10(df_dedup["JC"].max()), 20)
ax.hist(df_dedup["JC"], bins=bins, color="#4472C4", edgecolor="white", alpha=0.85); ax.set_xscale('log')
ax.axvline(stats_jc["medián"], color="red", linestyle="--", linewidth=1.5, label=f'Medián čistého trhu ({stats_jc["medián"]:,.0f})')
ax.axvline(stats_jc["průměr"], color="green", linestyle="--", linewidth=1.5, label=f'Vážený průměr trhu ({stats_jc["průměr"]:,.0f})')
ax.set_ylabel("Počet vkladů"); ax.set_xlabel("JC [Kč/m²] (Log)"); ax.set_title("Histogram jednotkových cen pozemků na trhu")
ax.xaxis.set_major_formatter(mticker.FuncFormatter(lambda x, _: f"{x:,.0f}")); ax.legend(); plt.tight_layout()
GRAPH_FILES["histogram_JC"] = p = os.path.join(SCRIPT_DIR, "graf_histogram_JC.png"); fig.savefig(p, dpi=150); plt.close(fig)

# 2. Časový vývoj (Opravený trendline a zabezpečená osa Y > 0)
fig, ax = plt.subplots(figsize=(12, 5))
sc = ax.scatter(df_clean["datum_podani"], df_clean["JC"], c=df_clean["JC"], cmap=CMAP_JC, norm=Normalize(vmin=df_clean["JC"].min(), vmax=df_clean["JC"].max()), s=60, alpha=0.85)

date_range = pd.date_range(df_clean["datum_podani"].min(), df_clean["datum_podani"].max(), periods=100)
cas_roky_pred = -((DATE_REF - date_range).days) / 365.25
jc_trend = np.exp(slope_k1 * cas_roky_pred + intercept_k1)

ax.plot(date_range, jc_trend, color="navy", linewidth=2, label="Log-lin trend (K1)")
for val, lbl, c, ls in [(stats_jc["medián"], "Medián", "red", "--"), (stats_jc["průměr"], "Průměr", "green", "--")]: 
    ax.axhline(val, linestyle=ls, color=c, label=f"{lbl} {val:,.0f}")
plt.colorbar(sc, ax=ax, label="JC [Kč/m²]")
ax.set_ylabel("JC [Kč/m²]")
ax.set_ylim(bottom=0) # Zajišťuje, že osa Y nespadne do minusu
ax.set_title("Časový vývoj")
ax.legend()
plt.tight_layout()
GRAPH_FILES["casovy_vyvoj"] = p = os.path.join(SCRIPT_DIR, "graf_casovy_vyvoj.png"); fig.savefig(p, dpi=150); plt.close(fig)

# 3. ÚP
up_s = df_clean.groupby("UP")["JC"].agg(["median", "mean", "std"]).reset_index()
fig, ax = plt.subplots(figsize=(10, 5)); x = np.arange(len(up_s))
ax.bar(x-0.2, up_s["median"], 0.4, label="Medián", color="#4472C4"); ax.bar(x+0.2, up_s["mean"], 0.4, label="Průměr", color="#ED7D31", alpha=0.85)
ax.errorbar(x+0.2, up_s["mean"], yerr=up_s["std"].fillna(0), fmt="none", ecolor="black", capsize=4); ax.set_xticks(x); ax.set_xticklabels(up_s["UP"], rotation=45, ha="right")
ax.set_ylabel("JC [Kč/m²]"); ax.set_title("Porovnání JC dle ÚP"); ax.legend(); plt.tight_layout()
GRAPH_FILES["JC_dle_UP"] = p = os.path.join(SCRIPT_DIR, "graf_JC_dle_UP.png"); fig.savefig(p, dpi=150); plt.close(fig)

# 4. Velikost
fig, ax = plt.subplots(figsize=(10, 5))
sc4 = ax.scatter(df_clean["vymera"], df_clean["JC"], c=df_clean["JC"], cmap=CMAP_JC, norm=Normalize(vmin=df_clean["JC"].min(), vmax=df_clean["JC"].max()), s=60, alpha=0.8)
v_r = np.linspace(df_clean["vymera"].min(), df_clean["vymera"].max(), 200)
ax.plot(v_r, np.exp(slope_k4 * np.log(v_r) + intercept_k4) if intercept_k4 else np.repeat(df_clean["JC"].median(), 200), color="navy", linewidth=2, label=f"Log-log trend (β={slope_k4:.3f})")
ax.set_xlabel("Výměra [m²]"); ax.set_ylabel("JC [Kč/m²]"); ax.set_title("Závislost JC na velikosti"); plt.colorbar(sc4); ax.legend(); plt.tight_layout()
GRAPH_FILES["JC_vs_vymera"] = p = os.path.join(SCRIPT_DIR, "graf_JC_vs_vymera.png"); fig.savefig(p, dpi=150); plt.close(fig)

# 5. Outliers
fig, ax = plt.subplots(figsize=(10, 5))
ax.scatter(df_clean["datum_podani"], df_clean["JC_norm"], color="#4472C4", s=40, label="Standardní trh", alpha=0.7)
if len(df_outliers)>0: ax.scatter(df_outliers["datum_podani"], df_outliers["JC_norm"], color="red", s=80, marker="X", label="Outlier", zorder=5)
ax.axhline(lower_bound_jc_norm, color="orange", linestyle="--", label=f"Dolní ochranné pásmo ({lower_bound_jc_norm:,.0f})")
ax.axhline(upper_bound_jc_norm, color="orange", linestyle="--", label=f"Horní ochranné pásmo ({upper_bound_jc_norm:,.0f})")
ax.set_title(f"Extrémy na Homogenizovaných Datech (Ceny převedeny na ÚP {OCE_UP})"); ax.legend(fontsize=8); ax.set_ylabel("Cena JC_norm"); plt.tight_layout()
GRAPH_FILES["outliery"] = p = os.path.join(SCRIPT_DIR, "graf_outliery.png"); fig.savefig(p, dpi=150); plt.close(fig)

# 6. Tornado
fig, ax = plt.subplots(figsize=(10, 5))
labels_map = {"K1": "K1 – Čas", "K2": "K2 – Lokalita", "K3": "K3 – Typ ÚP", "K4": "K4 – Velikost", "K5": "K5 – Tvar"}
m_x, m_X = float('inf'), float('-inf')
for i, k in enumerate(tornado_order):
    r = sens[k]
    lo, hi = min(r["dn"], r["up"]), max(r["dn"], r["up"])
    m_x, m_X = min(m_x, lo), max(m_X, hi)
    ax.barh(i, hi-etalon_vazeny, left=etalon_vazeny, color="#70AD47", height=0.5); ax.barh(i, lo-etalon_vazeny, left=etalon_vazeny, color="#FF0000", height=0.5)
r_x = max(m_X - m_x, 100); ax.set_xlim(m_x - r_x*0.25, m_X + r_x*0.25); ax.axvline(etalon_vazeny, color="black", linewidth=1.5)
ax.set_yticks(range(len(tornado_order))); ax.set_yticklabels([labels_map[k] for k in tornado_order]); ax.invert_yaxis(); ax.set_title("Citlivostní analýza"); plt.tight_layout()
GRAPH_FILES["tornado"] = p = os.path.join(SCRIPT_DIR, "graf_tornado.png"); fig.savefig(p, dpi=150); plt.close(fig)

# 7. Boxplot ÚP
if has_sub_up:
    fig, ax = plt.subplots(figsize=(8, 5)); df_clean['Skupina'] = np.where(df_clean['UP'].isin(OCE_UP_LIST), f'ÚP ({OCE_UP})', 'Ostatní ÚP')
    sns.boxplot(data=df_clean, x='Skupina', y='JC', ax=ax, palette=['#FF9999', '#99CCFF']); ax.set_title(f"Rozptyl JC: ÚP {OCE_UP} vůči zbytku"); plt.tight_layout()
    GRAPH_FILES["podmnozina_up"] = p = os.path.join(SCRIPT_DIR, "graf_podmnozina_up.png"); fig.savefig(p, dpi=150); plt.close(fig)

# 8. K5 PP Index
pp_vd = df_clean.loc[df_clean["pp_index"].notna() & (df_clean["pp_index"] > 0), "pp_index"]
if not pp_vd.empty:
    fig, ax = plt.subplots(figsize=(10, 5)); sns.histplot(pp_vd, kde=True, color="#9DC3E6", bins=20, ax=ax)
    ax.axvline(OCE_PP_INDEX, color="red", linestyle="--", linewidth=2, label=f"Oceňovaný ({OCE_PP_INDEX})"); ax.axvline(pp_vd.median(), color="green", linestyle=":", linewidth=2, label=f"Trh ({pp_vd.median():.2f})")
    ax.set_title("Rozložení tvarového indexu"); ax.legend(); plt.tight_layout()
    GRAPH_FILES["K5_tvar"] = p = os.path.join(SCRIPT_DIR, "graf_K5_tvar.png"); fig.savefig(p, dpi=150); plt.close(fig)

# 9. Mapy
LAT_COL, LON_COL = "refPoint_lat", "refPoint_lon"
valid_gps = df_dedup.dropna(subset=[LAT_COL, LON_COL])
if not valid_gps.empty:
    gdf = gpd.GeoDataFrame(valid_gps, geometry=gpd.points_from_xy(valid_gps[LON_COL], valid_gps[LAT_COL]), crs="EPSG:4326").to_crs("EPSG:3857")
    fig, ax = plt.subplots(figsize=(10, 8))
    gdf.plot(ax=ax, column="JC", cmap=CMAP_JC, norm=Normalize(vmin=df_dedup["JC"].min(), vmax=df_dedup["JC"].max()), markersize=80, edgecolor="black", alpha=0.8, legend=True)
    oce_gdf = gpd.GeoDataFrame(index=[0], crs="EPSG:4326", geometry=[gpd.points_from_xy([OCE_LON], [OCE_LAT])[0]]).to_crs("EPSG:3857")
    ax.scatter(oce_gdf.geometry.x, oce_gdf.geometry.y, color="blue", marker="^", s=300, edgecolors="white", zorder=10)
    ctx.add_basemap(ax, crs=gdf.crs.to_string(), source=ctx.providers.Esri.WorldImagery); ax.set_axis_off(); plt.tight_layout()
    GRAPH_FILES["prostorova_mapa"] = p = os.path.join(SCRIPT_DIR, "graf_prostorova_mapa.png"); fig.savefig(p, dpi=150); plt.close(fig)
    
    vg_clean = df_clean.dropna(subset=[LAT_COL, LON_COL])
    if len(vg_clean) > 4:
        try:
            gdf_c = gpd.GeoDataFrame(vg_clean, geometry=gpd.points_from_xy(vg_clean[LON_COL], vg_clean[LAT_COL]), crs="EPSG:4326").to_crs("EPSG:3857")
            x, y, z = gdf_c.geometry.x.values, gdf_c.geometry.y.values, gdf_c['JC'].values
            xi, yi = np.linspace(x.min()-1500, x.max()+1500, 300), np.linspace(y.min()-1500, y.max()+1500, 300)
            XI, YI = np.meshgrid(xi, yi)
            ZI = interp.griddata((x, y), z, (XI, YI), method='linear')
            fig, ax = plt.subplots(figsize=(10, 8))
            contour = ax.contourf(XI, YI, ZI, levels=30, cmap=CMAP_JC, alpha=0.85); plt.colorbar(contour, ax=ax)
            gdf_c.plot(ax=ax, color='black', markersize=15, alpha=0.4)
            ax.scatter(oce_gdf.geometry.x, oce_gdf.geometry.y, color="blue", marker="^", s=250, edgecolor="white", zorder=10)
            ctx.add_basemap(ax, crs=gdf_c.crs.to_string(), source=ctx.providers.Esri.WorldImagery); ax.set_axis_off(); plt.tight_layout()
            GRAPH_FILES["heatmap"] = p = os.path.join(SCRIPT_DIR, "graf_heatmap.png"); fig.savefig(p, dpi=150); plt.close(fig)
        except: pass

# --- MAPA INTERAKTIVNÍ ---
if not valid_gps.empty: map_center = [(valid_gps[LAT_COL].min() + valid_gps[LAT_COL].max())/2, (valid_gps[LON_COL].min() + valid_gps[LON_COL].max())/2]
else: map_center = [OCE_LAT, OCE_LON]
m = folium.Map(location=map_center, prefer_canvas=True); folium.TileLayer("OpenStreetMap").add_to(m)
jc_min, jc_max = df_dedup["JC"].min() if not pd.isna(df_dedup["JC"].min()) else 0, df_dedup["JC"].max() if not pd.isna(df_dedup["JC"].max()) else 1000
colormap = cm.LinearColormap(colors=['#1a9641', '#a6d96a', '#fdae61', '#d7191c'], vmin=jc_min, vmax=jc_max, caption='JC [Kč/m²]')
m.add_child(colormap)
for _, row in df_dedup.iterrows():
    if pd.isna(row.get(LAT_COL)): continue
    popup_txt = f"<b>Vklad: {row['cislo_vkladu']}</b><br>KÚ: {row.get('ku_nazev','–')}<br>UP: {row.get('UP','–')}<br>JC: <b>{row['JC']:,.0f} Kč/m²</b><br>Výměra: {row['vymera']:,.0f} m²"
    folium.CircleMarker(location=[row[LAT_COL], row[LON_COL]], radius=8, color="black", weight=1, fill=True, fill_color=colormap(row["JC"]), fill_opacity=0.9, popup=folium.Popup(popup_txt, max_width=250), tooltip=f"{row['JC']:,.0f} Kč/m²").add_to(m)
folium.Marker(location=[OCE_LAT, OCE_LON], popup=folium.Popup(f"<b>OCEŇOVANÁ NEMOVITOST</b><br>{OCE_UP}<br>{OCE_VYMERA:,.0f} m²", max_width=220), icon=folium.Icon(color="blue", icon="home", prefix="fa")).add_to(m)
m.save(OUT_MAP)

# --- GENERACE WORDU ---
print("[14] Generuji finální Word dokument …")
def set_cell_background(cell, hex_color): cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>'))
def add_hyperlink(paragraph, url, text):
    part = paragraph.part; r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink'); hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r'); rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color'); color.set(qn('w:val'), '0000FF'); rPr.append(color)
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '14'); rPr.append(sz)
    szCs = OxmlElement('w:szCs'); szCs.set(qn('w:val'), '14'); rPr.append(szCs)
    new_run.append(rPr); text_element = OxmlElement('w:t'); text_element.text = text; new_run.append(text_element)
    hyperlink.append(new_run); paragraph._p.append(hyperlink); return hyperlink
def add_heading(doc, text, level=1): h = doc.add_heading(text, level=level); h.runs[0].font.name = "Arial"; return h
def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph(text); r = p.runs[0] if p.runs else p.add_run(text)
    r.font.name = "Arial"; r.font.size = Pt(11)
    if bold: r.bold = True
    return p
def add_caption(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(8)
    r.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p
def add_table_from_df(doc, df_t):
    table = doc.add_table(rows=1, cols=len(df_t.columns)); table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df_t.columns):
        hdr_cells[i].text = str(col); hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(7) # Všechny tabulky font 7
        set_cell_background(hdr_cells[i], "2E74B5")
    for _, row in df_t.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val) if pd.notna(val) else "–"
            row_cells[i].paragraphs[0].runs[0].font.size = Pt(7) # Všechny tabulky font 7
    return table

doc = Document(); doc.styles["Normal"].font.name = "Arial"; doc.styles["Normal"].font.size = Pt(11)

# TITULNÍ STRANA
doc.add_heading("ANALÝZA TRHU NEMOVITOSTÍ", 0); doc.add_heading("Stanovení jednotkové obvyklé ceny pozemku – ETALON JC", 1); doc.add_paragraph("")
for label, val in [("Věc:", f"Pozemky k.ú. {OCE_KU}, {OCE_OKRES}"), ("Typ pozemku:", f"Stavební pozemek (UP={OCE_UP})"), ("Výměra:", f"{OCE_VYMERA:,.0f} m²"), ("Datum ocenění:", OCE_DATUM.strftime("%d. %m. %Y"))]:
    p = doc.add_paragraph(); r1 = p.add_run(f"{label}  "); r1.bold = True; r1.font.name = "Arial"; r2 = p.add_run(val); r2.font.name = "Arial"
doc.add_page_break()

# KAP 1: ÚVOD
add_heading(doc, "1. Úvod a předmět analýzy", 1)
add_paragraph(doc, f"Předmětem této analýzy trhu je stanovení jednotkové obvyklé ceny (dále jen „JC\") pozemku situovaného v katastrálním území {OCE_KU}, okres {OCE_OKRES}, ke dni {OCE_DATUM.strftime('%d. %m. %Y')}.\n\nOceňovaný pozemek má celkovou výměru {OCE_VYMERA:,.0f} m² a dle aktuálně platného územního plánu náleží do funkční plochy označené kódem „{OCE_UP}\". Tvar pozemku byl exaktně vyhodnocen s Polsby-Popper indexem na hodnotu {OCE_PP_INDEX}.\n\nCílem analýzy je prostřednictvím transparentního statistického modelu (tzv. White-Box) dle IVS identifikovat srovnatelné tržní transakce, matematicky odvodit korekční koeficienty zohledňující parametry nemovitosti a stanovit tzv. ETALON JC – referenční jednotkovou cenu průměrného pozemku.")

# KAP 2: METODIKA A DATA
add_heading(doc, "2. Metodika a zdrojová data", 1)
add_paragraph(doc, "Základem analýzy je databáze tržních transakcí s pozemky z katastru nemovitostí.")
if "prostorova_mapa" in GRAPH_FILES:
    doc.add_picture(GRAPH_FILES["prostorova_mapa"], width=Inches(6.5))
    add_caption(doc, "Obrázek 1: Prostorové a cenové rozložení vzorku")

add_paragraph(doc, "\nSurový datový vzorek (Všechny parcely, řazeno chronologicky):", bold=True)
df_raw_table = df.dropna(subset=["JC", "vymera"]).copy(); df_raw_table.sort_values(by="datum_podani", ascending=False, inplace=True)
tab_cols = ["Datum podání", "Vklad", "K.Ú.", "Parc.č.", "ÚP", "Výměra [m²]", "JC [Kč/m²]"]
table = doc.add_table(rows=1, cols=len(tab_cols)); table.style = "Table Grid"
hdr_cells = table.rows[0].cells
for i, name in enumerate(tab_cols):
    hdr_cells[i].text = name; hdr_cells[i].paragraphs[0].runs[0].font.bold = True; hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(7)
    set_cell_background(hdr_cells[i], "2E74B5")

def get_hex_color(val, vmin, vmax, cmap_name="RdYlGn_r"):
    if pd.isna(val): return "FFFFFF"
    norm_val = (val - vmin) / max(vmax - vmin, 1); rgba = plt.cm.get_cmap(cmap_name)(norm_val)
    return mcolors.to_hex(rgba).replace("#", "").upper()

vmin_jc, vmax_jc = df_raw_table["JC"].min(), df_raw_table["JC"].max()
vmin_vym, vmax_vym = df_raw_table["vymera"].min(), df_raw_table["vymera"].max()
prev_year, year_bg = None, "FFFFFF"
for _, row in df_raw_table.iterrows():
    r_cells = table.add_row().cells
    r_cells[0].text = str(row["datum_den"]); r_cells[1].text = str(row.get("cislo_vkladu", "–")); r_cells[2].text = str(row.get("ku_nazev", "–"))
    cell_parc = r_cells[3]; cell_parc.text = ""; paragraph = cell_parc.paragraphs[0]; paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d_parc = str(row.get("parcel_number", "–"))
    if pd.notna(row.get("ruian_parcela_id")): add_hyperlink(paragraph, f"https://nahlizenidokn.cuzk.gov.cz/ZobrazObjekt.aspx?&typ=parcela&id={str(int(row.get('ruian_parcela_id')))}", d_parc)
    else: paragraph.text = d_parc
    r_cells[4].text = str(row.get("UP", "–")); r_cells[5].text = f"{row.get('vymera', 0):,.0f}"; r_cells[6].text = f"{row.get('JC', 0):,.0f}"
    for cell in r_cells: cell.paragraphs[0].runs[0].font.size = Pt(7)
    curr_yr = str(row["datum_den"])[:4]
    if prev_year and curr_yr != prev_year: year_bg = "F2F2F2" if year_bg == "FFFFFF" else "FFFFFF"
    prev_year = curr_yr
    for i in range(5): set_cell_background(r_cells[i], year_bg)
    set_cell_background(r_cells[5], get_hex_color(row.get("vymera", 0), vmin_vym, vmax_vym, "Blues"))
    set_cell_background(r_cells[6], get_hex_color(row.get("JC", 0), vmin_jc, vmax_jc, "RdYlGn_r"))
doc.add_page_break()

# KAP 3: DETEKCE EXTRÉMŮ
add_heading(doc, "3. Zacílená detekce a zpracování extrémních hodnot", 1)
add_paragraph(doc, f"V rámci expertního nastavení statistického modelu byly pro tuto analýzu definovány regulační parametry:\n- Přísnost IQR filtru (Tukeyho násobek): {OUTLIER_IQR_MULT}\n- Tolerance pro velikost pozemku: Garantováno ponechání transakcí do {VYMERA_TOLERANCE_NASOBEK}x násobku oceňované výměry.\n- Absolutní minimální přípustná JC: {MIN_POVOLENA_JC:,.0f} Kč/m²")
add_paragraph(doc, f"\nTrh nemovitostí má tzv. multimodální rozdělení. Aby se zabránilo znehodnocení vzorku, využívá tento model inovativní metodu Cenové homogenizace (Target-Anchored IQR):\n\nA) Jednotková cena (JC): Všechny ceny v trhu jsou matematicky převedeny na cenovou hladinu oceňovaného územního plánu (ÚP = {OCE_UP}). Následný filtr bezpečně odřízne anomálie s pásmem upraveným na {lower_bound_jc_norm:,.0f} Kč/m² až {upper_bound_jc_norm:,.0f} Kč/m².\n\nB) Výměra pozemku: Kolem oceňované velikosti ({OCE_VYMERA:,.0f} m²) model dynamicky nastavil ochranné toleranční pásmo od {lower_bound_vym:,.0f} m² do {upper_bound_vym:,.0f} m².")
if "outliery" in GRAPH_FILES:
    doc.add_picture(GRAPH_FILES["outliery"], width=Inches(6.5))
    add_caption(doc, "Graf 1: Identifikace extrémů na homogenizovaných datech")
doc.add_page_break()

# KAP 4: POPISNÉ STATISTIKY
add_heading(doc, "4. Základní popisné statistiky (Očištěný vzorek)", 1)
add_paragraph(doc, f"Níže jsou uvedeny klíčové statistické ukazatele jednotkových cen pro čistý vzorek {len(df_clean)} transakcí před aplikací koeficientů.")
add_table_from_df(doc, pd.DataFrame([["Vážený průměr (Hrubý)", f"{stats_jc['průměr']:,.0f} Kč/m²"], ["Medián", f"{stats_jc['medián']:,.0f} Kč/m²"], ["Směrodatná odchylka", f"{stats_jc['std']:,.0f} Kč/m²"]], columns=["Statistický ukazatel", "Hodnota"]))
if "histogram_JC" in GRAPH_FILES:
    doc.add_paragraph("")
    doc.add_picture(GRAPH_FILES["histogram_JC"], width=Inches(6.5))
    add_caption(doc, "Graf 2: Histogram jednotkových cen (Zdravý trh)")
doc.add_page_break()

# KAP 5: WHITE-BOX MODELY
add_heading(doc, "5. Exaktní zdůvodnění a výpočet korekčních koeficientů (K1–K5)", 1)
add_paragraph(doc, f"Všechny koeficienty jsou bezpečnostně oříznuty na povolený interval [{K_CLIP_MIN}; {K_CLIP_MAX}].")
slovo_zmeny = "růst" if rocni_zmena_k1 > 0 else "pokles"

add_paragraph(doc, "K1 – Korekce na čas (Vývoj trhu)", bold=True)
add_paragraph(doc, f"Byla provedena log-lineární regrese v závislosti na čase. Z odhadnutého regresního sklonu (β = {slope_k1:.6f}) lze odvodit, že zkoumaný vzorek v čase vykazuje průměrný roční {slovo_zmeny} cen o {abs(rocni_zmena_k1):.2f} %. Koeficient K1 posunuje historickou cenu k datu ocenění vzorcem K1 = exp(β · t).")
if "casovy_vyvoj" in GRAPH_FILES:
    doc.add_picture(GRAPH_FILES["casovy_vyvoj"], width=Inches(6.5))
    add_caption(doc, "Graf 3: Časový vývoj jednotkových cen v tržním vzorku")

add_paragraph(doc, "\nK2 – Korekce na lokalitu", bold=True)
add_paragraph(doc, f"Tento koeficient nevyřazuje transakce z jiných katastrálních území (KÚ). Naopak, umožňuje jejich využití tím, že objektivně přepočítá jejich cenovou hladinu na hladinu oceňovaného KÚ. Počítá se jako poměr mediánu JC v oceňovaném KÚ ({OCE_KU}) a mediánu JC v KÚ, kde leží srovnávací transakce. (Vypočtený medián JC pro {OCE_KU} = {median_oce_ku:,.0f} Kč/m²).")

add_paragraph(doc, "\nK3 – Korekce na typ ÚP", bold=True)
add_paragraph(doc, f"Do čitatele vstupuje medián ceny pozemků s funkčním využitím {OCE_UP} a do jmenovatele medián JC plochy u porovnávané transakce. (Medián JC pro ÚP={OCE_UP} = {median_oce_up:,.0f} Kč/m²).")

if "JC_dle_UP" in GRAPH_FILES:
    doc.add_picture(GRAPH_FILES["JC_dle_UP"], width=Inches(6.5))
    add_caption(doc, "Graf 4: Porovnání mediánů JC dle územního plánu")
    
    # NOVÁ TABULKA K ÚP
    doc.add_paragraph("")
    add_paragraph(doc, "Tabulka 1: Statistika jednotkových cen dle typu ÚP")
    up_s_word = up_s.copy()
    up_s_word["median"] = up_s_word["median"].apply(lambda x: f"{x:,.0f} Kč/m²")
    up_s_word["mean"] = up_s_word["mean"].apply(lambda x: f"{x:,.0f} Kč/m²")
    up_s_word["std"] = up_s_word["std"].fillna(0).apply(lambda x: f"{x:,.0f} Kč/m²")
    up_s_word.columns = ["Územní plán (ÚP)", "Medián", "Průměr", "Směrodatná odchylka"]
    add_table_from_df(doc, up_s_word)

doc.add_page_break()

add_paragraph(doc, "K4 – Korekce na velikost (Diskont na velikost)", bold=True)
add_paragraph(doc, f"S rostoucí výměrou pozemku klesá jeho jednotková cena (klesající mezní užitek). Model využívá log-log (mocninnou) regresi s konstantní elasticitou. {'Záporná hodnota sklonu (elasticita β = ' + str(round(slope_k4, 4)) + ') tento diskont exaktně potvrzuje.' if slope_k4 < 0 else 'Skutečnost, že regresní sklon není záporný (β = ' + str(round(slope_k4, 4)) + '), naznačuje, že se sleva na velikost neprojevuje.'} Transakce jsou přepočteny na úroveň oceňované výměry {OCE_VYMERA:,.0f} m².")
if "JC_vs_vymera" in GRAPH_FILES:
    doc.add_picture(GRAPH_FILES["JC_vs_vymera"], width=Inches(6.5))
    add_caption(doc, "Graf 5: Závislost jednotkové ceny na výměře s proloženým log-log trendem")

if not pp_vd.empty: tvar_komentar = f"nad úrovní tržního mediánu (trh = {pp_vd.median():.2f}). Je kompaktnější, což vede k vyššímu koeficientu" if OCE_PP_INDEX > pp_vd.median() else f"pod úrovní tržního mediánu (trh = {pp_vd.median():.2f}). Vykazuje složitou geometrii, což vyžaduje cenovou penalizaci (K5 < 1)"
else: tvar_komentar = "nelze v tomto vzorku určit"

add_paragraph(doc, "\nK5 – Korekce na tvar pozemku (Polsby-Popper index)", bold=True)
add_paragraph(doc, f"Geometrická komplexita je objektivizována Polsby-Popper indexem (1,0 = dokonalý kruh). Oceňovaný pozemek má index {OCE_PP_INDEX}. Ve srovnání s trhem se nachází {tvar_komentar}.")
if "K5_tvar" in GRAPH_FILES:
    doc.add_picture(GRAPH_FILES["K5_tvar"], width=Inches(6.5))
    add_caption(doc, "Graf 6: Rozložení tvarového indexu v tržním vzorku")
doc.add_page_break()

# KAP 6: PŘEHLED KOEFICIENTŮ
add_heading(doc, "6. Přehled průměrných uplatněných koeficientů", 1)
mean_k1, mean_k2, mean_k3, mean_k4, mean_k5 = df_clean['K1'].mean(), df_clean['K2'].mean(), df_clean['K3'].mean(), df_clean['K4'].mean(), df_clean['K5'].mean()
mean_k_celkem = (df_clean['K1']*df_clean['K2']*df_clean['K3']*df_clean['K4']*df_clean['K5']).mean()
add_paragraph(doc, f"Pro absolutní transparentnost uvádí tabulka průměrné hodnoty koeficientů uplatněné na čistý vzorek {len(df_clean)} transakcí. Odhaluje, co primárně srazilo nebo navýšilo konečnou jednotkovou cenu.")
add_table_from_df(doc, pd.DataFrame([
    ["K1 - Korekce na čas", f"{mean_k1:.3f}"], ["K2 - Korekce na lokalitu", f"{mean_k2:.3f}"],
    ["K3 - Korekce na typ ÚP", f"{mean_k3:.3f}"], ["K4 - Korekce na velikost", f"{mean_k4:.3f}"],
    ["K5 - Korekce na tvar", f"{mean_k5:.3f}"], ["Celkový průměrný násobitel na transakce (K_celk)", f"{mean_k_celkem:.3f}"]
], columns=["Označení koeficientu", "Průměrná aplikovaná hodnota na tržní vzorek"]))
varovani_text = ""
if mean_k4 <= K_CLIP_MIN + 0.05 or mean_k4 >= K_CLIP_MAX - 0.05: varovani_text += f"\nUPOZORNĚNÍ: K4 (Velikost) dosahuje extrémních limitů kvůli propastnému rozdílu mezi oceňovanou výměrou a běžnými obchody trhu. Model aplikoval maximální povolený ořez ({K_CLIP_MIN} - {K_CLIP_MAX})."
if mean_k5 <= K_CLIP_MIN + 0.05 or mean_k5 >= K_CLIP_MAX - 0.05: varovani_text += f"\nUPOZORNĚNÍ: K5 (Tvar) dosahuje extrémních limitů. Pozemek má mimořádně nevýhodný tvar (PP = {OCE_PP_INDEX}), model uplatnil maximální povolenou penalizaci."
if varovani_text: add_paragraph(doc, varovani_text.strip(), bold=True)

# KAP 7: ETALON
add_heading(doc, "7. Stanovení ETALONU JC", 1)
add_paragraph(doc, f"Upravená JC pro každou transakci vznikla součinem její původní JC a všech vypočtených koeficientů. Na základě porovnávací metody byl z čistého vzorku {len(df_clean)} transakcí odvozen ETALON JC k datu {OCE_DATUM.strftime('%d. %m. %Y')} následovně:")
add_table_from_df(doc, pd.DataFrame([["ETALON JC – vážený průměr", f"{etalon_vazeny:,.0f} Kč/m²"], ["ETALON JC – medián", f"{etalon_median:,.0f} Kč/m²"], ["Směrodatná odchylka", f"{etalon_std:,.0f} Kč/m²"]], columns=["Statistický ukazatel", "Výsledná hodnota"]))
doc.add_page_break()

# KAP 8: CITLIVOSTNÍ ANALÝZA
add_heading(doc, "8. Citlivostní analýza modelu", 1)
add_paragraph(doc, f"Analýza testuje robustnost ETALONU vůči izolovaným změnám hlavních statistických parametrů pro K1–K5 o ±{int(SENSITIVITY_DELTA*100)} %.")
if "tornado" in GRAPH_FILES:
    doc.add_picture(GRAPH_FILES["tornado"], width=Inches(6.5))
    add_caption(doc, "Graf 7: Tornádový graf citlivostní analýzy")
doc.add_page_break()

# KAP 9: CÍLENÁ ANALÝZA ÚP
add_heading(doc, f"9. Cílená analýza podmnožiny: ÚP {OCE_UP}", 1)
if has_sub_up:
    add_paragraph(doc, f"Z celkového čistého vzorku ({len(df_clean)} vkladů) vyhovuje této klasifikaci přesně {len(df_sub_up)} transakcí. Níže je uvedeno statistické srovnání této podmnožiny s celkovým trhem.")
    add_table_from_df(doc, pd.DataFrame([["Vážený průměr", f"{stats_sub_jc['průměr']:,.0f} Kč/m²", f"{stats_jc['průměr']:,.0f} Kč/m²"], ["Medián", f"{stats_sub_jc['medián']:,.0f} Kč/m²", f"{stats_jc['medián']:,.0f} Kč/m²"]], columns=["Ukazatel", f"Podmnožina ({OCE_UP})", "Celý trh"]))
    if "podmnozina_up" in GRAPH_FILES:
        doc.add_paragraph("")
        doc.add_picture(GRAPH_FILES["podmnozina_up"], width=Inches(6.5))
        add_caption(doc, "Graf 8: Boxplot podmnožiny")
else: add_paragraph(doc, f"Pro typ územního plánu {OCE_UP} není ve vzorku dostatečný počet transakcí pro izolovanou statistiku.")
doc.add_page_break()

# KAP 10: HEATMAPA
if "heatmap" in GRAPH_FILES:
    add_heading(doc, "10. Prostorové rozložení cenové hladiny v území", 1)
    add_paragraph(doc, "Pro vizuální ověření lokalizačního vlivu na cenu byla vypracována teplotní mapa (Heatmapa). Zelené odstíny reprezentují území s nižší cenovou hladinou, zatímco žluté až červené oblasti signalizují prémiové ceny. Tento model demonstruje reprezentativnost stanovených průměrů pro okolí oceňované nemovitosti (modrý trojúhelník).")
    doc.add_picture(GRAPH_FILES["heatmap"], width=Inches(6.5))
    add_caption(doc, "Obrázek 2: Cenová heatmapa")

with pd.ExcelWriter(OUT_XLSX, engine="openpyxl") as writer:
    df_clean.to_excel(writer, sheet_name="Data_Očištěná", index=False)
    df_outliers.to_excel(writer, sheet_name="Extrémy", index=False)

doc.save(OUT_DOCX)
print(f"[14] Uložen kompletní Word dokument: {OUT_DOCX}")

print("\n" + "=" * 70)
print("ANALÝZA DOKONČENA")
print(f"  ETALON JC (vážený průměr) = {etalon_vazeny:,.0f} Kč/m²")
print(f"  Celkový průměrný koeficient aplikovaný na trh = {mean_k_celkem:.3f}")
print("=" * 70)