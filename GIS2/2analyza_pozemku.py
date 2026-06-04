# =============================================================================
# ANALÝZA TRHU POZEMKŮ – POROVNÁVACÍ METODA (IVS / Market Approach)
# Oceňovaná nemovitost: Všenory, Praha-západ, UP=BI, 3 048 m²
# Datum ocenění: 2026-05-29
# Autor skriptu: generováno automaticky pro znalecký posudek č. 040742-2026
# =============================================================================
#
# Požadované knihovny (nainstalujte před spuštěním):
#   pip install pandas openpyxl matplotlib seaborn scipy numpy folium shapely python-docx geopandas contextily
#
# Výstupní soubory se uloží do stejné složky jako tento skript.
# =============================================================================

import os
import warnings
import re
import numpy as np
import pandas as pd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
from matplotlib.colors import Normalize
from matplotlib.cm import ScalarMappable
import seaborn as sns
from scipy import stats
import scipy.interpolate as interp
from shapely.geometry import Polygon
import folium
from folium.plugins import MarkerCluster
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from openpyxl import load_workbook
from openpyxl.styles import Font, PatternFill, Alignment, numbers
from openpyxl.utils import get_column_letter
import datetime

warnings.filterwarnings("ignore")

# =============================================================================
# 0. KONFIGURACE – všechny vstupní parametry oceňované nemovitosti
# =============================================================================

# Cesta ke zdrojovému Excel souboru (upravte pokud spouštíte z jiné složky)
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
INPUT_FILE = os.path.join(SCRIPT_DIR, "AI.xlsx")
OUT_XLSX   = os.path.join(SCRIPT_DIR, "Vystup_Analyza.xlsx")
OUT_MAP    = os.path.join(SCRIPT_DIR, "mapa_pozemky.html")
OUT_DOCX   = os.path.join(SCRIPT_DIR, "Znalecka_Zprava.docx")

# Oceňovaná nemovitost
OCE_DATUM        = pd.Timestamp("2026-06-02")
OCE_KU           = "Smíchov"          # katastrální území
OCE_OKRES        = "Hlavní město Praha" # okres pro dynamické texty
OCE_UP           = "LR"               # funkční využití dle ÚP
OCE_VYMERA       = 138332.0           # m² – celková výměra pozemků
OCE_PP_INDEX     = 0.10               # Polsby-Popper index tvaru (obdélník ≈ 0,7)
OCE_LAT          = 50.07132229220769  # GPS WGS84
OCE_LON          = 14.364680438635116
OCE_JTSK_X       = 747066.07          # S-JTSK
OCE_JTSK_Y       = 1044237.95

# Prahové hodnoty pro detekci outlierů (Standardní lineární IQR)
OUTLIER_IQR_MULT = 1.5               # standardní Tukeyho násobek pro IQR

# Citlivostní analýza – procentní odchylka koeficientů
SENSITIVITY_DELTA = 0.10             # ±10 %

# Datum reference pro výpočet „čas v dnech" (K1)
DATE_REF = OCE_DATUM                 # vztahujeme ke dni ocenění = 0

# =============================================================================
# 1. NAČTENÍ DAT
# =============================================================================

print("=" * 70)
print(f"ANALÝZA TRHU POZEMKŮ – {OCE_OKRES} / {OCE_KU}")
print("=" * 70)
print(f"\n[1] Načítám data ze souboru: {INPUT_FILE}")

df_raw = pd.read_excel(INPUT_FILE, sheet_name="data", header=0)
print(f"    Celkem řádků (surová data): {len(df_raw)}")
print(f"    Sloupce: {list(df_raw.columns)}")

# =============================================================================
# 2. PŘÍPRAVA DAT
# =============================================================================

print("\n[2] Čistím a připravuji data …")

df = df_raw.copy()

# Standardizace názvů sloupců (odstraníme mezery)
df.columns = df.columns.str.strip()

# Převod data
df["datum_podani"] = pd.to_datetime(df["datum_podani"], errors="coerce")
df["datum_den"]    = df["datum_podani"].dt.date

# Jednotková cena
if "JC_calc" in df.columns:
    df["JC"] = pd.to_numeric(df["JC_calc"], errors="coerce")
elif "#JC [Kč/m2]" in df.columns:
    df["JC"] = pd.to_numeric(df["#JC [Kč/m2]"], errors="coerce")
else:
    df["JC"] = pd.to_numeric(df["cenovy_udaj"], errors="coerce") / pd.to_numeric(df["#CELKOVA_VYMERA"], errors="coerce")

df["vymera"]       = pd.to_numeric(df["#CELKOVA_VYMERA"], errors="coerce")
df["cena_celkem"]  = pd.to_numeric(df["cenovy_udaj"],     errors="coerce")
df["cas_v_dnech"] = (DATE_REF - df["datum_podani"]).dt.days

if "UP" in df.columns:
    df["UP"] = df["UP"].astype(str).str.strip()

# =============================================================================
# 3. DEDUPLIKACE NA ÚROVEŇ cislo_vkladu
# =============================================================================

print("[3] Deduplicita na úroveň vkladu (cislo_vkladu) …")

def dominant_up(sub):
    if sub["UP"].nunique() == 1:
        return sub["UP"].iloc[0]
    vymery = sub.groupby("UP")["vymera"].sum()
    total  = vymery.sum()
    best   = vymery.idxmax()
    return best if (vymery[best] / total) > 0.5 else "UP smíšené"

agg_dict = {
    "datum_podani":  "first",
    "cas_v_dnech":   "first",
    "JC":            "first",       
    "vymera":        "first",       
    "cena_celkem":   "first",
    "ku_nazev":      "first",
    "refPoint_lat":  "first",
    "refPoint_lon":  "first",
    "refPoint_x":    "first",
    "refPoint_y":    "first",
    "geometry_posList": "first",
}
agg_dict = {k: v for k, v in agg_dict.items() if k in df.columns}

df_dedup = df.groupby("cislo_vkladu").agg(agg_dict).reset_index()

up_per_vklad = df.groupby("cislo_vkladu").apply(dominant_up).reset_index()
up_per_vklad.columns = ["cislo_vkladu", "UP_dom"]
df_dedup = df_dedup.merge(up_per_vklad, on="cislo_vkladu", how="left")
df_dedup.rename(columns={"UP_dom": "UP"}, inplace=True)

df_dedup.dropna(subset=["JC", "vymera", "datum_podani"], inplace=True)
df_dedup = df_dedup[df_dedup["JC"] > 0].reset_index(drop=True)

print(f"    Unikátních vkladů po deduplikaci: {len(df_dedup)}")

# =============================================================================
# 4. POLSBY-POPPER INDEX (K5 tvar)
# =============================================================================

print("[4] Výpočet Polsby-Popper indexu z geometry_posList …")

def parse_geometry(geom_str):
    try:
        if pd.isna(geom_str) or str(geom_str).strip() == "": return None
        nums = list(map(float, str(geom_str).split()))
        if len(nums) < 6 or len(nums) % 2 != 0: return None
        coords = [(nums[i], nums[i+1]) for i in range(0, len(nums), 2)]
        poly = Polygon(coords)
        return poly if poly.is_valid else poly.buffer(0)
    except Exception:
        return None

if "geometry_posList" in df_dedup.columns:
    df_dedup["polygon"]   = df_dedup["geometry_posList"].apply(parse_geometry)
    df_dedup["pp_area"]   = df_dedup["polygon"].apply(lambda p: p.area if p is not None else np.nan)
    df_dedup["pp_perim"]  = df_dedup["polygon"].apply(lambda p: p.length if p is not None else np.nan)
    df_dedup["pp_index"]  = df_dedup.apply(
        lambda r: (4 * np.pi * r["pp_area"]) / (r["pp_perim"] ** 2)
        if (pd.notna(r["pp_perim"]) and r["pp_perim"] > 0) else np.nan, axis=1
    )
    n_pp = df_dedup["pp_index"].notna().sum()
    print(f"    PP index spočítán pro {n_pp} z {len(df_dedup)} vkladů.")
else:
    df_dedup["pp_index"] = np.nan
    print("    Sloupec geometry_posList nenalezen – PP index nastaven na NaN.")

# =============================================================================
# 5. DETEKCE OUTLIERŮ (Nová Standardní Lineární IQR Metoda s limitem)
# =============================================================================

print("[5] Detekce outlierů (Standardní lineární IQR s oříznutím) …")

Q1_jc = df_dedup["JC"].quantile(0.25)
Q3_jc = df_dedup["JC"].quantile(0.75)
IQR_jc = Q3_jc - Q1_jc

lower_bound = Q1_jc - (OUTLIER_IQR_MULT * IQR_jc)
upper_bound = Q3_jc + (OUTLIER_IQR_MULT * IQR_jc)

# Záchranná brzda pro extrémně levné nemovitosti - nesmí být záporné, ořízneme na 5. percentil
if lower_bound <= 0:
    lower_bound = max(df_dedup["JC"].quantile(0.05), 1.0)

df_dedup["outlier"] = (df_dedup["JC"] < lower_bound) | (df_dedup["JC"] > upper_bound)
df_outliers = df_dedup[df_dedup["outlier"]].copy()
df_clean    = df_dedup[~df_dedup["outlier"]].copy().reset_index(drop=True)

print(f"    Hranice outlierů: {lower_bound:,.0f} – {upper_bound:,.0f} Kč/m²")
print(f"    Detekováno outlierů: {len(df_outliers)}, v čistém vzorku zůstalo: {len(df_clean)}")

# =============================================================================
# 6. POPISNÉ STATISTIKY (Včetně analýzy podmnožiny ÚP)
# =============================================================================

print("[6] Popisné statistiky a analýza podmnožiny …")

def weighted_stats(series, weights=None):
    s = series.dropna()
    w = weights.loc[s.index].fillna(0) if weights is not None else None
    return {
        "počet":    len(s),
        "průměr":   np.average(s, weights=w) if w is not None and w.sum() > 0 else s.mean(),
        "min":      s.min(),
        "P5":       s.quantile(0.05),
        "Q1":       s.quantile(0.25),
        "medián":   s.median(),
        "Q3":       s.quantile(0.75),
        "P95":      s.quantile(0.95),
        "max":      s.max(),
        "std":      s.std(),
    }

stats_cas   = weighted_stats(df_clean["cas_v_dnech"])
stats_vym   = weighted_stats(df_clean["vymera"],   df_clean["vymera"])
stats_jc    = weighted_stats(df_clean["JC"],        df_clean["vymera"])

print(f"\n    Celý vzorek JC [Kč/m²] – průměr: {stats_jc['průměr']:,.0f}  |  medián: {stats_jc['medián']:,.0f}")

# Cílená statistika pro konkrétní ÚP
df_sub_up = df_clean[df_clean["UP"] == OCE_UP]
has_sub_up = len(df_sub_up) >= 3
if has_sub_up:
    stats_sub_jc = weighted_stats(df_sub_up["JC"], df_sub_up["vymera"])
    print(f"    Podmnožina UP={OCE_UP} detekována ({len(df_sub_up)} záznamů): průměr: {stats_sub_jc['průměr']:,.0f} | medián: {stats_sub_jc['medián']:,.0f}")
else:
    print(f"    Podmnožina UP={OCE_UP} má příliš málo dat pro izolovanou statistiku.")

# =============================================================================
# 7. VÝPOČET KOREKČNÍCH KOEFICIENTŮ K1–K5
# =============================================================================

print("[7] Výpočet korekčních koeficientů K1–K5 …")

# K1
df_reg_k1 = df_clean[df_clean["JC"] > 0].copy()
df_reg_k1["log_JC"] = np.log(df_reg_k1["JC"])
if len(df_reg_k1) > 1 and df_reg_k1["cas_v_dnech"].nunique() > 1:
    slope_k1, intercept_k1, r_k1, p_k1, _ = stats.linregress(df_reg_k1["cas_v_dnech"], df_reg_k1["log_JC"])
else: slope_k1, r_k1 = 0.0, 0.0
if pd.isna(slope_k1): slope_k1, r_k1 = 0.0, 0.0

df_clean["K1"] = np.exp(slope_k1 * df_clean["cas_v_dnech"]).clip(0.50, 2.00)
print(f"    K1 – regresní koeficient (slope): {slope_k1:.6f} | R²={r_k1**2:.3f}")

# K2
median_jc_per_ku = df_clean.groupby("ku_nazev")["JC"].median()
median_oce_ku    = median_jc_per_ku.get(OCE_KU, df_clean["JC"].median())
df_clean["K2"] = (median_oce_ku / df_clean["ku_nazev"].map(median_jc_per_ku)).clip(0.50, 2.00)
print(f"    K2 – medián JC [{OCE_KU}]: {median_oce_ku:,.0f} Kč/m²")

# K3
median_jc_per_up = df_clean.groupby("UP")["JC"].median()
median_oce_up    = median_jc_per_up.get(OCE_UP, df_clean["JC"].median())
df_clean["K3"] = (median_oce_up / df_clean["UP"].map(median_jc_per_up)).clip(0.50, 2.00)
print(f"    K3 – medián JC [UP={OCE_UP}]: {median_oce_up:,.0f} Kč/m²")

# K4
df_reg_k4 = df_clean[df_clean["JC"] > 0].copy()
df_reg_k4["log_JC"]   = np.log(df_reg_k4["JC"])
df_reg_k4["log_vym"]  = np.log(df_reg_k4["vymera"])
if len(df_reg_k4) > 1 and df_reg_k4["log_vym"].nunique() > 1:
    slope_k4, intercept_k4, r_k4, p_k4, _ = stats.linregress(df_reg_k4["log_vym"], df_reg_k4["log_JC"])
else: slope_k4, r_k4 = 0.0, 0.0
if pd.isna(slope_k4): slope_k4, r_k4 = 0.0, 0.0

df_clean["K4"] = ((OCE_VYMERA / df_clean["vymera"]) ** slope_k4).clip(0.50, 2.00)
print(f"    K4 – log-log slope: {slope_k4:.4f} | R²={r_k4**2:.3f}")

# K5
pp_valid = df_clean["pp_index"].notna() & (df_clean["pp_index"] > 0)
df_clean.loc[pp_valid,  "K5"] = (OCE_PP_INDEX / df_clean.loc[pp_valid, "pp_index"]).clip(0.50, 2.00)
df_clean.loc[~pp_valid, "K5"] = 1.0   
print(f"    K5 – PP index oceňované: {OCE_PP_INDEX} | K5 u {pp_valid.sum()} vkladů")


# =============================================================================
# 8. UPRAVENÁ JEDNOTKOVÁ CENA (porovnávací metoda – krok 1)
# =============================================================================

print("[8] Výpočet upravené JC …")
df_clean["JC_upravena"] = (
    df_clean["JC"] * df_clean["K1"] * df_clean["K2"]
    * df_clean["K3"] * df_clean["K4"] * df_clean["K5"]
)

# =============================================================================
# 9. ETALON JC (krok 2)
# =============================================================================

print("[9] Výpočet ETALONU JC …")

valid_mask = df_clean["JC_upravena"].notna() & df_clean["vymera"].notna()
if valid_mask.sum() > 0:
    etalon_vazeny = np.average(df_clean.loc[valid_mask, "JC_upravena"], weights=df_clean.loc[valid_mask, "vymera"])
else:
    etalon_vazeny = df_clean["JC"].median()
    if pd.isna(etalon_vazeny): etalon_vazeny = 1000.0

etalon_median  = df_clean["JC_upravena"].median()
if pd.isna(etalon_median): etalon_median = etalon_vazeny

etalon_std     = df_clean["JC_upravena"].std()
etalon_p25     = df_clean["JC_upravena"].quantile(0.25)
etalon_p75     = df_clean["JC_upravena"].quantile(0.75)

etalon_std = etalon_std if pd.notna(etalon_std) else 0.0
etalon_p25 = etalon_p25 if pd.notna(etalon_p25) else etalon_median
etalon_p75 = etalon_p75 if pd.notna(etalon_p75) else etalon_median

print(f"\n    *** ETALON JC ***")
print(f"    Vážený průměr : {etalon_vazeny:,.0f} Kč/m²")
print(f"    Medián        : {etalon_median:,.0f} Kč/m²")

# =============================================================================
# 10. CITLIVOSTNÍ ANALÝZA 
# =============================================================================

print("[10] Citlivostní analýza (testování parametrů) …")

def prepocitej_etalon_citlivost(df_base, zmeny_params):
    df_tmp = df_base.copy()
    
    s_k1 = slope_k1 * (1 + zmeny_params.get("slope_k1", 0))
    m_k2 = median_oce_ku * (1 + zmeny_params.get("median_k2", 0))
    m_k3 = median_oce_up * (1 + zmeny_params.get("median_k3", 0))
    s_k4 = slope_k4 * (1 + zmeny_params.get("slope_k4", 0))
    
    k1 = np.exp(s_k1 * df_tmp["cas_v_dnech"]).clip(0.50, 2.00)
    median_jc_per_ku = df_tmp.groupby("ku_nazev")["JC"].median()
    k2 = (m_k2 / df_tmp["ku_nazev"].map(median_jc_per_ku)).clip(0.50, 2.00)
    median_jc_per_up = df_tmp.groupby("UP")["JC"].median()
    k3 = (m_k3 / df_tmp["UP"].map(median_jc_per_up)).clip(0.50, 2.00)
    k4 = ((OCE_VYMERA / df_tmp["vymera"]) ** s_k4).clip(0.50, 2.00)
    
    p_k5 = OCE_PP_INDEX * (1 + zmeny_params.get("pp_k5", 0))
    pp_valid = df_tmp["pp_index"].notna() & (df_tmp["pp_index"] > 0)
    k5 = pd.Series(1.0, index=df_tmp.index)
    k5.loc[pp_valid] = (p_k5 / df_tmp.loc[pp_valid, "pp_index"]).clip(0.50, 2.00)
    
    jc_uprenava = df_tmp["JC"] * k1 * k2 * k3 * k4 * k5
    
    v_mask = jc_uprenava.notna() & df_tmp["vymera"].notna()
    if v_mask.sum() > 0: return np.average(jc_uprenava[v_mask], weights=df_tmp.loc[v_mask, "vymera"])
    return df_tmp["JC"].median()

param_map = {"K1": "slope_k1", "K2": "median_k2", "K3": "median_k3", "K4": "slope_k4", "K5": "pp_k5"}
sensitivity_results = {}
for k_label, p_name in param_map.items():
    vaz_up = prepocitej_etalon_citlivost(df_clean, {p_name: +SENSITIVITY_DELTA})
    vaz_dn = prepocitej_etalon_citlivost(df_clean, {p_name: -SENSITIVITY_DELTA})
    if pd.isna(vaz_up): vaz_up = etalon_vazeny
    if pd.isna(vaz_dn): vaz_dn = etalon_vazeny
    
    sensitivity_results[k_label] = {
        "base_vazeny": etalon_vazeny,
        "up_vazeny":   vaz_up,
        "dn_vazeny":   vaz_dn,
        "rozpeti":     abs(vaz_up - vaz_dn),
    }

tornado_order = sorted(sensitivity_results.keys(), key=lambda k: sensitivity_results[k]["rozpeti"], reverse=True)

# =============================================================================
# 11. GRAFY
# =============================================================================

print("[11] Generuji grafy …")

GRAPH_FILES = {}
sns.set_theme(style="whitegrid", font="Arial")

# ── 11.1 Histogram JC ──────────────────────────────────────────────────────
fig, ax = plt.subplots(figsize=(10, 5))
bins = np.logspace(np.log10(df_dedup["JC"].min() if df_dedup["JC"].min()>0 else 1), np.log10(df_dedup["JC"].max()), 20)
ax.hist(df_dedup["JC"], bins=bins, color="#4472C4", edgecolor="white", alpha=0.85)
ax.set_xscale('log')

ax.axvline(stats_jc["medián"],   color="red",   linestyle="--", linewidth=1.5, label=f'Medián {stats_jc["medián"]:,.0f}')
ax.axvline(stats_jc["průměr"],   color="green", linestyle="--", linewidth=1.5, label=f'Vážený průměr {stats_jc["průměr"]:,.0f}')
ax.axvline(stats_jc["Q1"],       color="orange",linestyle=":",  linewidth=1.2, label=f'Q1 {stats_jc["Q1"]:,.0f}')
ax.axvline(stats_jc["Q3"],       color="orange",linestyle=":",  linewidth=1.2, label=f'Q3 {stats_jc["Q3"]:,.0f}')
ax.axvline(lower_bound,          color="darkred", linestyle="-.", linewidth=1.5, label=f'Lin-IQR Dolní mez ({lower_bound:,.0f})')
ax.axvline(upper_bound,          color="darkred", linestyle="-.", linewidth=1.5, label=f'Lin-IQR Horní mez ({upper_bound:,.0f})')

ax.set_ylabel("Počet vkladů (frekvence)")
ax.set_xlabel("Jednotková cena [Kč/m²] (Logaritmické měřítko)")
ax.set_title("Histogram jednotkových cen pozemků")
ax.xaxis.set_major_formatter(mticker.FuncFormatter(lambda x, _: f"{x:,.0f}"))
ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda x, _: f"{x:,.0f}"))
ax.xaxis.set_minor_formatter(mticker.NullFormatter())
ax.tick_params(axis='x', rotation=45)
ax.legend(fontsize=9, loc="upper right")
plt.tight_layout()
p = os.path.join(SCRIPT_DIR, "graf_histogram_JC.png")
fig.savefig(p, dpi=150); plt.close(fig)
GRAPH_FILES["histogram_JC"] = p

# ── 11.2 Časový vývoj JC ──────────────────────────────────────────────────
fig, ax = plt.subplots(figsize=(12, 5))
norm = Normalize(vmin=df_clean["JC"].min(), vmax=df_clean["JC"].max())
cmap = plt.cm.RdYlGn
sc   = ax.scatter(df_clean["datum_podani"], df_clean["JC"],
                  c=df_clean["JC"], cmap=cmap, norm=norm, s=60, zorder=3, alpha=0.85)

x_num = df_clean["datum_podani"].map(pd.Timestamp.toordinal)
if len(x_num) > 1 and x_num.nunique() > 1:
    m, b, *_ = stats.linregress(x_num, df_clean["JC"])
else:
    m, b = 0.0, df_clean["JC"].median()
if pd.isna(m): m = 0.0
if pd.isna(b): b = df_clean["JC"].median()

x_range = pd.date_range(df_clean["datum_podani"].min(), df_clean["datum_podani"].max(), periods=100)
ax.plot(x_range, m * x_range.map(pd.Timestamp.toordinal) + b, color="navy", linewidth=2, label="Lineární trend")

for val, lbl, c, ls, lw in [
    (stats_jc["medián"], "Medián", "red", "--", 1.5), (stats_jc["průměr"], "Průměr", "green", "--", 1.5),
    (stats_jc["Q1"], "Q1", "orange", ":", 1.2), (stats_jc["Q3"], "Q3", "orange", ":", 1.2),
    (stats_jc["P5"], "P5", "gray", "-.", 1.0), (stats_jc["P95"], "P95", "gray", "-.", 1.0),
]:
    if pd.notna(val):
        ax.axhline(val, linestyle=ls, linewidth=lw, color=c, alpha=0.8, label=f"{lbl} {val:,.0f}")

plt.colorbar(sc, ax=ax, label="JC [Kč/m²]")
ax.set_xlabel("Datum podání")
ax.set_ylabel("Jednotková cena [Kč/m²]")
ax.set_title("Časový vývoj jednotkových cen pozemků")
ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda x, _: f"{x:,.0f}"))
ax.legend(fontsize=8, loc="upper left")
plt.tight_layout()
p = os.path.join(SCRIPT_DIR, "graf_casovy_vyvoj.png")
fig.savefig(p, dpi=150); plt.close(fig)
GRAPH_FILES["casovy_vyvoj"] = p

# ── 11.3 Analýza dle UP ───────────────────────────────────────────────────
up_stats = df_clean.groupby("UP")["JC"].agg(median="median", mean="mean", std="std", count="count").reset_index().sort_values("median", ascending=False)
fig, ax = plt.subplots(figsize=(10, 5))
x = np.arange(len(up_stats))
bars1 = ax.bar(x - 0.2, up_stats["median"], 0.4, label="Medián JC", color="#4472C4")
bars2 = ax.bar(x + 0.2, up_stats["mean"],   0.4, label="Průměr JC", color="#ED7D31", alpha=0.85)
std_vals = up_stats["std"].fillna(0)
ax.errorbar(x + 0.2, up_stats["mean"], yerr=std_vals, fmt="none", ecolor="black", capsize=4, linewidth=1)
ax.set_xticks(x)
ax.set_xticklabels(up_stats["UP"], rotation=45, ha="right")
ax.set_ylabel("Jednotková cena [Kč/m²]")
ax.set_title("Porovnání mediánu a průměru JC dle funkčního využití ÚP")
ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda x, _: f"{x:,.0f}"))
ax.legend()
for bar in bars1:
    val = bar.get_height()
    if pd.notna(val):
        ax.text(bar.get_x() + bar.get_width()/2, val + (val*0.02), f"{val:,.0f}", ha="center", va="bottom", fontsize=7)
plt.tight_layout()
p = os.path.join(SCRIPT_DIR, "graf_JC_dle_UP.png")
fig.savefig(p, dpi=150); plt.close(fig)
GRAPH_FILES["JC_dle_UP"] = p

# ── 11.4 Závislost JC na velikosti pozemku ───────────────────────────────
fig, ax = plt.subplots(figsize=(10, 5))
norm4 = Normalize(vmin=df_clean["JC"].min(), vmax=df_clean["JC"].max())
sc4   = ax.scatter(df_clean["vymera"], df_clean["JC"], c=df_clean["JC"], cmap="RdYlGn", norm=norm4, s=60, alpha=0.8)
vym_range = np.linspace(df_clean["vymera"].min(), df_clean["vymera"].max(), 200)
jc_trend  = np.exp(slope_k4 * np.log(vym_range) + intercept_k4) if pd.notna(slope_k4) and pd.notna(intercept_k4) else np.repeat(df_clean["JC"].median(), 200)
ax.plot(vym_range, jc_trend, color="navy", linewidth=2, label=f"Log-log trend (β={slope_k4:.3f})")
for val, lbl, c, ls, lw in [
    (stats_jc["medián"], "Medián", "red", "--", 1.5), (stats_jc["průměr"], "Průměr", "green", "--", 1.5),
    (stats_jc["Q1"], "Q1", "orange", ":", 1.2), (stats_jc["Q3"], "Q3", "orange", ":", 1.2),
]:
    if pd.notna(val):
        ax.axhline(val, linestyle=ls, linewidth=lw, color=c, alpha=0.8, label=f"{lbl} {val:,.0f}")
plt.colorbar(sc4, ax=ax, label="JC [Kč/m²]")
ax.set_xlabel("Celková výměra vkladu [m²]")
ax.set_ylabel("Jednotková cena [Kč/m²]")
ax.set_title("Závislost JC na velikosti pozemku (diskont na velikost)")
ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda x, _: f"{x:,.0f}"))
ax.xaxis.set_major_formatter(mticker.FuncFormatter(lambda x, _: f"{x:,.0f}"))
ax.legend(fontsize=8, loc="upper right")
plt.tight_layout()
p = os.path.join(SCRIPT_DIR, "graf_JC_vs_vymera.png")
fig.savefig(p, dpi=150); plt.close(fig)
GRAPH_FILES["JC_vs_vymera"] = p

# ── 11.5 Outliers – vizualizace ───────────────────────────────────────────
fig, ax = plt.subplots(figsize=(10, 5))
ax.scatter(df_clean["datum_podani"],    df_clean["JC"], color="#4472C4", s=40, label="Standardní transakce", alpha=0.7)
if len(df_outliers) > 0:
    ax.scatter(df_outliers["datum_podani"], df_outliers["JC"], color="red", s=80, marker="X", label="Outlier", zorder=5)
ax.axhline(lower_bound, color="orange", linestyle="--", linewidth=1.5, label=f"Dolní hranice ({lower_bound:,.0f})")
ax.axhline(upper_bound, color="orange", linestyle="--", linewidth=1.5, label=f"Horní hranice ({upper_bound:,.0f})")
ax.set_xlabel("Datum podání")
ax.set_ylabel("Jednotková cena [Kč/m²]")
ax.set_title("Identifikace extrémů v souboru dat (Lineární IQR metoda)")
ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda x, _: f"{x:,.0f}"))
ax.legend(fontsize=8)
plt.tight_layout()
p = os.path.join(SCRIPT_DIR, "graf_outliery.png")
fig.savefig(p, dpi=150); plt.close(fig)
GRAPH_FILES["outliery"] = p

# ── 11.6 Tornado chart ────────────────────────────────────────────────────
fig, ax = plt.subplots(figsize=(10, 5))
labels_map = {"K1": "K1 – Čas", "K2": "K2 – Lokalita", "K3": "K3 – Typ UP", "K4": "K4 – Velikost", "K5": "K5 – Tvar"}
colors_up, colors_dn = "#70AD47", "#FF0000"
min_x, max_x = float('inf'), float('-inf')

for i, k in enumerate(tornado_order):
    r  = sensitivity_results[k]
    if pd.isna(r["dn_vazeny"]) or pd.isna(r["up_vazeny"]): continue
    lo, hi = min(r["dn_vazeny"], r["up_vazeny"]), max(r["dn_vazeny"], r["up_vazeny"])
    min_x, max_x = min(min_x, lo), max(max_x, hi)
    
    ax.barh(i, hi - etalon_vazeny, left=etalon_vazeny, color=colors_up, alpha=0.75, height=0.5)
    ax.barh(i, lo - etalon_vazeny, left=etalon_vazeny, color=colors_dn, alpha=0.75, height=0.5)
    ax.text(hi + (hi*0.01), i, f"+10%: {hi:,.0f}", va="center", ha="left", fontsize=8)
    ax.text(lo - (lo*0.01), i, f"-10%: {lo:,.0f}", va="center", ha="right", fontsize=8)

if np.isinf(min_x) or np.isinf(max_x) or pd.isna(min_x) or pd.isna(max_x):
    bezpecny_etalon = etalon_vazeny if pd.notna(etalon_vazeny) else 1000.0
    min_x, max_x = bezpecny_etalon * 0.8, bezpecny_etalon * 1.2
if pd.isna(min_x): min_x = 800.0
if pd.isna(max_x): max_x = 1200.0

rozpeti_x = max(max_x - min_x, 100)
ax.set_xlim(min_x - (rozpeti_x * 0.25), max_x + (rozpeti_x * 0.25))
ax.axvline(etalon_vazeny, color="black", linewidth=1.5, linestyle="-")
ax.set_yticks(range(len(tornado_order)))
ax.set_yticklabels([labels_map[k] for k in tornado_order])
ax.set_xlabel("Etalon JC [Kč/m²]")
ax.set_title("Citlivostní analýza – Vliv vstupních parametrů (Tornádový graf)")
ax.xaxis.set_major_formatter(mticker.FuncFormatter(lambda x, _: f"{x:,.0f}"))
ax.invert_yaxis()
plt.tight_layout()
p = os.path.join(SCRIPT_DIR, "graf_tornado.png")
fig.savefig(p, dpi=150); plt.close(fig)
GRAPH_FILES["tornado"] = p

# ── 11.7 Cílená analýza podmnožiny (Boxplot) ──────────────────────────────
if has_sub_up:
    fig, ax = plt.subplots(figsize=(8, 5))
    df_boxplot = df_clean.copy()
    df_boxplot['Typ_Skupiny'] = np.where(df_boxplot['UP'] == OCE_UP, f'Oceňovaný ÚP ({OCE_UP})', 'Ostatní ÚP')
    
    # Použití boxplotu pro zobrazení rozptylu
    sns.boxplot(data=df_boxplot, x='Typ_Skupiny', y='JC', ax=ax, palette=['#FF9999', '#99CCFF'])
    ax.set_title(f"Rozptyl a hladina jednotkových cen: ÚP {OCE_UP} vůči zbytku trhu")
    ax.set_ylabel("Jednotková cena [Kč/m²]")
    ax.set_xlabel("")
    ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda x, _: f"{x:,.0f}"))
    plt.tight_layout()
    p = os.path.join(SCRIPT_DIR, "graf_podmnozina_up.png")
    fig.savefig(p, dpi=150); plt.close(fig)
    GRAPH_FILES["podmnozina_up"] = p
    print(f"    Uložen Boxplot podmnožiny: {p}")

# ── 11.8 Statická prostorová mapa pro Word (WGS84 verze) ──────────────────
import geopandas as gpd
import contextily as ctx
from matplotlib.colors import Normalize
from matplotlib.lines import Line2D

LAT_COL = "refPoint_lat" 
LON_COL = "refPoint_lon" 

fig, ax = plt.subplots(figsize=(10, 8))
valid_gps = df_dedup.dropna(subset=[LAT_COL, LON_COL])

if not valid_gps.empty:
    gdf = gpd.GeoDataFrame(
        valid_gps, geometry=gpd.points_from_xy(valid_gps[LON_COL], valid_gps[LAT_COL]), crs="EPSG:4326"
    ).to_crs("EPSG:3857")
    
    gdf.plot(
        ax=ax, column="JC", cmap="RdYlGn_r", norm=Normalize(vmin=df_dedup["JC"].min(), vmax=df_dedup["JC"].max()),
        markersize=80, edgecolor="black", alpha=0.8, legend=True,
        legend_kwds={'label': "Jednotková cena [Kč/m²]", 'shrink': 0.6}, zorder=2
    )

    oce_gdf = gpd.GeoDataFrame(
        index=[0], crs="EPSG:4326", geometry=[gpd.points_from_xy([OCE_LON], [OCE_LAT])[0]]
    ).to_crs("EPSG:3857")
    
    ax.scatter(oce_gdf.geometry.x, oce_gdf.geometry.y, color="blue", marker="^", s=300, edgecolors="white", zorder=10, label="Oceňovaná nemovitost")
    ctx.add_basemap(ax, crs=gdf.crs.to_string(), source=ctx.providers.Esri.WorldImagery)
    legend_elements = [Line2D([0], [0], marker='^', color='w', label='Oceňovaná nemovitost', markerfacecolor='blue', markeredgecolor='white', markersize=15)]
    ax.legend(handles=legend_elements, loc="upper right")
    ax.set_title("Prostorové rozložení transakcí")
    ax.set_axis_off()
    plt.tight_layout()
    p = os.path.join(SCRIPT_DIR, "graf_prostorova_mapa.png")
    fig.savefig(p, dpi=150); plt.close(fig)
    GRAPH_FILES["prostorova_mapa"] = p

# ── 11.9 Prostorová interpolace cen (Heatmap) ──────────────────────────────
valid_gps_clean = df_clean.dropna(subset=[LAT_COL, LON_COL]) # Pro heatmapu použijeme čistá data bez outlierů!

if not valid_gps_clean.empty and len(valid_gps_clean) > 4:
    try:
        gdf_clean = gpd.GeoDataFrame(
            valid_gps_clean, geometry=gpd.points_from_xy(valid_gps_clean[LON_COL], valid_gps_clean[LAT_COL]), crs="EPSG:4326"
        ).to_crs("EPSG:3857")
        
        oce_gdf = gpd.GeoDataFrame(
            index=[0], crs="EPSG:4326", geometry=[gpd.points_from_xy([OCE_LON], [OCE_LAT])[0]]
        ).to_crs("EPSG:3857")

        x, y = gdf_clean.geometry.x.values, gdf_clean.geometry.y.values
        z = gdf_clean['JC'].values
        
        # Grid pro interpolaci s přesahem
        margin = 1500 # 1.5 km přesah
        xi = np.linspace(x.min() - margin, x.max() + margin, 300)
        yi = np.linspace(y.min() - margin, y.max() + margin, 300)
        XI, YI = np.meshgrid(xi, yi)
        
        # Interpolace (linear)
        ZI = interp.griddata((x, y), z, (XI, YI), method='linear')
        
        fig, ax = plt.subplots(figsize=(10, 8))
        
        # Vykreslení contour plot (Heatmapa) se zvýšenou sytostí a počtem hladin pro lepší zřetelnost
        contour = ax.contourf(XI, YI, ZI, levels=30, cmap="RdYlGn_r", alpha=0.85)
        # Přidání jemných hranic mezi hladinami
        ax.contour(XI, YI, ZI, levels=15, colors='black', alpha=0.2, linewidths=0.5)
        
        cbar = plt.colorbar(contour, ax=ax, fraction=0.046, pad=0.04)
        cbar.set_label("Cenová hladina JC [Kč/m²]")
        cbar.ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda val, pos: f"{val:,.0f}"))
        
        # Původní body s mírnou průhledností pro kontext
        gdf_clean.plot(ax=ax, color='black', markersize=15, alpha=0.4, label="Čisté transakce")
        ax.scatter(oce_gdf.geometry.x, oce_gdf.geometry.y, color="blue", marker="^", s=250, edgecolor="white", zorder=10, label="Oceňovaná nemovitost")
        
        ctx.add_basemap(ax, crs=gdf_clean.crs.to_string(), source=ctx.providers.Esri.WorldImagery)
        ax.legend(loc="upper right")
        ax.set_title("Prostorová interpolace cenové hladiny (Teplotní mapa)")
        ax.set_axis_off()
        plt.tight_layout()
        
        p = os.path.join(SCRIPT_DIR, "graf_heatmap.png")
        fig.savefig(p, dpi=150); plt.close(fig)
        GRAPH_FILES["heatmap"] = p
        print(f"    Uložena heatmapa: {p}")
    except Exception as e:
        print(f"    Upozornění: Nepodařilo se vygenerovat heatmapu: {e}")


# =============================================================================
# 12. INTERAKTIVNÍ MAPA (Folium)
# =============================================================================

print("[12] Generuji interaktivní mapu …")
import branca.colormap as cm

valid_coords = df_dedup.dropna(subset=["refPoint_lat", "refPoint_lon"])
if not valid_coords.empty:
    sw = [valid_coords["refPoint_lat"].min(), valid_coords["refPoint_lon"].min()]
    ne = [valid_coords["refPoint_lat"].max(), valid_coords["refPoint_lon"].max()]
    map_center = [(sw[0] + ne[0]) / 2, (sw[1] + ne[1]) / 2]
else:
    map_center = [OCE_LAT, OCE_LON]
    sw, ne = None, None

m = folium.Map(location=map_center, prefer_canvas=True)
folium.TileLayer("OpenStreetMap", name="OpenStreetMap", show=True).add_to(m)
folium.TileLayer(
    tiles="https://ags.cuzk.gov.cz/arcgis1/rest/services/ORTOFOTO_WM/MapServer/tile/{z}/{y}/{x}",
    name="ČÚZK Ortofoto", attr="© ČÚZK", overlay=False, control=True, max_zoom=20, min_zoom=6, show=False,
).add_to(m)

jc_min = df_dedup["JC"].min() if not pd.isna(df_dedup["JC"].min()) else 0
jc_max = df_dedup["JC"].max() if not pd.isna(df_dedup["JC"].max()) else 1000
if jc_min == jc_max: jc_max += 1

colormap = cm.LinearColormap(colors=['#d7191c', '#fdae61', '#a6d96a', '#1a9641'], vmin=jc_min, vmax=jc_max, caption='Jednotková cena [Kč/m²]')
m.add_child(colormap)

for _, row in df_dedup.iterrows():
    lat, lon = row.get("refPoint_lat"), row.get("refPoint_lon")
    if pd.isna(lat) or pd.isna(lon): continue
    popup_txt = (f"<b>Vklad: {row['cislo_vkladu']}</b><br>KÚ: {row.get('ku_nazev','–')}<br>UP: {row.get('UP','–')}<br>"
                 f"JC: <b>{row['JC']:,.0f} Kč/m²</b><br>Výměra: {row['vymera']:,.0f} m²<br>Datum: {str(row.get('datum_podani',''))[:10]}")
    folium.CircleMarker(
        location=[lat, lon], radius=8, color="black", weight=1, fill=True, fill_color=colormap(row["JC"]),
        fill_opacity=0.9, popup=folium.Popup(popup_txt, max_width=250), tooltip=f"{row['JC']:,.0f} Kč/m²",
    ).add_to(m)

folium.Marker(
    location=[OCE_LAT, OCE_LON],
    popup=folium.Popup(f"<b>OCEŇOVANÁ NEMOVITOST</b><br>{OCE_KU}<br>UP: {OCE_UP}<br>Výměra: {OCE_VYMERA:,.0f} m²<br>PP index: {OCE_PP_INDEX}", max_width=220),
    icon=folium.Icon(color="red", icon="home", prefix="fa"), tooltip="Oceňovaná nemovitost",
).add_to(m)

if sw and ne: m.fit_bounds([sw, ne])
folium.LayerControl().add_to(m)
m.save(OUT_MAP)
print(f"    Uložena mapa: {OUT_MAP}")

# =============================================================================
# 13. EXCEL VÝSTUP
# =============================================================================

print("[13] Generuji Excel výstup …")
with pd.ExcelWriter(OUT_XLSX, engine="openpyxl") as writer:
    df_clean_out = df_clean[["cislo_vkladu","datum_podani","ku_nazev","UP","vymera","JC","K1","K2","K3","K4","K5","JC_upravena","pp_index"]].copy()
    df_clean_out.columns = ["Číslo vkladu","Datum podání","KÚ","UP","Výměra [m²]","JC [Kč/m²]","K1 čas","K2 lokalita","K3 typ UP","K4 velikost","K5 tvar","JC upravená [Kč/m²]","PP index"]
    df_clean_out.to_excel(writer, sheet_name="Data_Očištěná", index=False)

    pd.DataFrame([{"Ukazatel": l, **s} for l, s in [("Čas [dny]", stats_cas), ("Výměra [m²]", stats_vym), ("JC [Kč/m²]", stats_jc)]]).to_excel(writer, sheet_name="Statistika", index=False)
    pd.DataFrame([
        {"Koeficient": "K1 – Čas",     "Popis": "log-lin regrese", "Slope/Hodnota": round(slope_k1,6)},
        {"Koeficient": "K2 – Lokalita","Popis": f"Medián JC {OCE_KU}", "Slope/Hodnota": round(median_oce_ku,0)},
        {"Koeficient": "K3 – Typ UP",  "Popis": f"Medián JC UP={OCE_UP}", "Slope/Hodnota": round(median_oce_up,0)},
        {"Koeficient": "K4 – Velikost","Popis": "log-log regrese slope", "Slope/Hodnota": round(slope_k4,4)},
        {"Koeficient": "K5 – Tvar",    "Popis": f"PP index oceňované={OCE_PP_INDEX}", "Slope/Hodnota": OCE_PP_INDEX},
    ]).to_excel(writer, sheet_name="Analýza a porovnávací metoda", index=False)

    if len(df_outliers) > 0:
        df_out_out = df_outliers[["cislo_vkladu","datum_podani","ku_nazev","UP","vymera","JC"]].copy()
        df_out_out.columns = ["Číslo vkladu","Datum","KÚ","UP","Výměra [m²]","JC [Kč/m²]"]
    else: df_out_out = pd.DataFrame({"Poznámka": ["Žádné outliers nebyly detekovány."]})
    df_out_out.to_excel(writer, sheet_name="Extrémy", index=False)

    pd.DataFrame([
        {"Ukazatel": "Vážený průměr JC upravené [Kč/m²]",  "Hodnota": round(etalon_vazeny,  0)},
        {"Ukazatel": "Medián JC upravené [Kč/m²]",          "Hodnota": round(etalon_median,  0)},
        {"Ukazatel": "Směrodatná odchylka [Kč/m²]",         "Hodnota": round(etalon_std,     0)},
        {"Ukazatel": "Rozsah – Q1 [Kč/m²]",                 "Hodnota": round(etalon_p25,     0)},
        {"Ukazatel": "Rozsah – Q3 [Kč/m²]",                 "Hodnota": round(etalon_p75,     0)},
        {"Ukazatel": "Počet vkladů ve vzorku",               "Hodnota": len(df_clean)},
    ]).to_excel(writer, sheet_name="Etalon", index=False)

    pd.DataFrame([{
        "Koeficient": k, "Základ [Kč/m²]": round(v["base_vazeny"], 0), "+10% [Kč/m²]": round(v["up_vazeny"], 0),
        "-10% [Kč/m²]": round(v["dn_vazeny"], 0), "Rozpětí [Kč/m²]": round(v["rozpeti"], 0)
    } for k, v in sensitivity_results.items()]).to_excel(writer, sheet_name="Citlivostní analýza", index=False)

print(f"    Uložen Excel: {OUT_XLSX}")


# Pomocná funkce pro URL na parcely 
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT

def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '14') 
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), '14')
    rPr.append(szCs)
    new_run.append(rPr)
    text_element = OxmlElement('w:t')
    text_element.text = text
    new_run.append(text_element)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


# =============================================================================
# 14. WORD DOKUMENT – Znalecká zpráva
# =============================================================================

print("[14] Generuji Word dokument …")

doc = Document()

style_normal = doc.styles["Normal"]
style_normal.font.name = "Arial"
style_normal.font.size = Pt(11)

from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
def set_cell_background(cell, hex_color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.runs[0].font.name = "Arial"
    return h

def add_paragraph(doc, text, bold=False, italic=False):
    p = doc.add_paragraph(text)
    r = p.runs[0] if p.runs else p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(11)
    if bold:   r.bold   = True
    if italic: r.italic = True
    return p

def add_table_from_df(doc, df_t, header_bg="2E74B5"):
    table = doc.add_table(rows=1, cols=len(df_t.columns))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df_t.columns):
        hdr_cells[i].text = str(col)
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
    for _, row in df_t.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val) if not (isinstance(val, float) and np.isnan(val)) else "–"
            row_cells[i].paragraphs[0].runs[0].font.size = Pt(9)
    return table

# ── Titulní strana ──────────────────────────────────────────────────────────
doc.add_heading("ANALÝZA TRHU NEMOVITOSTÍ", 0)
doc.add_heading("Stanovení jednotkové obvyklé ceny pozemku – ETALON JC", 1)
doc.add_paragraph("")

intro_meta = [
    ("Věc:",           f"Pozemky k.ú. {OCE_KU}, {OCE_OKRES}"),
    ("Typ pozemku:",   f"Stavební pozemek pro bydlení (UP={OCE_UP})"),
    ("Výměra:",        f"{OCE_VYMERA:,.0f} m²"),
    ("Datum ocenění:", OCE_DATUM.strftime("%d. %m. %Y")),
    ("Zpracoval:",     "Automatizovaný analytický skript – doplňte znalce"),
    ("Číslo posudku:", "040742-2026"),
]
for label, val in intro_meta:
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}  "); r1.bold = True; r1.font.name = "Arial"; r1.font.size = Pt(11)
    r2 = p.add_run(val);                           r2.font.name = "Arial"; r2.font.size = Pt(11)

doc.add_page_break()

# ── 1. Úvod ──────────────────────────────────────────────────────────────────
add_heading(doc, "1. Úvod a předmět analýzy", 1)

uvod_text = f"""
Předmětem této analýzy trhu je stanovení jednotkové obvyklé ceny (dále jen „JC") pozemku
situovaného v katastrálním území {OCE_KU}, okres {OCE_OKRES}, ke dni {OCE_DATUM.strftime('%d. %m. %Y')}.

Oceňovaný pozemek má celkovou výměru {OCE_VYMERA:,.0f} m² a dle aktuálně platného územního plánu
příslušného sídlení útvaru náleží do funkční plochy označené kódem „{OCE_UP}". Tvar pozemku byl vyhodnocen jako vhodný, přičemž Polsby-Popper index tvaru byl stanoven na hodnotu {OCE_PP_INDEX}.

Cílem analýzy je prostřednictvím tržního přístupu dle Mezinárodních oceňovacích standardů (IVS 2022,
standard 105 – Market Approach) identifikovat a statisticky vyhodnotit srovnatelné tržní transakce,
odvodit korekční koeficienty zohledňující rozdíly mezi srovnávanými nemovitostmi a oceňovaným
pozemkem a nakonec stanovit tzv. ETALON JC – referenční jednotkovou cenu průměrného stavebního
pozemku pro bydlení v dané lokalitě, čase a kvalitativní třídě. Tato hodnota bude následně použita
jako výchozí veličina pro individuální ocenění konkrétního pozemku.
""".strip()
add_paragraph(doc, uvod_text)
doc.add_paragraph("")

# ── 2. Metodika a zdrojová data ───────────────────────────────────────────────
add_heading(doc, "2. Metodika a postup analýzy", 1)

metodika_text_1 = """
2.1  Tržní přístup dle IVS (Market Approach)

Tržní přístup oceňuje pozemek na základě porovnání s nedávnými prodeji obdobných pozemků
na trhu za předpokladu, že racionální kupující by za oceňovaný pozemek nezaplatil více,
než kolik by stála koupě srovnatelné alternativní nemovitosti (princip substituce). Metoda
je obecně považována za nejspolehlivější přístup k ocenění pozemků všude tam, kde existuje
dostatečný počet relevantních tržních transakcí.
""".strip()
add_paragraph(doc, metodika_text_1)
doc.add_paragraph("")

add_paragraph(doc, "2.2  Zdrojová data a jejich struktura")

zdroj_text = f"""
Datovým základem analýzy je databáze transakcí s pozemky vedená na základě zápisů
do katastru nemovitostí (tzv. vkladů). Následující vizualizace zobrazuje prostorové rozložení
identifikovaného datového vzorku vůči oceňované nemovitosti, přičemž barva bodu indikuje
jednotkovou cenovou hladinu transakce (zelená = vyšší JC, červená = nižší JC).
""".strip()
add_paragraph(doc, zdroj_text)
doc.add_paragraph("")

if "prostorova_mapa" in GRAPH_FILES:
    doc.add_picture(GRAPH_FILES["prostorova_mapa"], width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Obrázek 1: Prostorové a cenové rozložení vzorku").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

add_paragraph(doc, "Surový datový vzorek (Všechny parcely, řazeno chronologicky):", bold=True)
doc.add_paragraph("")

df_raw_table = df.dropna(subset=["JC", "vymera"]).copy()
df_raw_table.sort_values(by="datum_podani", ascending=False, inplace=True)

tab_cols = ["Datum podání", "Číslo vkladu", "Okres", "K.Ú.", "Parc.č.", "ÚP", "Výměra [m²]", "JC [Kč/m²]"]
table = doc.add_table(rows=1, cols=len(tab_cols))
table.style = "Table Grid"

hdr_cells = table.rows[0].cells
for i, name in enumerate(tab_cols):
    hdr_cells[i].text = name
    hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(8)
    set_cell_background(hdr_cells[i], "2E74B5")

import matplotlib.colors as mcolors
vmin_jc, vmax_jc = df_raw_table["JC"].min(), df_raw_table["JC"].max()
vmin_vym, vmax_vym = df_raw_table["vymera"].min(), df_raw_table["vymera"].max()

def get_hex_color(val, vmin, vmax, cmap_name="RdYlGn"):
    if pd.isna(val): return "FFFFFF"
    norm_val = (val - vmin) / max(vmax - vmin, 1)
    rgba = plt.cm.get_cmap(cmap_name)(norm_val)
    return mcolors.to_hex(rgba).replace("#", "").upper()

prev_year = None
year_bg = "FFFFFF"

for _, row in df_raw_table.iterrows():
    r_cells = table.add_row().cells
    
    d_datum = str(row["datum_den"])
    d_vklad = str(row.get("cislo_vkladu", "–"))
    d_okres = str(row.get("okres_nazev", "–"))
    d_ku    = str(row.get("ku_nazev", "–"))
    d_up    = str(row.get("UP", "–"))
    d_vym   = row.get("vymera", 0)
    d_jc    = row.get("JC", 0)
    
    r_cells[0].text = d_datum
    r_cells[1].text = d_vklad
    r_cells[2].text = d_okres
    r_cells[3].text = d_ku
    
    cell_parc = r_cells[4]
    cell_parc.text = ""
    paragraph = cell_parc.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    d_parc = str(row.get("parcel_number", "–")) 
    d_ruian_id = row.get("ruian_parcela_id")

    if pd.notna(d_ruian_id):
        clean_id = str(int(d_ruian_id))
        url = f"https://nahlizenidokn.cuzk.gov.cz/ZobrazObjekt.aspx?&typ=parcela&id={clean_id}"
        add_hyperlink(paragraph, url, d_parc)
    else:
        paragraph.text = d_parc
    
    r_cells[5].text = d_up
    r_cells[6].text = f"{d_vym:,.0f}"
    r_cells[7].text = f"{d_jc:,.0f}"
    
    for cell in r_cells:
        cell.paragraphs[0].runs[0].font.size = Pt(7)
        
    current_year = d_datum[:4]
    if prev_year is not None and current_year != prev_year:
        year_bg = "F2F2F2" if year_bg == "FFFFFF" else "FFFFFF"
    prev_year = current_year
    
    for i in range(6):
        set_cell_background(r_cells[i], year_bg)
        
    bg_vym = get_hex_color(d_vym, vmin_vym, vmax_vym, cmap_name="Blues")
    set_cell_background(r_cells[6], bg_vym)
    
    bg_jc = get_hex_color(d_jc, vmin_jc, vmax_jc, cmap_name="RdYlGn")
    set_cell_background(r_cells[7], bg_jc)
    
doc.add_paragraph("")

# ÚPRAVA TEXTU O DETEKCI EXTRÉMŮ (Standardní IQR)
metodika_text_2 = f"""
2.3  Deduplikace dat na úroveň vkladu

Veškeré statistické výpočty jsou prováděny nad deduplikovaným souborem, tj. pro každé
„cislo_vkladu" je uvažována právě jedna hodnota JC (která je totožná pro všechny parcely
v rámci téhož vkladu). Pokud v rámci jednoho vkladu figurují parcely s různým kódem UP,
je za dominantní UP považováno to, které plošně (dle výměry) převažuje nad 50 %; v opačném
případě je vklad zařazen do kategorie „UP smíšené".

2.4  Polsby-Popper index tvaru pozemku

Tvar pozemku výrazně ovlivňuje jeho využitelnost a tím i obvyklou cenu. Pro kvantifikaci tvaru
každé parcely je využíván tzv. Polsby-Popper index (PP), definovaný jako:

    PP = (4π · A) / P²

kde A je plocha a P obvod polygonu parcely. Hodnota PP = 1,0 odpovídá dokonalému kruhu (ideální
tvar), hodnota blízká 0 signalizuje extrémně štíhlý nebo složitý tvar. Pro výpočet PP indexu
jsou použity S-JTSK souřadnice z polygonu parcely (sloupec „geometry_posList") a knihovna Shapely.

2.5  Detekce a zpracování extrémních hodnot (outlierů)

Trh nemovitostí je charakteristický velkým rozptylem realizovaných cen. Transakce s neobvyklou JC 
(výrazně pod nebo nad průměrem trhu) mohou odrážet prodej podílu na nemovitosti, transakci mezi 
spřízněnými osobami, chybu v datech nebo jiné mimotržní vlivy.

Pro zajištění objektivity byla pro detekci extrémů použita statistická metoda interkvartilního 
rozpětí (Standardní IQR metoda). Metoda byla aplikována přímo na lineární data jednotkových cen.
Vypočtená spodní hranice byla navíc logicky omezena zdola (na úroveň 5. percentilu), aby 
nedocházelo k vyřazování přirozeně levných, ale standardních pozemků.

Outlierem je označena transakce, jejíž jednotková cena leží mimo takto definovaný interval. 
V aplikaci na tento konkrétní datový soubor to představuje matematicky odvozené rozmezí 
přípustných jednotkových cen od {lower_bound:,.0f} Kč/m² do {upper_bound:,.0f} Kč/m². 
Transakce mimo tento interval jsou z analytického vzorku vyřazeny.

2.6  Korekční koeficienty K1–K5

K zohlednění rozdílů mezi oceňovanou nemovitostí a srovnávacími transakcemi jsou odvozeny
koeficienty K1 až K5, každý vymezený metodou stanovenou v zadání:

  • K1 (čas):     log-lineární regrese ln(JC) ~ čas_v_dnech; koeficient je roven
                  exp(β · dny_od_transakce) a posunuje cenu transakce k datu ocenění.
                  Regresní sklon β = {slope_k1:.6f} (R² = {r_k1**2:.3f}).

  • K2 (lokalita): podíl mediánu JC katastrálního území oceňované nemovitosti ({OCE_KU})
                   a mediánu JC katastrálního území dané transakce.
                   Medián JC pro {OCE_KU}: {median_oce_ku:,.0f} Kč/m².

  • K3 (typ UP):  podíl mediánu JC ploch s funkčním využitím {OCE_UP} a mediánu JC plochy
                  UP dané transakce.
                  Medián JC pro UP={OCE_UP}: {median_oce_up:,.0f} Kč/m².

  • K4 (velikost): log-log regrese ln(JC) ~ ln(výměra); koeficient je roven
                   (výměra_oceňované / výměra_transakce)^β₄.
                   Regresní sklon β₄ = {slope_k4:.4f} (R² = {r_k4**2:.3f}).
                   {'Záporná hodnota sklonu potvrzuje platnost diskontu na velikost (s rostoucí výměrou klesá JC).' if slope_k4 < 0 else 
                    'Kladná hodnota sklonu naznačuje, že diskont na velikost v daném souboru není statisticky prokázán.'}

  • K5 (tvar):    podíl Polsby-Popper indexu oceňované nemovitosti ({OCE_PP_INDEX})
                  a PP indexu dané transakce.

Každý korekční koeficient je po výpočtu oříznut na interval [0,50; 2,00], čímž je zabráněno
neúměrným distorzím způsobeným extrémními hodnotami.

2.7  Výpočet upraveného JC a ETALONU

Upravená JC transakce: JC_upravená = JC · K1 · K2 · K3 · K4 · K5

ETALON JC je následně stanoven jako:
  (a) vážený průměr JC_upravené (váha = celková výměra vkladu), a
  (b) medián JC_upravené.
""".strip()
add_paragraph(doc, metodika_text_2)
doc.add_paragraph("")

# ── 3. Data, čištění a extrémy ───────────────────────────────────────────────
add_heading(doc, "3. Zpracování dat a identifikace extrémů", 1)

cisteni_text = f"""
Ze zdrojového souboru bylo načteno celkem {len(df_raw)} řádků odpovídajících parcelám
zahrnutým do evidovaných vkladů. Po deduplikaci na úroveň jednotlivých vkladů (cislo_vkladu)
a po odstranění záznamů s chybějícími nebo nenulovými hodnotami výměry, ceny nebo data
vznikl soubor {len(df_dedup)} unikátních transakcí.

Detekce extrémních hodnot (Standardní lineární IQR metoda) identifikovala {len(df_outliers)} vkladů
s atypickou JC (hranice: {lower_bound:,.0f} – {upper_bound:,.0f} Kč/m²). Tyto transakce
jsou z analytického vzorku vyloučeny. Výsledný čistý vzorek pro výpočty obsahuje
{len(df_clean)} vkladů.
""".strip()
add_paragraph(doc, cisteni_text)

if len(df_outliers) > 0:
    doc.add_paragraph("")
    add_paragraph(doc, "Přehled identifikovaných extrémů:", bold=True)
    ot = df_outliers[["cislo_vkladu","datum_podani","ku_nazev","UP","vymera","JC"]].copy()
    ot["datum_podani"] = ot["datum_podani"].dt.strftime("%Y-%m-%d")
    ot["vymera"] = ot["vymera"].map(lambda x: f"{x:,.0f}")
    ot["JC"]     = ot["JC"].map(lambda x: f"{x:,.0f}")
    ot.columns   = ["Číslo vkladu","Datum","KÚ","UP","Výměra [m²]","JC [Kč/m²]"]
    add_table_from_df(doc, ot)
else:
    add_paragraph(doc, "V analyzovaném souboru dat nebyly detekovány žádné extrémní hodnoty.")

doc.add_paragraph("")

doc.add_picture(GRAPH_FILES["outliery"], width=Inches(6))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("Graf 1: Identifikace extrémů v souboru dat").alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ── 4. Popisné statistiky a grafy ───────────────────────────────────────────

add_heading(doc, "4. Základní popisné statistiky", 1)

intro_stats_text = f"""
Níže jsou v přehledné tabulkové formě uvedeny klíčové statistické ukazatele jednotkových cen (JC) 
pro očištěný vzorek {len(df_clean)} tržních transakcí. Vážený průměr JC (kde váhou je celková výměra 
pozemků v daném vkladu) poskytuje robustní odhad průměrné tržní cenové hladiny. Směrodatná odchylka 
kvantifikuje absolutní míru rozptylu realizovaných cen kolem průměru.
""".strip()
add_paragraph(doc, intro_stats_text)
doc.add_paragraph("")

stats_tbl_df = pd.DataFrame([
    ["Vážený průměr", f"{stats_jc['průměr']:,.0f} Kč/m²"],
    ["Medián", f"{stats_jc['medián']:,.0f} Kč/m²"],
    ["Minimální hodnota (Min)", f"{stats_jc['min']:,.0f} Kč/m²"],
    ["5. percentil (P5)", f"{stats_jc['P5']:,.0f} Kč/m²"],
    ["První kvartil (Q1)", f"{stats_jc['Q1']:,.0f} Kč/m²"],
    ["Třetí kvartil (Q3)", f"{stats_jc['Q3']:,.0f} Kč/m²"],
    ["95. percentil (P95)", f"{stats_jc['P95']:,.0f} Kč/m²"],
    ["Maximální hodnota (Max)", f"{stats_jc['max']:,.0f} Kč/m²"],
    ["Směrodatná odchylka (Std. odchylka)", f"{stats_jc['std']:,.0f} Kč/m²"],
], columns=["Statistický ukazatel", "Jednotková cena JC"])

add_table_from_df(doc, stats_tbl_df)
doc.add_paragraph("")

doc.add_picture(GRAPH_FILES["histogram_JC"],   width=Inches(6))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("Graf 2: Histogram jednotkových cen").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("")

# ── 5. Odvození koeficientů ──────────────────────────────────────────────────
add_heading(doc, "5. Odvození korekčních koeficientů K1–K5", 1)

k1_text = f"""
K1 – Korekce na čas

Ceny nemovitostí se v čase mění vlivem vývoje nabídky a poptávky, inflace a dalších makroekonomických
faktorů. Pro zachycení tohoto vývoje byla provedena log-lineární regrese:

    ln(JC) = α + β₁ · t

kde t je vzdálenost transakce od data ocenění vyjádřená v kalendářních dnech (záporná hodnota
pro budoucí transakce, kladná pro minulé). Odhadnutý koeficient β₁ = {slope_k1:.6f}
(R² = {r_k1**2:.3f}). Korekční koeficient K1 pro každou transakci je roven exp(β₁ · t),
tzn. posunuje historickou cenu transakce na úroveň ceny platné k datu ocenění.
""".strip()
add_paragraph(doc, k1_text)
doc.add_paragraph("")

doc.add_picture(GRAPH_FILES["casovy_vyvoj"], width=Inches(6.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("Graf 3: Časový vývoj jednotkových cen").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("")

k2_text = f"""
K2 – Korekce na lokalitu

Lokalita je jedním z nejvýznamnějších determinantů hodnoty pozemku. Koeficient K2 je stanoven
jako podíl mediánu JC v katastrálním území oceňované nemovitosti ({OCE_KU}) a mediánu JC
katastrálního území dané transakce. Tento postup přísně vylučuje fyzickou vzdálenost v metrech
nebo regresní modelování vzdálenosti, neboť cenová diferenciace dle k.ú. lépe zachycuje lokální
tržní rozdíly než izotropní vzdálenostní funkce.

Medián JC pro k.ú. {OCE_KU}: {median_oce_ku:,.0f} Kč/m².
""".strip()
add_paragraph(doc, k2_text)
doc.add_paragraph("")

k3_text = f"""
K3 – Korekce na typ (funkční využití ÚP)

Funkční využití dle územního plánu determinuje přípustné způsoby využití pozemku, a tím přímo
ovlivňuje jeho výnosový potenciál. Koeficient K3 je stanoven jako podíl mediánu JC pozemků
s funkčním využitím oceňované nemovitosti (UP = {OCE_UP}) a mediánu JC pozemků s funkčním
využitím dané transakce.

Medián JC pro UP = {OCE_UP}: {median_oce_up:,.0f} Kč/m².
""".strip()
add_paragraph(doc, k3_text)
doc.add_paragraph("")

doc.add_picture(GRAPH_FILES["JC_dle_UP"], width=Inches(6))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("Graf 4: Porovnání mediánů JC dle územního plánu").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("")

k4_text = f"""
K4 – Korekce na velikost (diskont na velikost)

Empiricky bývá pozorováno, že s rostoucí výměrou pozemku klesá jeho JC – tzv. diskont na
velikost. Pro kvantifikaci tohoto efektu byla provedena log-log regrese:

    ln(JC) = α + β₄ · ln(výměra)

Odhadnutý koeficient β₄ = {slope_k4:.4f} (R² = {r_k4**2:.3f}).
{'Záporná hodnota β₄ potvrzuje existenci diskontu na velikost v analyzovaném souboru.' if slope_k4 < 0 else 'V daném souboru nebyla statisticky prokázána závislost JC na velikosti pozemku (β₄ > 0).'}

Na základě log-log vztahu byl pro oceňovanou výměru {OCE_VYMERA:,.0f} m² odvozena průměrná
hodnota koeficientu K4 = {df_clean['K4'].mean():.4f}.
""".strip()
add_paragraph(doc, k4_text)
doc.add_paragraph("")

doc.add_picture(GRAPH_FILES["JC_vs_vymera"], width=Inches(6.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("Graf 5: Závislost jednotkové ceny na výměře").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("")

k5_text = f"""
K5 – Korekce na tvar pozemku (Polsby-Popper index)

Tvar pozemku ovlivňuje jeho zastavitelnost a praktickou využitelnost. Koeficient K5 je
stanoven jako podíl PP indexu oceňované nemovitosti ({OCE_PP_INDEX}) a PP indexu dané
transakce. Hodnota PP indexu byla u {df_clean['pp_index'].notna().sum()} ze {len(df_clean)}
transakcí vypočtena z polygonu parcely; u zbývajících transakcí (chybějící nebo nevalidní
geometrie) byl koeficient K5 nastaven na neutrální hodnotu 1,0.
""".strip()
add_paragraph(doc, k5_text)

doc.add_page_break()

# ── 6. Závěr a ETALON ────────────────────────────────────────────────────────
add_heading(doc, "6. Výsledek – ETALON JC a závěr", 1)

zaver_text_1 = f"""
Na základě porovnávací metody dle IVS byl z čistého vzorku {len(df_clean)} tržních transakcí
odvozen ETALON JC – referenční jednotková cena pro průměrný stavební pozemek pro bydlení
(UP = {OCE_UP}) v katastrálním území {OCE_KU} k datu {OCE_DATUM.strftime('%d. %m. %Y')}:
""".strip()
add_paragraph(doc, zaver_text_1)
doc.add_paragraph("")

etalon_tbl_df = pd.DataFrame([
    ["ETALON JC – vážený průměr", f"{etalon_vazeny:,.0f} Kč/m²"],
    ["ETALON JC – medián",        f"{etalon_median:,.0f} Kč/m²"],
    ["Směrodatná odchylka",       f"{etalon_std:,.0f} Kč/m²"],
    ["Rozsah (Q1–Q3)",            f"{etalon_p25:,.0f} – {etalon_p75:,.0f} Kč/m²"],
    ["Počet transakcí ve vzorku", str(len(df_clean))],
], columns=["Statistický ukazatel", "Výsledná hodnota"])

add_table_from_df(doc, etalon_tbl_df)
doc.add_paragraph("")

zaver_text_2 = f"""
Hodnota ETALONU JC odráží průměrnou tržní cenovou hladinu pozemků srovnatelné kategorie
v dané lokalitě a čase, po eliminaci statisticky atypických transakcí a po korekci na
systematické rozdíly v čase, lokalitě, typu, velikosti a tvaru pozemků. Tato hodnota
slouží jako vstupní parametr pro individuální ocenění konkrétního oceňovaného pozemku,
při němž budou v samostatném výpočtu zohledněny jeho specifické vlastnosti.

Upozornění: Výsledná hodnota ETALONU JC je odvozena z dostupných tržních dat a metody
popsané v kapitole 2. Přesnost výsledku závisí na kvalitě, úplnosti a reprezentativnosti
vstupních dat. Znalec je povinen ověřit všechny vstupní předpoklady a v případě nutnosti
výsledek korigovat na základě vlastního odborného úsudku.
""".strip()
add_paragraph(doc, zaver_text_2)

doc.add_page_break()

# ── 7. Citlivostní analýza ────────────────────────────────────────────────────
add_heading(doc, "7. Citlivostní analýza", 1)

citlivost_text = f"""
Citlivostní analýza testuje robustnost výsledného ETALONU JC vůči izolovaným změnám 
vstupních parametrů a matematických driverů jednotlivých korekčních koeficientů K1–K5 o ±{int(SENSITIVITY_DELTA*100)} % (ceteris paribus). 
Tento postup identifikuje, které statistické předpoklady mají největší vliv na výslednou cenu.
""".strip()
add_paragraph(doc, citlivost_text)
doc.add_paragraph("")

headers = [
    "Označení", 
    "Sledovaný parametr koeficientu", 
    "Etalon při +10 % parametru [Kč/m²]", 
    "Výchozí Etalon (Základ) [Kč/m²]", 
    "Etalon při -10 % parametru [Kč/m²]", 
    "Celkové rozpětí (Volatilitia) [Kč/m²]"
]

table = doc.add_table(rows=1, cols=len(headers))
table.style = "Table Grid"

hdr_cells = table.rows[0].cells
for i, name in enumerate(headers):
    hdr_cells[i].text = name
    hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(8)
    set_cell_background(hdr_cells[i], "2E74B5")

param_desc = {
    "K1": "Sklon časové regrese (rychlost vývoje trhu v čase)",
    "K2": f"Hladina mediánu jednotkové ceny v cílovém k.ú. {OCE_KU}",
    "K3": f"Hladina mediánu jednotkové ceny pro funkční využití ÚP ({OCE_UP})",
    "K4": "Sklon log-log regrese výměry (koeficient diskontu na velikost)",
    "K5": "Polsby-Popper index tvaru oceňovaného pozemku"
}

for k in tornado_order:
    r = sensitivity_results[k]
    row_cells = table.add_row().cells
    
    row_cells[0].text = str(k)
    row_cells[1].text = param_desc[k]
    row_cells[2].text = f"{r['up_vazeny']:,.0f}"
    row_cells[3].text = f"{r['base_vazeny']:,.0f}"
    row_cells[4].text = f"{r['dn_vazeny']:,.0f}"
    row_cells[5].text = f"{r['rozpeti']:,.0f}"
    
    for cell in row_cells:
        cell.paragraphs[0].runs[0].font.size = Pt(7)
    
    set_cell_background(row_cells[0], "F2F2F2")
    set_cell_background(row_cells[1], "F2F2F2")
    set_cell_background(row_cells[2], "E2EFDA")
    set_cell_background(row_cells[3], "F2F2F2")
    set_cell_background(row_cells[4], "FCE4D6")
    
    if r['rozpeti'] > 500:
        set_cell_background(row_cells[5], "D9E1F2")
        row_cells[5].paragraphs[0].runs[0].font.bold = True
    else:
        set_cell_background(row_cells[5], "FFFFFF")

doc.add_paragraph("")

nejcitlivejsi = tornado_order[0]
add_paragraph(doc, f"""
Výsledek citlivostní analýzy ukazuje, že nejvyšší vliv na výsledný ETALON JC má koeficient
{nejcitlivejsi} (rozpětí {sensitivity_results[nejcitlivejsi]['rozpeti']:,.0f} Kč/m² při ±10% změně).
Koeficientu s největším vlivem je proto třeba věnovat zvýšenou pozornost při ověřování
správnosti vstupních předpokladů. Naopak nejméně citlivý koeficient je {tornado_order[-1]}.
""".strip())
doc.add_paragraph("")

doc.add_picture(GRAPH_FILES["tornado"], width=Inches(6))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("Graf 6: Tornádový graf citlivostní analýzy ETALONU JC").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()


# ── 8. Prostorová interpolace cen (Heatmap) ──────────────────────────────────
if "heatmap" in GRAPH_FILES:
    add_heading(doc, "8. Prostorové rozložení cenové hladiny v území", 1)
    heat_text = f"""
Pro vizuální ověření lokalizačního vlivu na cenu (koeficient K2) byla vypracována teplotní mapa (Heatmapa). Metodika využívá plošnou interpolaci naměřených jednotkových cen očištěného vzorku tržních transakcí. Zelené odstíny reprezentují území s nižší cenovou hladinou, zatímco žluté až červené oblasti signalizují prémiové ceny (čím sytější červená, tím vyšší hodnota JC). 

Tento model prostorové analýzy potvrzuje cenové gradienty napříč řešeným územím a demonstruje reprezentativnost stanovených průměrů pro okolí oceňované nemovitosti (vyznačena modrým trojúhelníkem).
    """.strip()
    add_paragraph(doc, heat_text)
    doc.add_paragraph("")
    
    doc.add_picture(GRAPH_FILES["heatmap"], width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Obrázek 2: Cenová mapa - Prostorová interpolace jednotkových cen").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

# ── 9. Cílená analýza podmnožiny (ÚP) ─────────────────────────────────────────
add_heading(doc, f"9. Cílená analýza podmnožiny: Územní plán {OCE_UP}", 1)

if has_sub_up:
    sub_text = f"""
Oceňovaný pozemek spadá dle platného územního plánu do funkčního využití označeného jako „{OCE_UP}“. Z celkového čistého vzorku ({len(df_clean)} vkladů) vyhovuje této specifické klasifikaci přesně {len(df_sub_up)} tržních transakcí.

Níže je uvedeno statistické srovnání této podmnožiny s celkovým trhem, které empiricky verifikuje úroveň korekčního koeficientu K3 (Korekce na typ ÚP). Rozptyl cen je vizualizován formou krabicového grafu (Boxplot), který jasně ohraničuje kvartilová rozpětí obou skupin.
    """.strip()
    add_paragraph(doc, sub_text)
    doc.add_paragraph("")
    
    sub_tbl = pd.DataFrame([
        ["Vážený průměr JC", f"{stats_sub_jc['průměr']:,.0f} Kč/m²", f"{stats_jc['průměr']:,.0f} Kč/m²"],
        ["Medián JC", f"{stats_sub_jc['medián']:,.0f} Kč/m²", f"{stats_jc['medián']:,.0f} Kč/m²"],
        ["Směrodatná odchylka", f"{stats_sub_jc['std']:,.0f} Kč/m²", f"{stats_jc['std']:,.0f} Kč/m²"]
    ], columns=["Statistický ukazatel", f"Podmnožina ({OCE_UP})", "Zbytek trhu (Celkem)"])
    
    add_table_from_df(doc, sub_tbl)
    doc.add_paragraph("")
    
    if "podmnozina_up" in GRAPH_FILES:
        doc.add_picture(GRAPH_FILES["podmnozina_up"], width=Inches(6))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("Graf 7: Boxplot - Porovnání rozptylu JC sledovaného územního plánu").alignment = WD_ALIGN_PARAGRAPH.CENTER
else:
    add_paragraph(doc, f"Pro typ územního plánu {OCE_UP} není v datovém vzorku dostatečný počet prokazatelných tržních transakcí pro plnohodnotnou izolovanou statistickou analýzu (detekovány méně než 3 transakce). Analýza proto v plném rozsahu aplikuje standardní metodiku a celotržní koeficienty odvozené z celého dostupného vzorku (viz kapitola 5).")


# Uložení
doc.save(OUT_DOCX)
print(f"    Uložen Word: {OUT_DOCX}")

# =============================================================================
# 15. ZÁVĚREČNÝ VÝPIS
# =============================================================================

print("\n" + "=" * 70)
print("ANALÝZA DOKONČENA – výsledné soubory:")
print(f"  Excel  : {OUT_XLSX}")
print(f"  Mapa   : {OUT_MAP}")
print(f"  Word   : {OUT_DOCX}")
print(f"  Grafy  : {SCRIPT_DIR}  (graf_*.png)")
print("=" * 70)
print(f"\n  ETALON JC (vážený průměr) = {etalon_vazeny:,.0f} Kč/m²")
print(f"  ETALON JC (medián)        = {etalon_median:,.0f} Kč/m²")
print("=" * 70)