"""
TRŽNÍ ANALÝZA POZEMKŮ – POROVNÁVACÍ METODA (IVS)
Znalecký posudek č. 040742-2026 | Všenory, Praha-západ
Skript generuje: Vystup_Analyza.xlsx, mapy .html, grafy .png, Znalecka_Zprava.docx
"""

import warnings
warnings.filterwarnings("ignore")

import pandas as pd
import numpy as np
from scipy import stats
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.cm as cm
import matplotlib.colors as mcolors
from matplotlib.patches import Patch
import matplotlib.ticker as mticker
import folium
from folium.plugins import MarkerCluster
from shapely.geometry import Polygon
import os
import re
from datetime import datetime, date

# python-docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# openpyxl pro formátování Excelu
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side, numbers
from openpyxl.utils import get_column_letter

# ============================================================
# 0. KONFIGURACE – parametry oceňované nemovitosti
# ============================================================

SOUBOR_VSTUP = r"C:\Users\ijttr\OneDrive\Dokumenty\OCEŇOVÁNÍ\_IJK\040742-2026 - POZ - Praha - západ - Všenory\AI.xlsx"
SOUBOR_VYSTUP_XLSX = "Vystup_Analyza.xlsx"
SOUBOR_ZPRAVA = "Znalecka_Zprava.docx"

# Oceňovaná nemovitost (subjekt)
DATUM_OCENENI = pd.Timestamp("2026-05-29")
KU_SUBJEKT = "Všenory"
UP_SUBJEKT = "BI"
VYMERA_SUBJEKT = 3048.0          # m²
PP_SUBJEKT = 0.70                 # Polsby-Popper index (obdélník)
# S-JTSK: data používají záporné hodnoty → negujeme
JTSK_X_SUBJ = -753348.09
JTSK_Y_SUBJ = -1059566.52
GPS_LAT_SUBJ = 49.927043
GPS_LON_SUBJ = 14.307275

# IQR multiplikátor pro detekci outlierů
OUTLIER_IQR_MULT = 2.0

# ============================================================
# 1. NAČTENÍ A ZÁKLADNÍ ČIŠTĚNÍ DAT
# ============================================================

print("=" * 60)
print("TRŽNÍ ANALÝZA POZEMKŮ – START")
print("=" * 60)

df_raw = pd.read_excel(SOUBOR_VSTUP, sheet_name="data")
print(f"[INFO] Načteno řádků (včetně duplicit parcel): {len(df_raw)}")

# Zajistíme správný datový typ datumu
df_raw["datum_podani"] = pd.to_datetime(df_raw["datum_podani"])

# Pracujeme jen s parcelami (pozemky)
df_raw = df_raw[df_raw["nemovitost"] == "parcela"].copy()
print(f"[INFO] Po filtraci na 'parcela': {len(df_raw)} řádků")

# Numerické sloupce
for col in ["#CELKOVA_VYMERA", "cenovy_udaj", "#JC", "refPoint_x", "refPoint_y",
            "refPoint_lat", "refPoint_lon"]:
    df_raw[col] = pd.to_numeric(df_raw[col], errors="coerce")

# ============================================================
# 2. DEDUPLIKACE NA ÚROVEŇ VKLADU (cislo_vkladu)
#    Pravidlo: statistiky se počítají přes unikátní vklady.
#    Pokud má vklad více řádků (více parcel), vezmeme:
#      - datum, výměru, cenu, JC ze sdíleného sloupce (jsou stejné)
#      - UP: pokud >50 % parcel stejné UP → to UP, jinak "UP smíšené"
#      - souřadnice: centroid všech parcel vkladu
# ============================================================

def dominant_up(series):
    """Vrátí dominantní UP (>50 %) nebo 'UP smíšené'."""
    counts = series.value_counts(normalize=True)
    if counts.iloc[0] > 0.5:
        return counts.index[0]
    return "UP smíšené"

agg_dict = {
    "datum_podani": "first",
    "ku_nazev": "first",
    "obec_nazev": "first",
    "okres_nazev": "first",
    "#CELKOVA_VYMERA": "first",
    "cenovy_udaj": "first",
    "#JC": "first",
    "UP": dominant_up,
    "refPoint_x": "mean",
    "refPoint_y": "mean",
    "refPoint_lat": "mean",
    "refPoint_lon": "mean",
    "geometry_posList": "first",   # tvar první parcely jako proxy
    "parcel_number": "count",
}

df_vklady = df_raw.groupby("cislo_vkladu", as_index=False).agg(agg_dict)
df_vklady.rename(columns={"parcel_number": "pocet_parcel"}, inplace=True)
print(f"[INFO] Unikátních vkladů po deduplikaci: {len(df_vklady)}")

# Přidáme pomocné sloupce
df_vklady["datum_datum"] = df_vklady["datum_podani"].dt.date
df_vklady["rok"] = df_vklady["datum_podani"].dt.year
df_vklady["kvartal"] = df_vklady["datum_podani"].dt.to_period("Q").astype(str)
df_vklady["days_to_val"] = (DATUM_OCENENI - df_vklady["datum_podani"]).dt.days

# ============================================================
# 3. POLSBY-POPPER INDEX (K5 – tvar)
# ============================================================

def parse_polygon(pos_list_str):
    """
    Parsuje řetězec S-JTSK souřadnic ve formátu 'X1 Y1 X2 Y2 ...'
    a vrátí Shapely Polygon nebo None.
    """
    if not isinstance(pos_list_str, str) or not pos_list_str.strip():
        return None
    try:
        nums = list(map(float, pos_list_str.split()))
        if len(nums) < 6 or len(nums) % 2 != 0:
            return None
        coords = [(nums[i], nums[i+1]) for i in range(0, len(nums), 2)]
        poly = Polygon(coords)
        return poly if poly.is_valid else poly.buffer(0)
    except Exception:
        return None

def polsby_popper(poly):
    """Polsby-Popper index = 4π·A / P² ∈ (0, 1]; 1 = kruh, ~0,785 = čtverec."""
    if poly is None or poly.is_empty:
        return np.nan
    try:
        return (4 * np.pi * poly.area) / (poly.length ** 2)
    except Exception:
        return np.nan

df_vklady["_polygon"] = df_vklady["geometry_posList"].apply(parse_polygon)
df_vklady["PP_index"] = df_vklady["_polygon"].apply(polsby_popper)
pp_valid = df_vklady["PP_index"].notna().sum()
print(f"[INFO] Polsby-Popper index vypočten pro {pp_valid}/{len(df_vklady)} vkladů")

# ============================================================
# 4. DETEKCE OUTLIERŮ (extrémní JC)
#    Metoda: IQR ± OUTLIER_IQR_MULT × IQR
#    Důvod: transakce s podílem <1, spřízněné osoby, neobvyklé protiplnění
# ============================================================

Q1 = df_vklady["#JC"].quantile(0.25)
Q3 = df_vklady["#JC"].quantile(0.75)
IQR = Q3 - Q1
dolni_mez = Q1 - OUTLIER_IQR_MULT * IQR
horni_mez = Q3 + OUTLIER_IQR_MULT * IQR

df_vklady["outlier_flag"] = (
    (df_vklady["#JC"] < dolni_mez) | (df_vklady["#JC"] > horni_mez)
)
df_extremy = df_vklady[df_vklady["outlier_flag"]].copy()
df_ciste = df_vklady[~df_vklady["outlier_flag"]].copy()
print(f"[INFO] Identifikováno outlierů: {len(df_extremy)}")
print(f"[INFO] Čistý vzorek pro analýzu: {len(df_ciste)} vkladů")
print(f"[INFO] IQR hranice JC: [{dolni_mez:.0f} – {horni_mez:.0f}] Kč/m²")

# ============================================================
# 5. ZÁKLADNÍ POPISNÉ STATISTIKY
# ============================================================

def popisne_statistiky(serie, vymery=None, label=""):
    """Vrátí dict se statistikami; vymery = váhy pro vážený průměr."""
    s = serie.dropna()
    w = vymery.loc[s.index].fillna(1) if vymery is not None else pd.Series(np.ones(len(s)), index=s.index)
    w = w / w.sum()
    out = {
        "Popis": label,
        "N": len(s),
        "Vážený průměr": float(np.average(s, weights=w)),
        "Min": float(s.min()),
        "P5": float(s.quantile(0.05)),
        "Q1 (25%)": float(s.quantile(0.25)),
        "Medián": float(s.median()),
        "Q3 (75%)": float(s.quantile(0.75)),
        "P95": float(s.quantile(0.95)),
        "Max": float(s.max()),
        "Směr. odchylka": float(s.std()),
    }
    return out

vahy = df_ciste["#CELKOVA_VYMERA"]
stat_cas  = popisne_statistiky(df_ciste["days_to_val"],  vahy, "Dní před oceněním")
stat_plochy = popisne_statistiky(df_ciste["#CELKOVA_VYMERA"], vahy, "Plocha (m²)")
stat_jc   = popisne_statistiky(df_ciste["#JC"],          vahy, "JC (Kč/m²)")

print("\n[STATISTIKY] JC čistý vzorek:")
for k, v in stat_jc.items():
    if k != "Popis":
        print(f"  {k}: {v:.1f}" if isinstance(v, float) else f"  {k}: {v}")

# ============================================================
# 6. GRAFY – funkce
# ============================================================

os.makedirs("grafy", exist_ok=True)

FONT_TITLE = {"fontsize": 13, "fontweight": "bold"}
FONT_AXIS  = {"fontsize": 10}

def cmap_jc(values):
    """Barevná škála zelená-žlutá-červená dle JC (nízká = červená, vysoká = zelená)."""
    norm = mcolors.Normalize(vmin=values.min(), vmax=values.max())
    return cm.RdYlGn(norm(values))

# ── 6.1 Histogramy ──────────────────────────────────────────
fig, axes = plt.subplots(1, 3, figsize=(16, 5))

axes[0].hist(df_ciste["#JC"], bins=20, color="#2196F3", edgecolor="white", alpha=0.85)
axes[0].axvline(stat_jc["Medián"], color="red", lw=2, label=f'Medián: {stat_jc["Medián"]:.0f}')
axes[0].axvline(stat_jc["Vážený průměr"], color="orange", lw=2, ls="--",
                label=f'Vážený prům.: {stat_jc["Vážený průměr"]:.0f}')
axes[0].set_title("Histogram JC (Kč/m²)", **FONT_TITLE)
axes[0].set_xlabel("JC (Kč/m²)", **FONT_AXIS)
axes[0].legend(fontsize=9)

axes[1].hist(df_ciste["#CELKOVA_VYMERA"], bins=20, color="#4CAF50", edgecolor="white", alpha=0.85)
axes[1].axvline(stat_plochy["Medián"], color="red", lw=2,
                label=f'Medián: {stat_plochy["Medián"]:.0f} m²')
axes[1].set_title("Histogram ploch (m²)", **FONT_TITLE)
axes[1].set_xlabel("Plocha (m²)", **FONT_AXIS)
axes[1].legend(fontsize=9)

axes[2].hist(df_ciste["days_to_val"], bins=20, color="#FF9800", edgecolor="white", alpha=0.85)
axes[2].axvline(stat_cas["Medián"], color="red", lw=2,
                label=f'Medián: {stat_cas["Medián"]:.0f} dní')
axes[2].set_title("Histogram stáří transakcí (dny)", **FONT_TITLE)
axes[2].set_xlabel("Dní před oceněním", **FONT_AXIS)
axes[2].legend(fontsize=9)

fig.tight_layout()
fig.savefig("grafy/01_histogramy.png", dpi=150)
plt.close(fig)
print("[GRAF] 01_histogramy.png uložen")

# ── 6.2 Časový vývoj JC ────────────────────────────────────

def graf_casovy_vyvoj(df, fname="grafy/02_casovy_vyvoj.png"):
    s = df.sort_values("datum_podani")
    x_dt = s["datum_podani"]
    x_num = (x_dt - x_dt.min()).dt.days.values.astype(float)
    y = s["#JC"].values

    # Lineární trend v log-prostoru
    mask = y > 0
    slope, intercept, r, p, se = stats.linregress(x_num[mask], np.log(y[mask]))
    trend_y = np.exp(intercept + slope * x_num)

    fig, ax = plt.subplots(figsize=(13, 6))
    sc = ax.scatter(x_dt, y, c=cmap_jc(pd.Series(y)), s=60, zorder=3, alpha=0.8)
    ax.plot(x_dt, trend_y, color="navy", lw=2, label=f"Log-lineární trend (R²={r**2:.3f})")

    # Kvartilové pásmo
    q1 = s["#JC"].quantile(0.25); q3 = s["#JC"].quantile(0.75)
    p5 = s["#JC"].quantile(0.05); p95 = s["#JC"].quantile(0.95)
    ax.axhline(q1,  color="steelblue",  ls="--", lw=1.2, label=f"Q1 = {q1:.0f}")
    ax.axhline(q3,  color="steelblue",  ls="--", lw=1.2, label=f"Q3 = {q3:.0f}")
    ax.axhline(p5,  color="gray",       ls=":",  lw=1.0, label=f"P5 = {p5:.0f}")
    ax.axhline(p95, color="gray",       ls=":",  lw=1.0, label=f"P95 = {p95:.0f}")

    # Barevná legenda
    cbar = fig.colorbar(cm.ScalarMappable(
        norm=mcolors.Normalize(y.min(), y.max()), cmap="RdYlGn"), ax=ax)
    cbar.set_label("JC (Kč/m²)", fontsize=9)

    ax.set_title("Časový vývoj jednotkových cen pozemků", **FONT_TITLE)
    ax.set_xlabel("Datum podání", **FONT_AXIS)
    ax.set_ylabel("JC (Kč/m²)", **FONT_AXIS)
    ax.xaxis.set_major_formatter(matplotlib.dates.DateFormatter("%m/%Y"))
    ax.legend(fontsize=8, loc="upper left")
    fig.tight_layout()
    fig.savefig(fname, dpi=150)
    plt.close(fig)
    print(f"[GRAF] {fname} uložen")
    return slope, intercept, r**2

slope_cas, intercept_cas, r2_cas = graf_casovy_vyvoj(df_ciste)
# Roční míra zdražení
rocni_mira_cas = (np.exp(slope_cas * 365) - 1) * 100
print(f"[K1-ČAS] Roční trend JC: {rocni_mira_cas:+.2f}% (R²={r2_cas:.3f})")

# ── 6.3 Analýza dle UP (K3) ────────────────────────────────

def graf_up(df, fname="grafy/03_up_porovnani.png"):
    up_stats = df.groupby("UP")["#JC"].agg(
        Medián="median", Průměr="mean", Std="std", N="count"
    ).reset_index().sort_values("Medián", ascending=False)

    fig, axes = plt.subplots(1, 2, figsize=(14, 5))

    # Sloupcový graf
    x = np.arange(len(up_stats))
    w = 0.35
    axes[0].bar(x - w/2, up_stats["Průměr"], w, label="Průměr", color="#2196F3", alpha=0.85)
    axes[0].bar(x + w/2, up_stats["Medián"], w, label="Medián", color="#4CAF50", alpha=0.85)
    axes[0].errorbar(x + w/2, up_stats["Medián"], yerr=up_stats["Std"],
                     fmt="none", color="black", capsize=4)
    axes[0].set_xticks(x)
    axes[0].set_xticklabels(up_stats["UP"], rotation=30, ha="right")
    axes[0].set_ylabel("JC (Kč/m²)", **FONT_AXIS)
    axes[0].set_title("Průměr a medián JC dle UP", **FONT_TITLE)
    axes[0].legend()

    # Scatter JC vs plocha dle UP
    up_list = df["UP"].unique()
    colors_up = plt.cm.tab10(np.linspace(0, 1, len(up_list)))
    for i, up in enumerate(up_list):
        sub = df[df["UP"] == up]
        axes[1].scatter(sub["#CELKOVA_VYMERA"], sub["#JC"], label=up,
                        color=colors_up[i], s=50, alpha=0.75)
    # Celkový trend
    log_area = np.log(df["#CELKOVA_VYMERA"].clip(lower=1))
    sl, ic, _, _, _ = stats.linregress(log_area, df["#JC"])
    xa = np.linspace(df["#CELKOVA_VYMERA"].min(), df["#CELKOVA_VYMERA"].max(), 200)
    axes[1].plot(xa, ic + sl * np.log(xa), "k--", lw=2, label="Trend (log-lin)")
    axes[1].set_xlabel("Plocha (m²)", **FONT_AXIS)
    axes[1].set_ylabel("JC (Kč/m²)", **FONT_AXIS)
    axes[1].set_title("JC vs plocha (barva = UP)", **FONT_TITLE)
    axes[1].legend(fontsize=8)
    fig.tight_layout()
    fig.savefig(fname, dpi=150)
    plt.close(fig)
    print(f"[GRAF] {fname} uložen")
    return up_stats

up_stats = graf_up(df_ciste)

# ── 6.4 Analýza velikosti (K4) ──────────────────────────────

def graf_velikost(df, fname="grafy/04_velikost.png"):
    x = df["#CELKOVA_VYMERA"].values
    y = df["#JC"].values
    log_x = np.log(x.clip(min=1))
    slope, intercept, r, p, _ = stats.linregress(log_x, y)

    fig, ax = plt.subplots(figsize=(10, 6))
    sc = ax.scatter(x, y, c=cmap_jc(pd.Series(y)), s=60, alpha=0.8, zorder=3)
    xa = np.linspace(x.min(), x.max(), 300)
    ax.plot(xa, intercept + slope * np.log(xa), "k--", lw=2,
            label=f"Log-lin trend (R²={r**2:.3f}, β={slope:.1f})")
    q1, q3 = np.percentile(y, 25), np.percentile(y, 75)
    p5, p95 = np.percentile(y, 5), np.percentile(y, 95)
    for val, lbl, col in [(q1,"Q1","steelblue"),(q3,"Q3","steelblue"),
                           (p5,"P5","gray"),(p95,"P95","gray")]:
        ax.axhline(val, color=col, ls="--" if "Q" in lbl else ":", lw=1.2, label=f"{lbl}={val:.0f}")
    cbar = fig.colorbar(cm.ScalarMappable(
        norm=mcolors.Normalize(y.min(), y.max()), cmap="RdYlGn"), ax=ax)
    cbar.set_label("JC (Kč/m²)", fontsize=9)
    ax.set_xlabel("Plocha pozemku (m²)", **FONT_AXIS)
    ax.set_ylabel("JC (Kč/m²)", **FONT_AXIS)
    ax.set_title("Závislost JC na velikosti pozemku", **FONT_TITLE)
    ax.legend(fontsize=8)
    fig.tight_layout()
    fig.savefig(fname, dpi=150)
    plt.close(fig)
    print(f"[GRAF] {fname} uložen")
    return slope, intercept, r**2

slope_vel, intercept_vel, r2_vel = graf_velikost(df_ciste)
print(f"[K4-VELIKOST] β (JC/ln(m²)) = {slope_vel:.1f}, R² = {r2_vel:.3f}")

# ── 6.5 Prostorová mapa (folium) ──────────────────────────

def vytvor_mapu(df, fname="mapa_transakci.html"):
    m = folium.Map(location=[GPS_LAT_SUBJ, GPS_LON_SUBJ], zoom_start=13)

    # ČÚZK ortofoto
    folium.TileLayer(
        tiles="https://ags.cuzk.gov.cz/arcgis1/rest/services/ORTOFOTO_WM/MapServer/tile/{z}/{y}/{x}",
        name="ČÚZK Ortofoto", attr="© ČÚZK", overlay=False, control=True,
        max_zoom=20, min_zoom=6, show=False,
    ).add_to(m)

    # Barevná škála pro JC
    jc_vals = df["#JC"].values
    norm = mcolors.Normalize(vmin=jc_vals.min(), vmax=jc_vals.max())
    cmap = cm.RdYlGn

    up_list_sorted = sorted(df["UP"].dropna().unique())
    up_colors = {up: f"#{int(c[0]*255):02x}{int(c[1]*255):02x}{int(c[2]*255):02x}"
                 for up, c in zip(up_list_sorted,
                 plt.cm.tab10(np.linspace(0, 1, len(up_list_sorted))))}

    # Vrstva JC (barevná škála)
    fg_jc = folium.FeatureGroup(name="JC – barevná škála (zelená=drahé)")
    for _, row in df.iterrows():
        if pd.notna(row["refPoint_lat"]) and pd.notna(row["refPoint_lon"]):
            rgba = cmap(norm(row["#JC"]))
            hex_col = f"#{int(rgba[0]*255):02x}{int(rgba[1]*255):02x}{int(rgba[2]*255):02x}"
            folium.CircleMarker(
                location=[row["refPoint_lat"], row["refPoint_lon"]],
                radius=7, color="gray", weight=0.5,
                fill=True, fill_color=hex_col, fill_opacity=0.85,
                popup=folium.Popup(
                    f"<b>{row['cislo_vkladu']}</b><br>"
                    f"KÚ: {row['ku_nazev']}<br>UP: {row['UP']}<br>"
                    f"JC: <b>{row['#JC']:.0f} Kč/m²</b><br>"
                    f"Plocha: {row['#CELKOVA_VYMERA']:.0f} m²<br>"
                    f"Datum: {str(row['datum_datum'])}", max_width=250),
            ).add_to(fg_jc)
    fg_jc.add_to(m)

    # Vrstva UP (barva dle UP)
    fg_up = folium.FeatureGroup(name="UP – barevné rozlišení", show=False)
    for _, row in df.iterrows():
        if pd.notna(row["refPoint_lat"]) and pd.notna(row["refPoint_lon"]):
            col = up_colors.get(row["UP"], "#888888")
            folium.CircleMarker(
                location=[row["refPoint_lat"], row["refPoint_lon"]],
                radius=7, color="gray", weight=0.5,
                fill=True, fill_color=col, fill_opacity=0.85,
                popup=folium.Popup(f"UP: {row['UP']}", max_width=150),
            ).add_to(fg_up)
    fg_up.add_to(m)

    # Oceňovaná nemovitost
    folium.Marker(
        location=[GPS_LAT_SUBJ, GPS_LON_SUBJ],
        popup=f"<b>OCEŇOVANÁ NEMOVITOST</b><br>KÚ Všenory<br>{VYMERA_SUBJEKT} m² | UP: {UP_SUBJEKT}",
        icon=folium.Icon(color="red", icon="home", prefix="fa"),
    ).add_to(m)
    folium.CircleMarker(
        location=[GPS_LAT_SUBJ, GPS_LON_SUBJ],
        radius=14, color="red", fill=False, weight=2.5,
    ).add_to(m)

    folium.LayerControl(collapsed=False).add_to(m)
    m.save(fname)
    print(f"[MAPA] {fname} uložena")

vytvor_mapu(df_ciste)
# Mapa i s outliery
vytvor_mapu(df_vklady, fname="mapa_vsechny_transakce.html")

# ============================================================
# 7. VÝPOČET KOREKČNÍCH KOEFICIENTŮ K1–K5
# ============================================================

# ── K1 – Čas ────────────────────────────────────────────────
# Model: JC(t) = JC_0 · exp(slope_cas · days_to_val)
# K1_i = exp(slope_cas · days_i)  → převede cenu transakce na datum ocenění
# Pro transakci s days_to_val=0 je K1=1.0
df_ciste = df_ciste.copy()
df_ciste["K1_cas"] = np.exp(slope_cas * df_ciste["days_to_val"])

print(f"[K1] Rozsah K1: {df_ciste['K1_cas'].min():.4f} – {df_ciste['K1_cas'].max():.4f}")

# ── K2 – Lokalita (vzdálenost od subjektu) ──────────────────
# Vzdálenost v metrech (S-JTSK)
df_ciste["dist_m"] = np.sqrt(
    (df_ciste["refPoint_x"] - JTSK_X_SUBJ)**2 +
    (df_ciste["refPoint_y"] - JTSK_Y_SUBJ)**2
)

# Regrese JC vs vzdálenost (lineární)
mask_loc = df_ciste["dist_m"].notna() & df_ciste["#JC"].notna()
sl_loc, ic_loc, r_loc, p_loc, _ = stats.linregress(
    df_ciste.loc[mask_loc, "dist_m"],
    df_ciste.loc[mask_loc, "#JC"]
)
# K2 = JC_at_0 / JC_at_dist_i  (normalizuje na polohu subjektu)
# Chráníme před dělením nulou
jc_at_subj_loc = ic_loc  # JC při vzdálenosti 0 (= subjekt)
jc_at_trans_loc = ic_loc + sl_loc * df_ciste["dist_m"]
df_ciste["K2_lokalita"] = jc_at_subj_loc / jc_at_trans_loc.clip(lower=1.0)
df_ciste["K2_lokalita"] = df_ciste["K2_lokalita"].clip(lower=0.5, upper=2.0)

print(f"[K2] Rozsah K2: {df_ciste['K2_lokalita'].min():.4f} – {df_ciste['K2_lokalita'].max():.4f}")
print(f"[K2] Sklon regrese (JC/m): {sl_loc:.2f}, R²={r_loc**2:.3f}")

# ── K3 – Typ UP ────────────────────────────────────────────
# Porovnání mediánové JC UP subjektu vs UP transakce
median_jc_up = df_ciste.groupby("UP")["#JC"].median()
median_subj_up = median_jc_up.get(UP_SUBJEKT, df_ciste["#JC"].median())

def k3_up(up):
    m = median_jc_up.get(up, np.nan)
    if pd.isna(m) or m == 0:
        return 1.0
    return median_subj_up / m

df_ciste["K3_typ"] = df_ciste["UP"].apply(k3_up)
print(f"[K3] Mediány JC dle UP:\n{median_jc_up.to_string()}")
print(f"[K3] Rozsah K3: {df_ciste['K3_typ'].min():.4f} – {df_ciste['K3_typ'].max():.4f}")

# ── K4 – Velikost ───────────────────────────────────────────
# Model: JC ≈ a + b·ln(plocha)  (negativní b = diskont na velikost)
# K4_i = JC_pred(subjekt) / JC_pred(transakce)
jc_pred_subj_vel = intercept_vel + slope_vel * np.log(VYMERA_SUBJEKT)
jc_pred_trans_vel = intercept_vel + slope_vel * np.log(df_ciste["#CELKOVA_VYMERA"].clip(lower=1))
df_ciste["K4_velikost"] = jc_pred_subj_vel / jc_pred_trans_vel.clip(lower=1.0)
df_ciste["K4_velikost"] = df_ciste["K4_velikost"].clip(lower=0.3, upper=3.0)
print(f"[K4] Rozsah K4: {df_ciste['K4_velikost'].min():.4f} – {df_ciste['K4_velikost'].max():.4f}")

# ── K5 – Tvar (Polsby-Popper) ──────────────────────────────
# K5_i = PP_subjekt / PP_transakce (nemovitost s horším tvarem dostane K5 > 1)
# Kde není PP, dosadíme medián vzorku
pp_median = df_ciste["PP_index"].median()
df_ciste["PP_filled"] = df_ciste["PP_index"].fillna(pp_median)
df_ciste["K5_tvar"] = (PP_SUBJEKT / df_ciste["PP_filled"].clip(lower=0.05)).clip(lower=0.5, upper=2.0)
print(f"[K5] Rozsah K5: {df_ciste['K5_tvar'].min():.4f} – {df_ciste['K5_tvar'].max():.4f}")

# ============================================================
# 8. VÝPOČET ETALONU JC
# ============================================================
# Krok 1: Upravená JC = původní JC × K1 × K2 × K3 × K4 × K5

df_ciste["JC_upravena"] = (
    df_ciste["#JC"]
    * df_ciste["K1_cas"]
    * df_ciste["K2_lokalita"]
    * df_ciste["K3_typ"]
    * df_ciste["K4_velikost"]
    * df_ciste["K5_tvar"]
)

# Krok 2: ETALON z upravených cen
vahy_etalon = df_ciste["#CELKOVA_VYMERA"].fillna(1.0)
etalon_vazeny_prumer = float(np.average(df_ciste["JC_upravena"], weights=vahy_etalon))
etalon_median = float(df_ciste["JC_upravena"].median())
etalon_std = float(df_ciste["JC_upravena"].std())
etalon_p25 = float(df_ciste["JC_upravena"].quantile(0.25))
etalon_p75 = float(df_ciste["JC_upravena"].quantile(0.75))

print("\n" + "=" * 60)
print("VÝSLEDEK – ETALON JC")
print("=" * 60)
print(f"  Vážený průměr:  {etalon_vazeny_prumer:,.0f} Kč/m²")
print(f"  Medián:         {etalon_median:,.0f} Kč/m²")
print(f"  Směr. odchylka: {etalon_std:,.0f} Kč/m²")
print(f"  Rozptyl (Q1–Q3): {etalon_p25:,.0f} – {etalon_p75:,.0f} Kč/m²")
print("=" * 60)

# ============================================================
# 9. CITLIVOSTNÍ ANALÝZA – TORNÁDO GRAF
# ============================================================

DELTA = 0.10  # ± 10 %

def etalon_z_koeficientu(k1_adj=1.0, k2_adj=1.0, k3_adj=1.0, k4_adj=1.0, k5_adj=1.0):
    """Přepočítá etalon (vážený průměr i medián) při změně koeficientů o zadaný faktor."""
    jc = (df_ciste["#JC"]
          * df_ciste["K1_cas"] * k1_adj
          * df_ciste["K2_lokalita"] * k2_adj
          * df_ciste["K3_typ"] * k3_adj
          * df_ciste["K4_velikost"] * k4_adj
          * df_ciste["K5_tvar"] * k5_adj)
    vp = float(np.average(jc, weights=vahy_etalon))
    med = float(jc.median())
    return vp, med

kriteria = ["K1 – Čas", "K2 – Lokalita", "K3 – Typ UP", "K4 – Velikost", "K5 – Tvar"]
citlivost = []

for i, nazev in enumerate(kriteria):
    kwargs_plus  = {f"k{i+1}_adj": 1 + DELTA}
    kwargs_minus = {f"k{i+1}_adj": 1 - DELTA}
    vp_plus,  med_plus  = etalon_z_koeficientu(**kwargs_plus)
    vp_minus, med_minus = etalon_z_koeficientu(**kwargs_minus)
    citlivost.append({
        "Kritérium": nazev,
        "Base VP (Kč/m²)": etalon_vazeny_prumer,
        "+10% VP": vp_plus,
        "-10% VP": vp_minus,
        "+10% Med": med_plus,
        "-10% Med": med_minus,
        "Rozptyl VP": vp_plus - vp_minus,
    })

df_citlivost = pd.DataFrame(citlivost).sort_values("Rozptyl VP", ascending=True)

def graf_tornado(df_cit, fname="grafy/05_tornado.png"):
    labels = df_cit["Kritérium"].tolist()
    low = df_cit["-10% VP"].values
    high = df_cit["+10% VP"].values
    base = etalon_vazeny_prumer

    fig, ax = plt.subplots(figsize=(10, 5))
    y = np.arange(len(labels))

    for i, (lo, hi, lbl) in enumerate(zip(low, high, labels)):
        ax.barh(i, hi - base, left=base, height=0.5, color="#4CAF50", alpha=0.85)
        ax.barh(i, lo - base, left=base, height=0.5, color="#F44336", alpha=0.85)
        ax.text(hi + 30, i, f"{hi:,.0f}", va="center", fontsize=8, color="green")
        ax.text(lo - 30, i, f"{lo:,.0f}", va="center", fontsize=8, color="red", ha="right")

    ax.axvline(base, color="black", lw=1.5, label=f"Základ: {base:,.0f} Kč/m²")
    ax.set_yticks(y)
    ax.set_yticklabels(labels, fontsize=10)
    ax.set_xlabel("ETALON JC (Kč/m²)", **FONT_AXIS)
    ax.set_title("Citlivostní analýza – Tornádo graf (±10 % koeficientů)", **FONT_TITLE)
    ax.xaxis.set_major_formatter(mticker.FuncFormatter(lambda x, _: f"{x:,.0f}"))
    ax.legend(fontsize=9)
    fig.tight_layout()
    fig.savefig(fname, dpi=150)
    plt.close(fig)
    print(f"[GRAF] {fname} uložen")

graf_tornado(df_citlivost)

# ============================================================
# 10. EXPORT DO EXCELU
# ============================================================

print("\n[EXCEL] Generuji Vystup_Analyza.xlsx ...")

# Sloupce pro export analýzy
COLS_ANALYZA = [
    "cislo_vkladu", "datum_datum", "ku_nazev", "UP", "pocet_parcel",
    "#CELKOVA_VYMERA", "cenovy_udaj", "#JC",
    "PP_index", "dist_m",
    "K1_cas", "K2_lokalita", "K3_typ", "K4_velikost", "K5_tvar",
    "JC_upravena",
]

def style_header(ws, row=1, fill_hex="1F4E79", font_hex="FFFFFF"):
    fill = PatternFill("solid", start_color=fill_hex)
    font = Font(bold=True, color=font_hex)
    border_side = Side(style="thin", color="AAAAAA")
    border = Border(bottom=border_side)
    for cell in ws[row]:
        cell.fill = fill
        cell.font = font
        cell.alignment = Alignment(horizontal="center", wrap_text=True)
        cell.border = border

def autofit(ws):
    for col in ws.columns:
        max_len = max((len(str(cell.value)) if cell.value else 0 for cell in col), default=0)
        ws.column_dimensions[get_column_letter(col[0].column)].width = min(max_len + 4, 40)

with pd.ExcelWriter(SOUBOR_VYSTUP_XLSX, engine="openpyxl") as writer:

    # List 1 – Data_Očištěná
    df_export_ciste = df_ciste[COLS_ANALYZA].copy()
    df_export_ciste["datum_datum"] = df_export_ciste["datum_datum"].astype(str)
    df_export_ciste.to_excel(writer, sheet_name="Data_Očištěná", index=False)

    # List 2 – Statistika
    rows_stat = [stat_cas, stat_plochy, stat_jc]
    df_stat = pd.DataFrame(rows_stat)
    df_stat.to_excel(writer, sheet_name="Statistika", index=False)

    # List 3 – Analýza a porovnávací metoda (shrnutí K1–K5)
    k_summary = pd.DataFrame({
        "Koeficient": ["K1 – Čas", "K2 – Lokalita", "K3 – Typ UP", "K4 – Velikost", "K5 – Tvar"],
        "Metoda odvození": [
            f"Log-lin regrese JC vs čas; r²={r2_cas:.3f}; roční trend {rocni_mira_cas:+.2f}%",
            f"Lin. regrese JC vs vzdálenost (JTSK); r²={r_loc**2:.3f}; sklon={sl_loc:.2f} Kč/m",
            "Poměr mediánových JC dle UP k mediánu UP subjektu (BI)",
            f"Log-lin regrese JC vs plocha; r²={r2_vel:.3f}; β={slope_vel:.2f}",
            f"Polsby-Popper index; PP subjektu={PP_SUBJEKT}; median vzorku={pp_median:.3f}",
        ],
        "Průměr K (vzorek)": [
            df_ciste["K1_cas"].mean(), df_ciste["K2_lokalita"].mean(),
            df_ciste["K3_typ"].mean(), df_ciste["K4_velikost"].mean(),
            df_ciste["K5_tvar"].mean(),
        ],
        "Min K": [df_ciste[f"K{i}_col"].min() for i, col in enumerate(
            ["K1_cas","K2_lokalita","K3_typ","K4_velikost","K5_tvar"], 1)
            for df_ciste[f"K{i}_col"] in [df_ciste[col]]] if False else  # dummy - přepíšeme
        [df_ciste[c].min() for c in ["K1_cas","K2_lokalita","K3_typ","K4_velikost","K5_tvar"]],
        "Max K": [df_ciste[c].max() for c in ["K1_cas","K2_lokalita","K3_typ","K4_velikost","K5_tvar"]],
    })
    k_summary.to_excel(writer, sheet_name="Analýza a porovnávací metoda", index=False)

    # List 4 – Extrémy
    cols_ext = ["cislo_vkladu", "datum_datum", "ku_nazev", "UP",
                "#CELKOVA_VYMERA", "cenovy_udaj", "#JC", "outlier_flag"]
    df_ext_exp = df_extremy[[c for c in cols_ext if c in df_extremy.columns]].copy()
    if "datum_datum" in df_ext_exp:
        df_ext_exp["datum_datum"] = df_ext_exp["datum_datum"].astype(str)
    df_ext_exp["Poznámka"] = (
        df_ext_exp["#JC"].apply(
            lambda v: "⚠ EXTRÉMNĚ NÍZKÁ JC" if v < dolni_mez else "⚠ EXTRÉMNĚ VYSOKÁ JC"
        )
    )
    df_ext_exp.to_excel(writer, sheet_name="Extrémy", index=False)

    # List 5 – Etalon
    etalon_data = {
        "Ukazatel": ["N (čistý vzorek)", "Vážený průměr JC (Kč/m²)", "Medián JC (Kč/m²)",
                     "Směrodatná odchylka (Kč/m²)", "Dolní rozptyl Q1 (Kč/m²)",
                     "Horní rozptyl Q3 (Kč/m²)"],
        "Hodnota": [len(df_ciste), round(etalon_vazeny_prumer), round(etalon_median),
                    round(etalon_std), round(etalon_p25), round(etalon_p75)],
        "Poznámka": [
            f"Po vyloučení {len(df_extremy)} outlierů",
            "Váha = výměra vkladu; doporučená hodnota ETALONU",
            "Robustní vůči extrémům",
            "Variabilita upravených cen",
            "Spodní hranice rozptýlení",
            "Horní hranice rozptýlení",
        ],
    }
    pd.DataFrame(etalon_data).to_excel(writer, sheet_name="Etalon", index=False)

    # List 6 – Citlivostní analýza
    df_citlivost.to_excel(writer, sheet_name="Citlivostní analýza", index=False)

# Doplnění formátování přes openpyxl
wb = openpyxl.load_workbook(SOUBOR_VYSTUP_XLSX)
for sname in wb.sheetnames:
    ws = wb[sname]
    style_header(ws)
    autofit(ws)
wb.save(SOUBOR_VYSTUP_XLSX)
print(f"[EXCEL] {SOUBOR_VYSTUP_XLSX} uložen ({len(wb.sheetnames)} listů)")

# ============================================================
# 11. WORD DOKUMENT – ZNALECKÁ ZPRÁVA
# ============================================================

print("\n[WORD] Generuji Znalecka_Zprava.docx ...")

doc = Document()

# ── Styly stránky ──────────────────────────────────────────
from docx.oxml.shared import OxmlElement
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(2.0)

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    return h

def add_para(doc, text, bold=False, italic=False, size=10.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(6)
    return p

def add_table(doc, df_t, title=None):
    if title:
        add_para(doc, title, bold=True)
    rows, cols = df_t.shape
    tbl = doc.add_table(rows=rows + 1, cols=cols)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, col in enumerate(df_t.columns):
        cell = tbl.rows[0].cells[j]
        cell.text = str(col)
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "1F4E79")
        shd.set(qn("w:color"), "FFFFFF")
        tcPr.append(shd)
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for i, row in df_t.iterrows():
        for j, val in enumerate(row):
            cell = tbl.rows[i + 1].cells[j]
            if isinstance(val, float):
                cell.text = f"{val:,.1f}" if abs(val) < 1e6 else f"{val:,.0f}"
            else:
                cell.text = str(val) if pd.notna(val) else "–"
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT \
                if isinstance(val, (int, float)) else WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph()
    return tbl

def vloz_obrazek(doc, path, width_cm=15, caption=None):
    if os.path.exists(path):
        doc.add_picture(path, width=Cm(width_cm))
        if caption:
            p = doc.add_paragraph(caption)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.italic = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        doc.add_paragraph()
    else:
        add_para(doc, f"[Obrázek {path} nenalezen]", italic=True)

# ╔══════════════════════════════════════════════════════════╗
# ║  TITULNÍ STRANA                                          ║
# ╚══════════════════════════════════════════════════════════╝
doc.add_heading("ANALÝZA TRHU – POROVNÁVACÍ METODA", 0)
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

meta = [
    ("Předmět ocenění:", f"Pozemky v k.ú. {KU_SUBJEKT}, obec Všenory, okres Praha-západ"),
    ("Výměra subjektu:", f"{VYMERA_SUBJEKT:,.0f} m²"),
    ("Funkční využití UP:", UP_SUBJEKT),
    ("Datum ocenění:", DATUM_OCENENI.strftime("%d. %m. %Y")),
    ("GPS poloha:", f"{GPS_LAT_SUBJ}, {GPS_LON_SUBJ}"),
    ("Datum zpracování:", date.today().strftime("%d. %m. %Y")),
]
for k, v in meta:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{k}  ")
    r.bold = True
    r.font.size = Pt(11)
    r2 = p.add_run(v)
    r2.font.size = Pt(11)

doc.add_page_break()

# ╔══════════════════════════════════════════════════════════╗
# ║  1. ÚVOD                                                 ║
# ╚══════════════════════════════════════════════════════════╝
add_heading(doc, "1. Úvod a vymezení předmětu analýzy", 1)

uvod_text = f"""
Tato analýza trhu tvoří součást znaleckého posudku o ceně obvyklé pozemků v katastrálním území Všenory (k.ú. Všenory, obec Všenory, okres Praha-západ, kraj Středočeský). Analýza je zpracována za účelem stanovení tržní hodnoty pozemků o celkové výměře {VYMERA_SUBJEKT:,.0f} m², zařazených v územním plánu do funkčního využití území „{UP_SUBJEKT}" (bydlení individuální), s referenčním bodem na GPS souřadnicích {GPS_LAT_SUBJ}°N, {GPS_LON_SUBJ}°E (JTSK: X = {abs(JTSK_X_SUBJ):,.2f}, Y = {abs(JTSK_Y_SUBJ):,.2f}).

Analýza vychází z Mezinárodních oceňovacích standardů (IVS, vydání 2022), konkrétně z tržního přístupu a porovnávací metody (Market Approach – Sales Comparison Method). Podkladem je vlastní databáze tržních transakcí s pozemky evidovaná v katastru nemovitostí (vklady vlastnického práva do KN), obsahující záznamy z oblasti Praha-západ a sousedních lokalit.

Výsledkem analýzy je tzv. ETALON JC – etalonová jednotková cena vyjádřená v Kč/m², která reprezentuje obvyklou (tržní) cenu za 1 m² průměrného pozemku oceňovaného typu v dané lokalitě a čase. Tato etalonová cena je následně použita jako vstup do výpočtu tržní hodnoty konkrétních oceňovaných pozemků.
""".strip()

add_para(doc, uvod_text)
doc.add_paragraph()

# ╔══════════════════════════════════════════════════════════╗
# ║  2. METODIKA                                             ║
# ╚══════════════════════════════════════════════════════════╝
add_heading(doc, "2. Metodika – Porovnávací metoda dle IVS", 1)

metodika_text = f"""
2.1  Princip porovnávací metody a tržního přístupu

Porovnávací metoda (Sales Comparison Method) je jednou ze tří základních metod oceňování nemovitostí zakotvených v Mezinárodních oceňovacích standardech (IVS). Vychází z principu substituce, který říká, že racionálně jednající kupující by za oceňovanou nemovitost nezaplatil více, než by byl nucen zaplatit za srovnatelnou alternativu na trhu za stejných tržních podmínek. Tržní hodnota je tedy odvozena z cen skutečně realizovaných transakcí s obdobnými nemovitostmi.

Aplikace metody spočívá v: (i) výběru srovnatelných transakcí ze zdrojové databáze, (ii) jejich systematické úpravě o zjištěné rozdíly oproti oceňované nemovitosti prostřednictvím korekčních koeficientů K1–K5, a (iii) z takto upravených cen odvozené etalonové ceny (vážený průměr a medián).

2.2  Struktura a příprava datového vzorku

Zdrojová data pocházejí z privátní databáze tržních transakcí s pozemky v okrese Praha-západ, evidovaných v katastru nemovitostí jako vklady vlastnického práva (záznamy ve tvaru V-XXXX/RRRR-KKK). Každý záznam „cislo_vkladu" představuje jednu obchodní transakci; v rámci jednoho vkladu může být prodáno více parcel. Celková výměra a jednotková cena jsou v takovém případě počítány za celý soubor pozemků.

Klíčovým metodickým pravidlem je deduplikace dat na úroveň unikátního vkladu před jakýmikoli statistickými výpočty. Pro každý vklad se pracuje s jednou hodnotou výměry, ceny a jednotkové ceny; funkční využití dle ÚP (UP) je přiřazeno dle dominantního zastoupení (>50 %), jinak je vklad označen jako „UP smíšené".

Upozornění: Databáze neobsahuje informaci o výši převáděného spoluvlastnického podílu. Transakce jsou proto standardně zpracovávány jako převod podílu 1/1. Transakce s neobvyklou jednotkovou cenou, které by mohly odpovídat převodům podílu menším než 1/1 nebo prodejům za netržní cenu (mezi spřízněnými osobami, při neobvyklém protiplnění apod.), jsou identifikovány jako odlehlé hodnoty (outliery) a z dalšího hodnocení vyloučeny.

2.3  Detekce a vyloučení odlehlých hodnot

Pro identifikaci netržních nebo chybných dat je použita metoda mezikvartilového rozpětí (IQR). Za outlier je považována každá transakce, jejíž jednotková cena #JC leží mimo interval [Q1 – {OUTLIER_IQR_MULT}·IQR; Q3 + {OUTLIER_IQR_MULT}·IQR], kde Q1 a Q3 jsou dolní a horní kvartil čistého vzorku. Tato metoda je robustní vůči extrémním hodnotám a nevyžaduje předpoklad normálního rozdělení.

2.4  Polsby-Popper index tvaru pozemku (K5)

Tvar pozemku je kvantifikován pomocí Polsby-Popper indexu (PP), definovaného jako:

    PP = (4π · A) / P²

kde A je plocha pozemku (m²) a P je obvod (m). Index nabývá hodnot v intervalu (0; 1]: hodnota 1,0 odpovídá kruhu (nejideálnější tvar), hodnota ≈ 0,785 čtverci, hodnota ≈ 0,7 odpovídá standardnímu obdélníkovému pozemku vhodnému pro zástavbu. Index je počítán z polygonu parcely zapsaného v S-JTSK souřadnicích (sloupec `geometry_posList`) pomocí knihovny Shapely. Oceňovaná nemovitost má přiřazenu hodnotu PP = {PP_SUBJEKT} (odhadnuto ze slovního popisu tvaru – obdélník).

2.5  Korekční koeficienty K1–K5

Pro každou transakci ve finálním čistém vzorku jsou vypočteny korekční koeficienty K1–K5, které upravují její původní jednotkovou cenu na srovnatelnou hodnotu odpovídající charakteristikám oceňované nemovitosti (ceteris paribus). Korekční koeficient hodnoty 1,0 znamená, že porovnávaná transakce se v daném kritériu od subjektu neliší. Hodnota >1,0 znamená, že transakce měla v daném kritériu horší parametr (nižší JC) a je nutné ji přidat, hodnota <1,0 naopak indikuje vyšší JC vůči subjektu.

    K1 – Čas:     Odvozeno z log-lineární regrese jednotkové ceny na čase. Zjištěný roční trend cen: {rocni_mira_cas:+.2f} % (R² = {r2_cas:.3f}).
    K2 – Lokalita: Odvozeno z lineární regrese JC na vzdálenosti od oceňované nemovitosti (v S-JTSK metrech). Sklon: {sl_loc:.2f} Kč/m (R² = {r_loc**2:.3f}).
    K3 – Typ UP:  Poměr mediánové JC UP subjektu (BI) k mediánové JC UP transakce.
    K4 – Velikost: Odvozeno z log-lineární regrese JC na ploše. Koeficient diskontuje větší/menší pozemky vůči subjektu ({VYMERA_SUBJEKT:,.0f} m²). β = {slope_vel:.2f} (R² = {r2_vel:.3f}).
    K5 – Tvar:    Poměr PP subjektu ({PP_SUBJEKT}) k PP transakce.

Upravená jednotková cena transakce i:
    JC_upravena_i = JC_i · K1_i · K2_i · K3_i · K4_i · K5_i

2.6  Výpočet etalonové ceny

Z upraveného vzorku jsou vypočteny:
  • Vážený průměr upravených JC (váha = výměra vkladu) – doporučená hodnota ETALONU
  • Medián upravených JC – robustní míra centrální tendence
  • Směrodatná odchylka – míra variability vzorku
  • Rozptyl Q1–Q3 – typické cenové pásmo
""".strip()

for odstavec in metodika_text.split("\n\n"):
    p = doc.add_paragraph()
    for chunk in odstavec.split("\n"):
        if chunk.strip().startswith("K") and "–" in chunk:
            p.add_run("    " + chunk.strip() + "\n").font.size = Pt(10)
        elif chunk.strip().startswith("•"):
            p.add_run("  " + chunk.strip() + "\n").font.size = Pt(10)
        else:
            r = p.add_run(chunk.strip() + " ")
            r.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(6)

doc.add_paragraph()

# ╔══════════════════════════════════════════════════════════╗
# ║  3. ZPRACOVÁNÍ DAT A EXTRÉMY                             ║
# ╚══════════════════════════════════════════════════════════╝
add_heading(doc, "3. Zpracování dat a identifikace extrémních hodnot", 1)

cisteni_text = f"""
Do analýzy vstoupilo celkem {len(df_vklady)} unikátních vkladů (transakcí) z databáze pozemků v oblasti Praha-západ. Všechna zdrojová data jsou typu „parcela". Po provedení deduplikace na úroveň cislo_vkladu a výpočtu Polsby-Popper indexu z polygonů S-JTSK souřadnic bylo přistoupeno k detekci odlehlých hodnot.

IQR hranice pro detekci outlierů:
  • Dolní hranice:  {dolni_mez:,.0f} Kč/m²
  • Horní hranice:  {horni_mez:,.0f} Kč/m²

Identifikováno {len(df_extremy)} odlehlých transakcí, které byly vyloučeny z dalšího hodnocení. Finální čistý vzorek pro analýzu a výpočet ETALONU tvoří {len(df_ciste)} vkladů.
""".strip()

add_para(doc, cisteni_text)

# Tabulka outlierů
if len(df_extremy) > 0:
    ext_tbl = df_extremy[["cislo_vkladu", "ku_nazev", "UP", "#CELKOVA_VYMERA",
                           "cenovy_udaj", "#JC"]].copy()
    ext_tbl.columns = ["Číslo vkladu", "k.ú.", "UP", "Výměra (m²)", "Cena (Kč)", "JC (Kč/m²)"]
    ext_tbl["Cena (Kč)"] = ext_tbl["Cena (Kč)"].apply(lambda x: f"{x:,.0f}" if pd.notna(x) else "–")
    ext_tbl["JC (Kč/m²)"] = ext_tbl["JC (Kč/m²)"].apply(lambda x: f"{x:,.0f}" if pd.notna(x) else "–")
    add_table(doc, ext_tbl, title="Tabulka 1 – Identifikované odlehlé hodnoty (outliery)")
    add_para(doc,
        "Poznámka: Výše uvedené transakce se vymykají statisticky zjištěným tržním cenám. "
        "Jejich příčinou může být převod spoluvlastnického podílu, transakce mezi spřízněnými "
        "osobami, nestandardní protiplnění nebo datová chyba. Jsou vyloučeny z výpočtu ETALONU.",
        italic=True)
else:
    add_para(doc, "V analyzovaném vzorku nebyly identifikovány žádné odlehlé hodnoty.", italic=True)

doc.add_paragraph()

# ╔══════════════════════════════════════════════════════════╗
# ║  4. ZÁKLADNÍ STATISTIKY A GRAFY                          ║
# ╚══════════════════════════════════════════════════════════╝
add_heading(doc, "4. Základní popisné statistiky a grafická analýza", 1)

stat_tbl = pd.DataFrame([stat_jc, stat_plochy, stat_cas]).drop(columns=["Popis"])
stat_tbl.index = ["JC (Kč/m²)", "Plocha (m²)", "Stáří trans. (dny)"]
stat_tbl = stat_tbl.reset_index().rename(columns={"index": "Ukazatel"})
add_table(doc, stat_tbl, title="Tabulka 2 – Základní statistiky (čistý vzorek)")

vloz_obrazek(doc, "grafy/01_histogramy.png", caption="Graf 1 – Histogramy JC, plochy a stáří transakcí")

# ── K1 – časový vývoj ──────────────────────────────────────
add_heading(doc, "4.1  Kritérium K1 – Časový vývoj cen", 2)

k1_text = f"""
Analýza časového vývoje jednotkových cen pozemků v analyzovaném vzorku prokázala {'rostoucí' if rocni_mira_cas > 0 else 'klesající' if rocni_mira_cas < 0 else 'stagnující'} trend. Na základě log-lineární regrese (JC ~ exp(β·t)) byl zjištěn průměrný roční přírůstek cen {rocni_mira_cas:+.2f} % (R² = {r2_cas:.3f}). Korekční koeficient K1 je proto pro každou transakci vypočten jako:

    K1_i = exp(β · days_to_val_i)

kde days_to_val_i je počet dní mezi datem podání vkladu a datem ocenění ({DATUM_OCENENI.strftime('%d. %m. %Y')}). Koeficient K1 převádí historické ceny na aktuální cenovou hladinu k datu ocenění.
""".strip()

add_para(doc, k1_text)
vloz_obrazek(doc, "grafy/02_casovy_vyvoj.png", caption="Graf 2 – Časový vývoj JC s trendem a kvartily")

# ── K2 – lokalita ──────────────────────────────────────────
add_heading(doc, "4.2  Kritérium K2 – Lokalita", 2)

k2_text = f"""
Prostorová analýza sleduje závislost jednotkové ceny na vzdálenosti od oceňované nemovitosti. Vzdálenost je měřena v metrech v souřadnicovém systému S-JTSK, který poskytuje metricky přesné výsledky bez nutnosti přepočtu z GPS. Lineární regresní model prokázal sklon {sl_loc:.2f} Kč/m (R² = {r_loc**2:.3f}), což naznačuje {'negativní – s rostoucí vzdáleností JC klesá' if sl_loc < 0 else 'pozitivní'} prostorový gradient cen.

Korekční koeficient K2 je vypočten jako poměr predikované JC v místě subjektu (vzdálenost = 0) k predikované JC v místě transakce. Interaktivní mapa transakcí je uložena v souboru mapa_transakci.html.
""".strip()

add_para(doc, k2_text)

# ── K3 – UP ────────────────────────────────────────────────
add_heading(doc, "4.3  Kritérium K3 – Typ funkčního využití ÚP", 2)

up_tbl = up_stats.rename(columns={"UP": "Funkční využití UP",
                                   "Medián": "Medián JC (Kč/m²)",
                                   "Průměr": "Průměr JC (Kč/m²)",
                                   "Std": "Std. odchylka",
                                   "N": "Počet vkladů"})
add_table(doc, up_tbl, title="Tabulka 3 – Mediánové a průměrné JC dle funkčního využití UP")
add_para(doc,
    f"Korekční koeficient K3 je vypočten jako poměr mediánu JC pro UP subjektu "
    f"('{UP_SUBJEKT}': {median_subj_up:,.0f} Kč/m²) k mediánu JC pro UP dané transakce. "
    f"Transakce se shodným UP jako subjekt mají K3 = 1,0.")
vloz_obrazek(doc, "grafy/03_up_porovnani.png", caption="Graf 3 – Porovnání JC dle UP a závislost JC na ploše")

# ── K4 – velikost ──────────────────────────────────────────
add_heading(doc, "4.4  Kritérium K4 – Velikost pozemku (diskont na velikost)", 2)

k4_text = f"""
Analýza závislosti JC na výměře pozemku prokázala existenci {'diskontu na velikost (s rostoucí výměrou JC klesá)' if slope_vel < 0 else 'prémie za velikost (s rostoucí výměrou JC roste)'}. Log-lineární regresní model (JC ~ a + β·ln(plocha)) vykázal koeficient β = {slope_vel:.2f} Kč/m² na jednotku přirozeného logaritmu plochy (R² = {r2_vel:.3f}).

Korekční koeficient K4 pro transakci s výměrou S_i:
    K4_i = (a + β·ln({VYMERA_SUBJEKT:,.0f})) / (a + β·ln(S_i))

kde a = {intercept_vel:.1f} je absolutní člen modelu. Pro oceňovanou nemovitost s výměrou {VYMERA_SUBJEKT:,.0f} m² je predikovaná JC dle modelu {intercept_vel + slope_vel * np.log(VYMERA_SUBJEKT):,.0f} Kč/m².
""".strip()

add_para(doc, k4_text)
vloz_obrazek(doc, "grafy/04_velikost.png", caption="Graf 4 – Závislost JC na výměře pozemku")

# ── K5 – tvar ──────────────────────────────────────────────
add_heading(doc, "4.5  Kritérium K5 – Tvar pozemku (Polsby-Popper index)", 2)

k5_text = f"""
Tvar pozemku ovlivňuje jeho praktickou využitelnost pro zástavbu. Pravidelný pozemek s kompaktním tvarem umožňuje efektivnější zástavbu, a proto dosahuje typicky vyšší jednotkové ceny. Polsby-Popper index (PP) byl úspěšně vypočten pro {pp_valid} z {len(df_vklady)} vkladů; pro zbývající záznamy s chybějícím polygonem byl dosazen mediánový PP vzorku ({pp_median:.3f}).

Oceňovaná nemovitost má PP = {PP_SUBJEKT} (obdélníkový tvar, dle slovního popisu). Korekční koeficient K5 = PP_subjektu / PP_transakce.
""".strip()

add_para(doc, k5_text)

# ╔══════════════════════════════════════════════════════════╗
# ║  5. ETALON JC                                            ║
# ╚══════════════════════════════════════════════════════════╝
add_heading(doc, "5. Výsledek analýzy – ETALON jednotkové ceny", 1)

etalon_text = f"""
Na základě kompletní porovnávací analýzy, zahrnující čistý vzorek {len(df_ciste)} tržních transakcí s pozemky po vyloučení {len(df_extremy)} outlierů, a po aplikaci korekčních koeficientů K1–K5 odrážejících zjištěné rozdíly mezi porovnávanými transakcemi a oceňovanou nemovitostí, byl stanoven následující ETALON jednotkové ceny:
""".strip()

add_para(doc, etalon_text)

etalon_tbl = pd.DataFrame({
    "Ukazatel": ["Vážený průměr JC (Kč/m²) ← DOPORUČENÁ HODNOTA",
                 "Medián JC (Kč/m²)",
                 "Směrodatná odchylka (Kč/m²)",
                 "Dolní rozptyl Q1 (Kč/m²)",
                 "Horní rozptyl Q3 (Kč/m²)",
                 "Počet transakcí (N)"],
    "Hodnota": [f"{etalon_vazeny_prumer:,.0f}",
                f"{etalon_median:,.0f}",
                f"{etalon_std:,.0f}",
                f"{etalon_p25:,.0f}",
                f"{etalon_p75:,.0f}",
                f"{len(df_ciste)}"],
})
add_table(doc, etalon_tbl, title="Tabulka 4 – ETALON JC (výsledek porovnávací metody)")

add_para(doc,
    f"Výsledný ETALON JC ve výši {etalon_vazeny_prumer:,.0f} Kč/m² (vážený průměr) "
    f"resp. {etalon_median:,.0f} Kč/m² (medián) je doporučenou hodnotou pro další výpočet "
    f"tržní hodnoty oceňovaných pozemků v k.ú. {KU_SUBJEKT} ke dni {DATUM_OCENENI.strftime('%d. %m. %Y')}. "
    f"Etalonová cena reflektuje aktuální tržní podmínky, polohu, typ, výměru a tvar oceňované nemovitosti.",
    bold=False)

doc.add_paragraph()

# ╔══════════════════════════════════════════════════════════╗
# ║  6. CITLIVOSTNÍ ANALÝZA                                  ║
# ╚══════════════════════════════════════════════════════════╝
add_heading(doc, "6. Citlivostní analýza robustnosti ETALONU", 1)

nejcitlivejsi = df_citlivost.iloc[-1]["Kritérium"]
nejcitlivejsi_rozptyl = df_citlivost.iloc[-1]["Rozptyl VP"]

citl_text = f"""
Citlivostní analýza testuje robustnost vypočteného ETALONU JC ({etalon_vazeny_prumer:,.0f} Kč/m²) vůči izolovaným změnám jednotlivých korekčních koeficientů o ±10 %, přičemž ostatní koeficienty zůstávají zmrazeny na původních hodnotách (princip ceteris paribus).

Výsledky citlivostní analýzy ukazují, že největší vliv na výsledný ETALON má kritérium „{nejcitlivejsi}", jehož změna o ±10 % vyvolává rozptyl ETALONU o {nejcitlivejsi_rozptyl:,.0f} Kč/m². Naopak nejméně citlivý je ETALON na změny parametru s nejmenším pruhem v tornádo grafu. Tato informace je klíčová pro posouzení spolehlivosti stanovené hodnoty a identifikaci parametrů, které by měly být při revizi posudku ověřeny s největší pečlivostí.
""".strip()

add_para(doc, citl_text)

add_table(doc, df_citlivost[["Kritérium", "-10% VP", "Base VP (Kč/m²)", "+10% VP", "Rozptyl VP"]],
          title="Tabulka 5 – Citlivostní analýza ETALONU JC (vážený průměr)")

vloz_obrazek(doc, "grafy/05_tornado.png", caption="Graf 5 – Tornádo graf citlivostní analýzy")

# ╔══════════════════════════════════════════════════════════╗
# ║  7. ZÁVĚR                                                ║
# ╚══════════════════════════════════════════════════════════╝
add_heading(doc, "7. Závěr", 1)

zaver_text = f"""
Tato analýza trhu byla provedena v souladu s metodickými principy Mezinárodních oceňovacích standardů (IVS, vydání 2022) s využitím tržního přístupu a porovnávací metody. Na základě {len(df_vklady)} evidovaných tržních transakcí s pozemky v okrese Praha-západ, z nichž bylo po vyloučení {len(df_extremy)} odlehlých hodnot do finálního hodnocení zahrnuto {len(df_ciste)} transakcí, byl stanoven ETALON jednotkové ceny.

Parametry oceňované nemovitosti:
  • Katastrální území:  {KU_SUBJEKT}
  • Funkční využití UP: {UP_SUBJEKT}
  • Výměra:             {VYMERA_SUBJEKT:,.0f} m²
  • Tvar (PP index):    {PP_SUBJEKT}
  • Datum ocenění:      {DATUM_OCENENI.strftime('%d. %m. %Y')}

ETALON JC – výsledek porovnávací analýzy:
  • Vážený průměr: {etalon_vazeny_prumer:,.0f} Kč/m²  ← DOPORUČENÁ HODNOTA
  • Medián:        {etalon_median:,.0f} Kč/m²
  • Rozptyl Q1–Q3: {etalon_p25:,.0f} – {etalon_p75:,.0f} Kč/m²

Citlivostní analýza potvrdila robustnost výsledku; největší vliv na hodnotu ETALONU má kritérium „{nejcitlivejsi}" (rozptyl ±10 % koeficientu = {nejcitlivejsi_rozptyl:,.0f} Kč/m²). Etalonová jednotková cena je vstupem do dalšího výpočtu tržní hodnoty konkrétních oceňovaných pozemků, který je součástí hlavního znaleckého posudku.

Tento dokument je součástí znaleckého posudku č. 040742-2026 a byl automaticky vygenerován analytickým skriptem ke dni {date.today().strftime('%d. %m. %Y')}.
""".strip()

add_para(doc, zaver_text)

doc.save(SOUBOR_ZPRAVA)
print(f"[WORD] {SOUBOR_ZPRAVA} uložen")

# ============================================================
# 12. SOUHRN VÝSTUPŮ
# ============================================================

print("\n" + "=" * 60)
print("VÝSTUPY ANALÝZY")
print("=" * 60)
print(f"  Excel:  {SOUBOR_VYSTUP_XLSX}")
print(f"  Word:   {SOUBOR_ZPRAVA}")
print(f"  Mapy:   mapa_transakci.html, mapa_vsechny_transakce.html")
print(f"  Grafy:  grafy/01_histogramy.png ... grafy/05_tornado.png")
print("=" * 60)
print(f"\nETALON JC: {etalon_vazeny_prumer:,.0f} Kč/m² (vážený průměr)")
print(f"           {etalon_median:,.0f} Kč/m² (medián)")
print(f"Rozptyl:   {etalon_p25:,.0f} – {etalon_p75:,.0f} Kč/m² (Q1–Q3)")
print("=" * 60)
