# =============================================================================
# ANALÝZA TRHU POZEMKŮ – POROVNÁVACÍ METODA (IVS / Market Approach)
# =============================================================================

import os
import warnings
import numpy as np
import pandas as pd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
from matplotlib.colors import Normalize
import seaborn as sns
from scipy import stats
import scipy.interpolate as interp
from shapely.geometry import Polygon
import folium
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from docx.opc.constants import RELATIONSHIP_TYPE as RT

warnings.filterwarnings("ignore")

# =============================================================================
# 0. KONFIGURACE – všechny vstupní parametry oceňované nemovitosti
# =============================================================================

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
INPUT_FILE = os.path.join(SCRIPT_DIR, "AI.xlsx")
OUT_XLSX   = os.path.join(SCRIPT_DIR, "Vystup_Analyza.xlsx")
OUT_MAP    = os.path.join(SCRIPT_DIR, "mapa_pozemky.html")
OUT_DOCX   = os.path.join(SCRIPT_DIR, "Znalecka_Zprava.docx")

# Oceňovaná nemovitost
OCE_DATUM        = pd.Timestamp("2026-06-03")
OCE_KU           = "Smíchov"          
OCE_OKRES        = "Hlavní město Praha" 

#OCE_UP           = "LR"               
OCE_UP           = "OB"               

#OCE_VYMERA       = 138332.0           
OCE_VYMERA       = 900.0           

#OCE_PP_INDEX     = 0.10               
OCE_PP_INDEX     = 0.70               

OCE_LAT          = 50.07132229220769  
OCE_LON          = 14.364680438635116

# Parametry metodiky
OUTLIER_IQR_MULT = 1.5               
SENSITIVITY_DELTA = 0.10             
DATE_REF = OCE_DATUM                 
CMAP_JC = "RdYlGn_r"  # Globální paleta (Zelená = Levné, Červená = Drahé)

# =============================================================================
# 1. NAČTENÍ DAT
# =============================================================================

print("=" * 70)
print(f"ANALÝZA TRHU POZEMKŮ – {OCE_OKRES} / {OCE_KU}")
print("=" * 70)
print(f"\n[1] Načítám data ze souboru: {INPUT_FILE}")

df_raw = pd.read_excel(INPUT_FILE, sheet_name="data", header=0)
print(f"    Celkem řádků (surová data): {len(df_raw)}")

# =============================================================================
# 2. PŘÍPRAVA DAT
# =============================================================================

print("\n[2] Čistím a připravuji data …")
df = df_raw.copy()
df.columns = df.columns.str.strip()
df["datum_podani"] = pd.to_datetime(df["datum_podani"], errors="coerce")
df["datum_den"]    = df["datum_podani"].dt.date

if "JC_calc" in df.columns:
    df["JC"] = pd.to_numeric(df["JC_calc"], errors="coerce")
elif "#JC [Kč/m2]" in df.columns:
    df["JC"] = pd.to_numeric(df["#JC [Kč/m2]"], errors="coerce")
else:
    df["JC"] = pd.to_numeric(df["cenovy_udaj"], errors="coerce") / pd.to_numeric(df["#CELKOVA_VYMERA"], errors="coerce")

df["vymera"]       = pd.to_numeric(df["#CELKOVA_VYMERA"], errors="coerce")
df["cena_celkem"]  = pd.to_numeric(df["cenovy_udaj"],     errors="coerce")
df["cas_v_dnech"] = (DATE_REF - df["datum_podani"]).dt.days
if "UP" in df.columns: df["UP"] = df["UP"].astype(str).str.strip()

# =============================================================================
# 3. DEDUPLIKACE NA ÚROVEŇ cislo_vkladu
# =============================================================================

print("[3] Deduplicita na úroveň vkladu (cislo_vkladu) …")
def dominant_up(sub):
    if sub["UP"].nunique() == 1: return sub["UP"].iloc[0]
    vymery = sub.groupby("UP")["vymera"].sum()
    return vymery.idxmax() if (vymery.max() / vymery.sum()) > 0.5 else "UP smíšené"

agg_dict = {
    "datum_podani":  "first", "cas_v_dnech":   "first", "JC": "first",       
    "vymera":        "first", "cena_celkem":   "first", "ku_nazev": "first",
    "okres_nazev":   "first", "parcel_number": "first", "ruian_parcela_id": "first",
    "refPoint_lat":  "first", "refPoint_lon":  "first", "geometry_posList": "first",
}
agg_dict = {k: v for k, v in agg_dict.items() if k in df.columns}
df_dedup = df.groupby("cislo_vkladu").agg(agg_dict).reset_index()

up_per_vklad = df.groupby("cislo_vkladu").apply(dominant_up).reset_index()
up_per_vklad.columns = ["cislo_vkladu", "UP_dom"]
df_dedup = df_dedup.merge(up_per_vklad, on="cislo_vkladu", how="left")
df_dedup.rename(columns={"UP_dom": "UP"}, inplace=True)
df_dedup.dropna(subset=["JC", "vymera", "datum_podani"], inplace=True)
df_dedup = df_dedup[df_dedup["JC"] > 0].reset_index(drop=True)
print(f"    Unikátních vkladů po deduplikaci: {len(df_dedup)}")

# =============================================================================
# 4. POLSBY-POPPER INDEX (K5 tvar)
# =============================================================================

print("[4] Výpočet Polsby-Popper indexu z geometry_posList …")
def parse_geometry(geom_str):
    try:
        if pd.isna(geom_str) or str(geom_str).strip() == "": return None
        nums = list(map(float, str(geom_str).split()))
        if len(nums) < 6 or len(nums) % 2 != 0: return None
        coords = [(nums[i], nums[i+1]) for i in range(0, len(nums), 2)]
        poly = Polygon(coords)
        return poly if poly.is_valid else poly.buffer(0)
    except Exception: return None

if "geometry_posList" in df_dedup.columns:
    df_dedup["polygon"]   = df_dedup["geometry_posList"].apply(parse_geometry)
    df_dedup["pp_area"]   = df_dedup["polygon"].apply(lambda p: p.area if p is not None else np.nan)
    df_dedup["pp_perim"]  = df_dedup["polygon"].apply(lambda p: p.length if p is not None else np.nan)
    df_dedup["pp_index"]  = df_dedup.apply(lambda r: (4 * np.pi * r["pp_area"]) / (r["pp_perim"] ** 2) if (pd.notna(r["pp_perim"]) and r["pp_perim"] > 0) else np.nan, axis=1)
else: df_dedup["pp_index"] = np.nan

# =============================================================================
# 5. DETEKCE OUTLIERŮ (Metoda cenové homogenizace / ÚP-Anchored IQR)
# =============================================================================

print("[5] Detekce outlierů (Metoda cenové homogenizace dle oceňovaného ÚP) …")

# Předběžné robustní mediány
median_jc_per_up_raw = df_dedup.groupby("UP")["JC"].median()
median_oce_up_raw = median_jc_per_up_raw.get(OCE_UP, df_dedup["JC"].median())

# A) Cenová Homogenizace: Přepočet všech cen v trhu na cenovou hladinu cílového ÚP
df_dedup["JC_norm"] = df_dedup.apply(
    lambda r: r["JC"] * (median_oce_up_raw / median_jc_per_up_raw[r["UP"]]), axis=1
)

# IQR nad HOMOGENIZOVANÝMI daty (odřízne skutečné extrémy bez ohledu na převahu levné zeleně)
Q1_jc = df_dedup["JC_norm"].quantile(0.25)
Q3_jc = df_dedup["JC_norm"].quantile(0.75)
IQR_jc = Q3_jc - Q1_jc
if IQR_jc == 0: IQR_jc = median_oce_up_raw * 0.20 # Pojistka při extrémní homogenitě

lower_bound_jc_norm = Q1_jc - (OUTLIER_IQR_MULT * IQR_jc)
upper_bound_jc_norm = Q3_jc + (OUTLIER_IQR_MULT * IQR_jc)
if lower_bound_jc_norm <= 0: lower_bound_jc_norm = max(df_dedup["JC_norm"].quantile(0.05), 1.0)

# B) Výměra: Ochranné logaritmické pásmo roztažené na zadanou výměru
log_vym = np.log(df_dedup["vymera"])
Q1_v = log_vym.quantile(0.25)
Q3_v = log_vym.quantile(0.75)
IQR_v = Q3_v - Q1_v

L_v_log = Q1_v - (OUTLIER_IQR_MULT * IQR_v)
U_v_log = Q3_v + (OUTLIER_IQR_MULT * IQR_v)

log_oce_vym = np.log(OCE_VYMERA)
# Dynamické roztažení tolerančního pásma (záruka nevyřazení srovnatelné plochy)
L_v_log = min(L_v_log, log_oce_vym - np.log(10)) 
U_v_log = max(U_v_log, log_oce_vym + np.log(10))

lower_bound_vym = np.exp(L_v_log)
upper_bound_vym = np.exp(U_v_log)

# Aplikace filtrů
outlier_jc = (df_dedup["JC_norm"] < lower_bound_jc_norm) | (df_dedup["JC_norm"] > upper_bound_jc_norm)
outlier_vym = (df_dedup["vymera"] < lower_bound_vym) | (df_dedup["vymera"] > upper_bound_vym)

df_dedup["outlier"] = outlier_jc | outlier_vym
df_outliers = df_dedup[df_dedup["outlier"]].copy()
df_clean    = df_dedup[~df_dedup["outlier"]].copy().reset_index(drop=True)

print(f"    Přípustné pásmo HOMOGENIZOVANÉ ceny: {lower_bound_jc_norm:,.0f} – {upper_bound_jc_norm:,.0f} Kč/m²")
print(f"    Přípustné pásmo Výměry: {lower_bound_vym:,.0f} – {upper_bound_vym:,.0f} m²")
print(f"    Celkem vyloučeno extrémů: {len(df_outliers)}. Zůstalo v čistém vzorku: {len(df_clean)}")

# =============================================================================
# 6. POPISNÉ STATISTIKY
# =============================================================================

print("[6] Popisné statistiky a analýza podmnožiny …")
def weighted_stats(series, weights=None):
    s = series.dropna()
    w = weights.loc[s.index].fillna(0) if weights is not None else None
    return {
        "počet": len(s), "průměr": np.average(s, weights=w) if w is not None and w.sum() > 0 else s.mean(),
        "min": s.min(), "P5": s.quantile(0.05), "Q1": s.quantile(0.25), "medián": s.median(),
        "Q3": s.quantile(0.75), "P95": s.quantile(0.95), "max": s.max(), "std": s.std()
    }

stats_cas = weighted_stats(df_clean["cas_v_dnech"])
stats_vym = weighted_stats(df_clean["vymera"], df_clean["vymera"])
stats_jc  = weighted_stats(df_clean["JC"], df_clean["vymera"])

df_sub_up = df_clean[df_clean["UP"] == OCE_UP]
has_sub_up = len(df_sub_up) >= 3 
if has_sub_up: 
    stats_sub_jc = weighted_stats(df_sub_up["JC"], df_sub_up["vymera"])

# =============================================================================
# 7. VÝPOČET KOREKČNÍCH KOEFICIENTŮ K1–K5 (Uvolněné limity)
# =============================================================================

print("[7] Výpočet korekčních koeficientů K1–K5 …")

# K1 (ČAS)
df_reg_k1 = df_clean[df_clean["JC"] > 0].copy()
df_reg_k1["log_JC"] = np.log(df_reg_k1["JC"])
df_reg_k1["cas_roky"] = -df_reg_k1["cas_v_dnech"] / 365.25 # Chronologický čas 

if len(df_reg_k1) > 1 and df_reg_k1["cas_roky"].nunique() > 1:
    slope_k1, intercept_k1, r_k1, p_k1, _ = stats.linregress(df_reg_k1["cas_roky"], df_reg_k1["log_JC"])
else: 
    slope_k1, r_k1 = 0.0, 0.0
if pd.isna(slope_k1): slope_k1, r_k1 = 0.0, 0.0

rocni_zmena_k1 = (np.exp(slope_k1) - 1) * 100
df_clean["K1"] = np.exp(slope_k1 * (df_clean["cas_v_dnech"] / 365.25)).clip(0.05, 20.00)
print(f"    K1 – regresní roční sklon: {slope_k1:.6f} (změna {rocni_zmena_k1:.2f} % p.a.) | R²={r_k1**2:.3f}")

# K2 (LOKALITA)
median_jc_per_ku = df_clean.groupby("ku_nazev")["JC"].median()
median_oce_ku    = median_jc_per_ku.get(OCE_KU, df_clean["JC"].median())
df_clean["K2"] = (median_oce_ku / df_clean["ku_nazev"].map(median_jc_per_ku)).clip(0.05, 20.00)

# K3 (ÚP - Maximálně uvolněný limit pro obrovské cenové skoky zeleň vs. stavba)
median_jc_per_up = df_clean.groupby("UP")["JC"].median()
median_oce_up    = median_jc_per_up.get(OCE_UP, df_clean["JC"].median())
df_clean["K3"] = (median_oce_up / df_clean["UP"].map(median_jc_per_up)).clip(0.01, 150.00)

# K4 (VELIKOST)
df_reg_k4 = df_clean[df_clean["JC"] > 0].copy()
df_reg_k4["log_JC"]  = np.log(df_reg_k4["JC"])
df_reg_k4["log_vym"] = np.log(df_reg_k4["vymera"])
if len(df_reg_k4) > 1 and df_reg_k4["log_vym"].nunique() > 1:
    slope_k4, intercept_k4, r_k4, p_k4, _ = stats.linregress(df_reg_k4["log_vym"], df_reg_k4["log_JC"])
else: slope_k4, r_k4 = 0.0, 0.0
if pd.isna(slope_k4): slope_k4, r_k4 = 0.0, 0.0

df_clean["K4"] = ((OCE_VYMERA / df_clean["vymera"]) ** slope_k4).clip(0.05, 20.00)

# K5 (TVAR)
pp_valid_data = df_clean.loc[df_clean["pp_index"].notna() & (df_clean["pp_index"] > 0), "pp_index"]
pp_valid = df_clean["pp_index"].notna() & (df_clean["pp_index"] > 0)
df_clean.loc[pp_valid,  "K5"] = (OCE_PP_INDEX / df_clean.loc[pp_valid, "pp_index"]).clip(0.05, 20.00)
df_clean.loc[~pp_valid, "K5"] = 1.0   

# =============================================================================
# 8. UPRAVENÁ JEDNOTKOVÁ CENA
# =============================================================================

print("[8] Výpočet upravené JC …")
df_clean["JC_upravena"] = df_clean["JC"] * df_clean["K1"] * df_clean["K2"] * df_clean["K3"] * df_clean["K4"] * df_clean["K5"]

# =============================================================================
# 9. ETALON JC
# =============================================================================

print("[9] Výpočet ETALONU JC …")
valid_mask = df_clean["JC_upravena"].notna() & df_clean["vymera"].notna()
if valid_mask.sum() > 0:
    etalon_vazeny = np.average(df_clean.loc[valid_mask, "JC_upravena"], weights=df_clean.loc[valid_mask, "vymera"])
else:
    etalon_vazeny = df_clean["JC"].median()
    if pd.isna(etalon_vazeny): etalon_vazeny = 1000.0

etalon_median  = df_clean["JC_upravena"].median() if pd.notna(df_clean["JC_upravena"].median()) else etalon_vazeny
etalon_std     = df_clean["JC_upravena"].std() if pd.notna(df_clean["JC_upravena"].std()) else 0.0
etalon_p25     = df_clean["JC_upravena"].quantile(0.25) if pd.notna(df_clean["JC_upravena"].quantile(0.25)) else etalon_median
etalon_p75     = df_clean["JC_upravena"].quantile(0.75) if pd.notna(df_clean["JC_upravena"].quantile(0.75)) else etalon_median

print(f"\n    *** ETALON JC ***\n    Vážený průměr : {etalon_vazeny:,.0f} Kč/m²")

# =============================================================================
# 10. CITLIVOSTNÍ ANALÝZA 
# =============================================================================

print("[10] Citlivostní analýza …")
def prepocitej_etalon_citlivost(df_base, zmeny_params):
    df_tmp = df_base.copy()
    s_k1 = slope_k1 * (1 + zmeny_params.get("slope_k1", 0))
    m_k2 = median_oce_ku * (1 + zmeny_params.get("median_k2", 0))
    m_k3 = median_oce_up * (1 + zmeny_params.get("median_k3", 0))
    s_k4 = slope_k4 * (1 + zmeny_params.get("slope_k4", 0))
    p_k5 = OCE_PP_INDEX * (1 + zmeny_params.get("pp_k5", 0))
    
    k1 = np.exp(s_k1 * (df_tmp["cas_v_dnech"] / 365.25)).clip(0.05, 20.00)
    median_jc_per_ku = df_tmp.groupby("ku_nazev")["JC"].median()
    k2 = (m_k2 / df_tmp["ku_nazev"].map(median_jc_per_ku)).clip(0.05, 20.00)
    median_jc_per_up = df_tmp.groupby("UP")["JC"].median()
    k3 = (m_k3 / df_tmp["UP"].map(median_jc_per_up)).clip(0.01, 150.00)
    k4 = ((OCE_VYMERA / df_tmp["vymera"]) ** s_k4).clip(0.05, 20.00)
    
    k5 = pd.Series(1.0, index=df_tmp.index)
    v_pp = df_tmp["pp_index"].notna() & (df_tmp["pp_index"] > 0)
    k5.loc[v_pp] = (p_k5 / df_tmp.loc[v_pp, "pp_index"]).clip(0.05, 20.00)
    
    jc_uprenava = df_tmp["JC"] * k1 * k2 * k3 * k4 * k5
    v_mask = jc_uprenava.notna() & df_tmp["vymera"].notna()
    if v_mask.sum() > 0: return np.average(jc_uprenava[v_mask], weights=df_tmp.loc[v_mask, "vymera"])
    return df_tmp["JC"].median()

param_map = {"K1": "slope_k1", "K2": "median_k2", "K3": "median_k3", "K4": "slope_k4", "K5": "pp_k5"}
sensitivity_results = {}
for k_label, p_name in param_map.items():
    vaz_up = prepocitej_etalon_citlivost(df_clean, {p_name: +SENSITIVITY_DELTA})
    vaz_dn = prepocitej_etalon_citlivost(df_clean, {p_name: -SENSITIVITY_DELTA})
    if pd.isna(vaz_up): vaz_up = etalon_vazeny
    if pd.isna(vaz_dn): vaz_dn = etalon_vazeny
    sensitivity_results[k_label] = {"base_vazeny": etalon_vazeny, "up_vazeny": vaz_up, "dn_vazeny": vaz_dn, "rozpeti": abs(vaz_up - vaz_dn)}
tornado_order = sorted(sensitivity_results.keys(), key=lambda k: sensitivity_results[k]["rozpeti"], reverse=True)

# =============================================================================
# 11. GRAFY (Zelená = Levné, Červená = Drahé)
# =============================================================================

print("[11] Generuji grafy …")
GRAPH_FILES = {}
sns.set_theme(style="whitegrid", font="Arial")

# ── 11.1 Histogram JC (Zobrazuje skutečné hodnoty, ale pro ořez používá norm.) ──
fig, ax = plt.subplots(figsize=(10, 5))
bins = np.logspace(np.log10(df_dedup["JC"].min() if df_dedup["JC"].min()>0 else 1), np.log10(df_dedup["JC"].max()), 20)
ax.hist(df_dedup["JC"], bins=bins, color="#4472C4", edgecolor="white", alpha=0.85)
ax.set_xscale('log')
ax.axvline(stats_jc["medián"], color="red", linestyle="--", linewidth=1.5, label=f'Medián čistého trhu ({stats_jc["medián"]:,.0f})')
ax.axvline(stats_jc["průměr"], color="green", linestyle="--", linewidth=1.5, label=f'Vážený průměr čistého trhu ({stats_jc["průměr"]:,.0f})')
ax.set_ylabel("Počet vkladů"); ax.set_xlabel("Jednotková cena [Kč/m²] (Log měřítko)"); ax.set_title("Histogram jednotkových cen pozemků na trhu")
ax.xaxis.set_major_formatter(mticker.FuncFormatter(lambda x, _: f"{x:,.0f}")); ax.xaxis.set_minor_formatter(mticker.NullFormatter()); ax.legend(fontsize=9)
plt.tight_layout(); p = os.path.join(SCRIPT_DIR, "graf_histogram_JC.png"); fig.savefig(p, dpi=150); plt.close(fig); GRAPH_FILES["histogram_JC"] = p

# ── 11.2 Časový vývoj JC ──
fig, ax = plt.subplots(figsize=(12, 5))
sc = ax.scatter(df_clean["datum_podani"], df_clean["JC"], c=df_clean["JC"], cmap=CMAP_JC, norm=Normalize(vmin=df_clean["JC"].min(), vmax=df_clean["JC"].max()), s=60, zorder=3, alpha=0.85)
x_num = df_clean["datum_podani"].map(pd.Timestamp.toordinal)
m, b = stats.linregress(x_num, df_clean["JC"])[:2] if len(x_num) > 1 else (0.0, df_clean["JC"].median())
x_range = pd.date_range(df_clean["datum_podani"].min(), df_clean["datum_podani"].max(), periods=100)
ax.plot(x_range, m * x_range.map(pd.Timestamp.toordinal) + b, color="navy", linewidth=2, label="Lineární trend")
for val, lbl, c, ls in [(stats_jc["medián"], "Medián", "red", "--"), (stats_jc["průměr"], "Průměr", "green", "--")]: ax.axhline(val, linestyle=ls, color=c, label=f"{lbl} {val:,.0f}")
plt.colorbar(sc, ax=ax, label="JC [Kč/m²]"); ax.set_ylabel("JC [Kč/m²]"); ax.set_title("Časový vývoj jednotkových cen pozemků"); ax.legend(); plt.tight_layout()
p = os.path.join(SCRIPT_DIR, "graf_casovy_vyvoj.png"); fig.savefig(p, dpi=150); plt.close(fig); GRAPH_FILES["casovy_vyvoj"] = p

# ── 11.3 ÚP ──
up_stats = df_clean.groupby("UP")["JC"].agg(median="median", mean="mean", std="std").reset_index().sort_values("median", ascending=False)
fig, ax = plt.subplots(figsize=(10, 5))
x = np.arange(len(up_stats))
ax.bar(x - 0.2, up_stats["median"], 0.4, label="Medián", color="#4472C4"); ax.bar(x + 0.2, up_stats["mean"], 0.4, label="Průměr", color="#ED7D31", alpha=0.85)
ax.errorbar(x + 0.2, up_stats["mean"], yerr=up_stats["std"].fillna(0), fmt="none", ecolor="black", capsize=4, linewidth=1)
ax.set_xticks(x); ax.set_xticklabels(up_stats["UP"], rotation=45, ha="right"); ax.set_title("Porovnání JC dle ÚP"); ax.legend(); plt.tight_layout()
p = os.path.join(SCRIPT_DIR, "graf_JC_dle_UP.png"); fig.savefig(p, dpi=150); plt.close(fig); GRAPH_FILES["JC_dle_UP"] = p

# ── 11.4 Velikost ──
fig, ax = plt.subplots(figsize=(10, 5))
sc4 = ax.scatter(df_clean["vymera"], df_clean["JC"], c=df_clean["JC"], cmap=CMAP_JC, norm=Normalize(vmin=df_clean["JC"].min(), vmax=df_clean["JC"].max()), s=60, alpha=0.8)
vym_range = np.linspace(df_clean["vymera"].min(), df_clean["vymera"].max(), 200)
ax.plot(vym_range, np.exp(slope_k4 * np.log(vym_range) + intercept_k4) if intercept_k4 else np.repeat(df_clean["JC"].median(), 200), color="navy", linewidth=2, label=f"Log-log trend (β={slope_k4:.3f})")
for val, lbl, c, ls, lw in [(stats_jc["medián"], "Medián", "red", "--", 1.5), (stats_jc["průměr"], "Průměr", "green", "--", 1.5)]:
    if pd.notna(val): ax.axhline(val, linestyle=ls, linewidth=lw, color=c, alpha=0.8, label=f"{lbl} {val:,.0f}")
ax.set_xlabel("Výměra [m²]"); ax.set_ylabel("JC [Kč/m²]"); ax.set_title("Závislost JC na velikosti"); plt.colorbar(sc4); ax.legend(); plt.tight_layout()
p = os.path.join(SCRIPT_DIR, "graf_JC_vs_vymera.png"); fig.savefig(p, dpi=150); plt.close(fig); GRAPH_FILES["JC_vs_vymera"] = p

# ── 11.5 Outliers ──
fig, ax = plt.subplots(figsize=(10, 5))
ax.scatter(df_clean["datum_podani"], df_clean["JC_norm"], color="#4472C4", s=40, label="Standardní trh", alpha=0.7)
if len(df_outliers)>0: ax.scatter(df_outliers["datum_podani"], df_outliers["JC_norm"], color="red", s=80, marker="X", label="Outlier", zorder=5)
ax.axhline(lower_bound_jc_norm, color="orange", linestyle="--", label=f"Dolní ochranné pásmo ({lower_bound_jc_norm:,.0f})")
ax.axhline(upper_bound_jc_norm, color="orange", linestyle="--", label=f"Horní ochranné pásmo ({upper_bound_jc_norm:,.0f})")
ax.set_title(f"Extrémy na Homogenizovaných Datech (Ceny převedeny na ÚP {OCE_UP})"); ax.legend(fontsize=8); ax.set_ylabel("Cena JC_norm"); plt.tight_layout()
p = os.path.join(SCRIPT_DIR, "graf_outliery.png"); fig.savefig(p, dpi=150); plt.close(fig); GRAPH_FILES["outliery"] = p

# ── 11.6 Tornado ──
fig, ax = plt.subplots(figsize=(10, 5))
labels_map = {"K1": "K1 – Čas", "K2": "K2 – Lokalita", "K3": "K3 – Typ ÚP", "K4": "K4 – Velikost", "K5": "K5 – Tvar"}
min_x, max_x = float('inf'), float('-inf')
for i, k in enumerate(tornado_order):
    r = sensitivity_results[k]
    if pd.isna(r["dn_vazeny"]) or pd.isna(r["up_vazeny"]): continue
    lo, hi = min(r["dn_vazeny"], r["up_vazeny"]), max(r["dn_vazeny"], r["up_vazeny"])
    min_x, max_x = min(min_x, lo), max(max_x, hi)
    ax.barh(i, hi - etalon_vazeny, left=etalon_vazeny, color="#70AD47", height=0.5); ax.barh(i, lo - etalon_vazeny, left=etalon_vazeny, color="#FF0000", height=0.5)
if np.isinf(min_x) or np.isinf(max_x) or pd.isna(min_x) or pd.isna(max_x):
    bezpecny_etalon = etalon_vazeny if pd.notna(etalon_vazeny) else 1000.0
    min_x, max_x = bezpecny_etalon * 0.8, bezpecny_etalon * 1.2
rozpeti_x = max(max_x - min_x, 100)
ax.set_xlim(min_x - rozpeti_x*0.25, max_x + rozpeti_x*0.25); ax.axvline(etalon_vazeny, color="black", linewidth=1.5)
ax.set_yticks(range(len(tornado_order))); ax.set_yticklabels([labels_map[k] for k in tornado_order]); ax.invert_yaxis(); ax.set_title("Citlivostní analýza"); plt.tight_layout()
p = os.path.join(SCRIPT_DIR, "graf_tornado.png"); fig.savefig(p, dpi=150); plt.close(fig); GRAPH_FILES["tornado"] = p

# ── 11.7 Boxplot ÚP ──
if has_sub_up:
    fig, ax = plt.subplots(figsize=(8, 5))
    df_clean['Skupina'] = np.where(df_clean['UP'] == OCE_UP, f'ÚP ({OCE_UP})', 'Ostatní ÚP')
    sns.boxplot(data=df_clean, x='Skupina', y='JC', ax=ax, palette=['#FF9999', '#99CCFF']); ax.set_title(f"Rozptyl JC: ÚP {OCE_UP} vůči zbytku"); plt.tight_layout()
    p = os.path.join(SCRIPT_DIR, "graf_podmnozina_up.png"); fig.savefig(p, dpi=150); plt.close(fig); GRAPH_FILES["podmnozina_up"] = p

# ── 11.8 K5 PP Index ──
if not pp_valid_data.empty:
    fig, ax = plt.subplots(figsize=(10, 5))
    sns.histplot(pp_valid_data, kde=True, color="#9DC3E6", bins=20, ax=ax)
    ax.axvline(OCE_PP_INDEX, color="red", linestyle="--", linewidth=2, label=f"Oceňovaný pozemek ({OCE_PP_INDEX})")
    ax.axvline(pp_valid_data.median(), color="green", linestyle=":", linewidth=2, label=f"Medián trhu ({pp_valid_data.median():.2f})")
    ax.set_title("Rozložení tvarového indexu (Polsby-Popper)"); ax.legend(); plt.tight_layout()
    p = os.path.join(SCRIPT_DIR, "graf_K5_tvar.png"); fig.savefig(p, dpi=150); plt.close(fig); GRAPH_FILES["K5_tvar"] = p

# ── 11.9 Prostorová Mapa a Heatmapa ──
import geopandas as gpd
import contextily as ctx
from matplotlib.lines import Line2D

LAT_COL, LON_COL = "refPoint_lat", "refPoint_lon" 
valid_gps = df_dedup.dropna(subset=[LAT_COL, LON_COL])

if not valid_gps.empty:
    gdf = gpd.GeoDataFrame(valid_gps, geometry=gpd.points_from_xy(valid_gps[LON_COL], valid_gps[LAT_COL]), crs="EPSG:4326").to_crs("EPSG:3857")
    
    # Statická mapa
    fig, ax = plt.subplots(figsize=(10, 8))
    gdf.plot(ax=ax, column="JC", cmap=CMAP_JC, norm=Normalize(vmin=df_dedup["JC"].min(), vmax=df_dedup["JC"].max()), markersize=80, edgecolor="black", alpha=0.8, legend=True)
    oce_gdf = gpd.GeoDataFrame(index=[0], crs="EPSG:4326", geometry=[gpd.points_from_xy([OCE_LON], [OCE_LAT])[0]]).to_crs("EPSG:3857")
    ax.scatter(oce_gdf.geometry.x, oce_gdf.geometry.y, color="blue", marker="^", s=300, edgecolors="white", zorder=10)
    ctx.add_basemap(ax, crs=gdf.crs.to_string(), source=ctx.providers.Esri.WorldImagery); ax.set_axis_off(); plt.tight_layout()
    p = os.path.join(SCRIPT_DIR, "graf_prostorova_mapa.png"); fig.savefig(p, dpi=150); plt.close(fig); GRAPH_FILES["prostorova_mapa"] = p
    
    # Heatmapa (čistá data)
    valid_gps_clean = df_clean.dropna(subset=[LAT_COL, LON_COL])
    if len(valid_gps_clean) > 4:
        try:
            gdf_clean = gpd.GeoDataFrame(valid_gps_clean, geometry=gpd.points_from_xy(valid_gps_clean[LON_COL], valid_gps_clean[LAT_COL]), crs="EPSG:4326").to_crs("EPSG:3857")
            x, y, z = gdf_clean.geometry.x.values, gdf_clean.geometry.y.values, gdf_clean['JC'].values
            xi, yi = np.linspace(x.min() - 1500, x.max() + 1500, 300), np.linspace(y.min() - 1500, y.max() + 1500, 300)
            XI, YI = np.meshgrid(xi, yi)
            ZI = interp.griddata((x, y), z, (XI, YI), method='linear')
            fig, ax = plt.subplots(figsize=(10, 8))
            contour = ax.contourf(XI, YI, ZI, levels=30, cmap=CMAP_JC, alpha=0.85)
            plt.colorbar(contour, ax=ax, fraction=0.046, pad=0.04)
            gdf_clean.plot(ax=ax, color='black', markersize=15, alpha=0.4)
            ax.scatter(oce_gdf.geometry.x, oce_gdf.geometry.y, color="blue", marker="^", s=250, edgecolor="white", zorder=10)
            ctx.add_basemap(ax, crs=gdf_clean.crs.to_string(), source=ctx.providers.Esri.WorldImagery); ax.set_axis_off(); plt.tight_layout()
            p = os.path.join(SCRIPT_DIR, "graf_heatmap.png"); fig.savefig(p, dpi=150); plt.close(fig); GRAPH_FILES["heatmap"] = p
        except: pass

# =============================================================================
# 12. INTERAKTIVNÍ MAPA (Folium)
# =============================================================================
print("[12] Generuji interaktivní mapu …")
import branca.colormap as cm  # <--- PŘIDEJTE TENTO ŘÁDEK SEM

valid_coords = df_dedup.dropna(subset=["refPoint_lat", "refPoint_lon"])

if not valid_coords.empty:
    sw = [valid_coords["refPoint_lat"].min(), valid_coords["refPoint_lon"].min()]
    ne = [valid_coords["refPoint_lat"].max(), valid_coords["refPoint_lon"].max()]
    map_center = [(sw[0] + ne[0]) / 2, (sw[1] + ne[1]) / 2]
else:
    sw, ne = None, None
    map_center = [OCE_LAT, OCE_LON]

m = folium.Map(location=map_center, prefer_canvas=True)
folium.TileLayer("OpenStreetMap").add_to(m)

jc_min = df_dedup["JC"].min() if not pd.isna(df_dedup["JC"].min()) else 0
jc_max = df_dedup["JC"].max() if not pd.isna(df_dedup["JC"].max()) else 1000
if jc_min == jc_max: jc_max += 1

colormap = cm.LinearColormap(colors=['#1a9641', '#a6d96a', '#fdae61', '#d7191c'], vmin=jc_min, vmax=jc_max, caption='JC [Kč/m²]')
m.add_child(colormap)

for _, row in df_dedup.iterrows():
    if pd.isna(row.get("refPoint_lat")): continue
    popup_txt = (f"<b>Vklad: {row['cislo_vkladu']}</b><br>KÚ: {row.get('ku_nazev','–')}<br>UP: {row.get('UP','–')}<br>JC: <b>{row['JC']:,.0f} Kč/m²</b><br>Výměra: {row['vymera']:,.0f} m²")
    folium.CircleMarker(location=[row["refPoint_lat"], row["refPoint_lon"]], radius=8, color="black", weight=1, fill=True, fill_color=colormap(row["JC"]), fill_opacity=0.9, popup=folium.Popup(popup_txt, max_width=250), tooltip=f"{row['JC']:,.0f} Kč/m²").add_to(m)

folium.Marker(location=[OCE_LAT, OCE_LON], popup=folium.Popup(f"<b>OCEŇOVANÁ NEMOVITOST</b><br>{OCE_UP}<br>{OCE_VYMERA:,.0f} m²", max_width=220), icon=folium.Icon(color="blue", icon="home", prefix="fa")).add_to(m)

if sw and ne: 
    m.fit_bounds([sw, ne])
    
m.save(OUT_MAP)
print(f"    Uložena mapa: {OUT_MAP}")

# =============================================================================
# 13. EXCEL VÝSTUP
# =============================================================================
print("[13] Generuji Excel výstup …")
with pd.ExcelWriter(OUT_XLSX, engine="openpyxl") as writer:
    df_clean.to_excel(writer, sheet_name="Data_Očištěná", index=False)
    df_outliers.to_excel(writer, sheet_name="Extrémy", index=False)

# =============================================================================
# 14. WORD DOKUMENT A POMOCNÉ FUNKCE
# =============================================================================
print("[14] Generuji detailní Word dokument se všemi grafy…")

def set_cell_background(cell, hex_color):
    cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>'))

def add_hyperlink(paragraph, url, text):
    part = paragraph.part; r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink'); hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r'); rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color'); color.set(qn('w:val'), '0000FF'); rPr.append(color)
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '14'); rPr.append(sz)
    szCs = OxmlElement('w:szCs'); szCs.set(qn('w:val'), '14'); rPr.append(szCs)
    new_run.append(rPr); text_element = OxmlElement('w:t'); text_element.text = text; new_run.append(text_element)
    hyperlink.append(new_run); paragraph._p.append(hyperlink)
    return hyperlink

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level); h.runs[0].font.name = "Arial"; return h

def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph(text); r = p.runs[0] if p.runs else p.add_run(text)
    r.font.name = "Arial"; r.font.size = Pt(11)
    if bold: r.bold = True
    return p

def add_table_from_df(doc, df_t):
    table = doc.add_table(rows=1, cols=len(df_t.columns)); table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df_t.columns):
        hdr_cells[i].text = str(col); hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
        set_cell_background(hdr_cells[i], "2E74B5")
    for _, row in df_t.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val) if pd.notna(val) else "–"
            row_cells[i].paragraphs[0].runs[0].font.size = Pt(9)
    return table

doc = Document()
doc.styles["Normal"].font.name = "Arial"; doc.styles["Normal"].font.size = Pt(11)

# TITULNÍ STRANA
doc.add_heading("ANALÝZA TRHU NEMOVITOSTÍ", 0)
doc.add_heading("Stanovení jednotkové obvyklé ceny pozemku – ETALON JC", 1)
doc.add_paragraph("")
for label, val in [("Věc:", f"Pozemky k.ú. {OCE_KU}, {OCE_OKRES}"), ("Typ pozemku:", f"Stavební pozemek (UP={OCE_UP})"), ("Výměra:", f"{OCE_VYMERA:,.0f} m²"), ("Datum ocenění:", OCE_DATUM.strftime("%d. %m. %Y"))]:
    p = doc.add_paragraph(); r1 = p.add_run(f"{label}  "); r1.bold = True; r1.font.name = "Arial"; r2 = p.add_run(val); r2.font.name = "Arial"
doc.add_page_break()

# KAP 1: ÚVOD
add_heading(doc, "1. Úvod a předmět analýzy", 1)
add_paragraph(doc, f"Předmětem této analýzy trhu je stanovení jednotkové obvyklé ceny (dále jen „JC\") pozemku situovaného v katastrálním území {OCE_KU}, okres {OCE_OKRES}, ke dni {OCE_DATUM.strftime('%d. %m. %Y')}.\n\nOceňovaný pozemek má celkovou výměru {OCE_VYMERA:,.0f} m² a dle aktuálně platného územního plánu náleží do funkční plochy označené kódem „{OCE_UP}\". Tvar pozemku byl exaktně vyhodnocen s Polsby-Popper indexem na hodnotu {OCE_PP_INDEX}.\n\nCílem analýzy je prostřednictvím transparentního statistického modelu (tzv. White-Box) dle IVS identifikovat srovnatelné tržní transakce, matematicky odvodit korekční koeficienty zohledňující parametry nemovitosti a stanovit tzv. ETALON JC – referenční jednotkovou cenu průměrného pozemku.")

# KAP 2: METODIKA A DATA
add_heading(doc, "2. Metodika a zdrojová data", 1)
add_paragraph(doc, "Základem analýzy je databáze tržních transakcí s pozemky z katastru nemovitostí.")
if "prostorova_mapa" in GRAPH_FILES:
    doc.add_picture(GRAPH_FILES["prostorova_mapa"], width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Obrázek 1: Prostorové a cenové rozložení vzorku").alignment = WD_ALIGN_PARAGRAPH.CENTER

add_paragraph(doc, "\nSurový datový vzorek (Všechny parcely, řazeno chronologicky):", bold=True)
df_raw_table = df.dropna(subset=["JC", "vymera"]).copy()
df_raw_table.sort_values(by="datum_podani", ascending=False, inplace=True)
tab_cols = ["Datum podání", "Vklad", "K.Ú.", "Parc.č.", "ÚP", "Výměra [m²]", "JC [Kč/m²]"]
table = doc.add_table(rows=1, cols=len(tab_cols))
table.style = "Table Grid"
hdr_cells = table.rows[0].cells
for i, name in enumerate(tab_cols):
    hdr_cells[i].text = name; hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_cell_background(hdr_cells[i], "2E74B5")

def get_hex_color(val, vmin, vmax, cmap_name="RdYlGn_r"):
    import matplotlib.colors as mcolors  
    if pd.isna(val): return "FFFFFF"
    norm_val = (val - vmin) / max(vmax - vmin, 1); rgba = plt.cm.get_cmap(cmap_name)(norm_val)
    return mcolors.to_hex(rgba).replace("#", "").upper()

vmin_jc, vmax_jc = df_raw_table["JC"].min(), df_raw_table["JC"].max()
vmin_vym, vmax_vym = df_raw_table["vymera"].min(), df_raw_table["vymera"].max()
prev_year, year_bg = None, "FFFFFF"
for _, row in df_raw_table.iterrows():
    r_cells = table.add_row().cells
    r_cells[0].text = str(row["datum_den"]); r_cells[1].text = str(row.get("cislo_vkladu", "–")); r_cells[2].text = str(row.get("ku_nazev", "–"))
    cell_parc = r_cells[3]; cell_parc.text = ""; paragraph = cell_parc.paragraphs[0]; paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d_parc = str(row.get("parcel_number", "–"))
    if pd.notna(row.get("ruian_parcela_id")): add_hyperlink(paragraph, f"https://nahlizenidokn.cuzk.gov.cz/ZobrazObjekt.aspx?&typ=parcela&id={str(int(row.get('ruian_parcela_id')))}", d_parc)
    else: paragraph.text = d_parc
    r_cells[4].text = str(row.get("UP", "–")); r_cells[5].text = f"{row.get('vymera', 0):,.0f}"; r_cells[6].text = f"{row.get('JC', 0):,.0f}"
    for cell in r_cells: cell.paragraphs[0].runs[0].font.size = Pt(7)
    curr_yr = str(row["datum_den"])[:4]
    if prev_year and curr_yr != prev_year: year_bg = "F2F2F2" if year_bg == "FFFFFF" else "FFFFFF"
    prev_year = curr_yr
    for i in range(5): set_cell_background(r_cells[i], year_bg)
    set_cell_background(r_cells[5], get_hex_color(row.get("vymera", 0), vmin_vym, vmax_vym, "Blues"))
    set_cell_background(r_cells[6], get_hex_color(row.get("JC", 0), vmin_jc, vmax_jc, "RdYlGn_r"))
doc.add_page_break()

# KAP 3: DETEKCE EXTRÉMŮ
add_heading(doc, "3. Zacílená detekce a zpracování extrémních hodnot", 1)
add_paragraph(doc, f"Trh nemovitostí má tzv. multimodální rozdělení (obsahuje diametrálně odlišné typy pozemků, např. levná pole a drahé stavební parcely). Pokud by byl pro detekci extrémů použit pouze jeden plošný globální filtr, došlo by ke smíchání nekompatibilních trhů a nežádoucímu odříznutí relevantních transakcí.\n\nAby se zabránilo znehodnocení srovnávacího vzorku, využívá tento model inovativní metodu Cenové homogenizace (Target-Anchored IQR):\n\nA) Jednotková cena (JC): Všechny ceny v trhu jsou před kontrolou extrémů matematicky převedeny (homogenizovány) na cenovou hladinu oceňovaného územního plánu (ÚP = {OCE_UP}). Následný IQR filtr bezpečně odřízne skutečné lokální anomálie s přípustným pásmem upraveným na {lower_bound_jc_norm:,.0f} Kč/m² až {upper_bound_jc_norm:,.0f} Kč/m².\n\nB) Výměra pozemku: Kolem oceňované extrémní velikosti ({OCE_VYMERA:,.0f} m²) model dynamicky nastavil elastické logaritmické toleranční pásmo výměr od {lower_bound_vym:,.0f} m² do {upper_bound_vym:,.0f} m², čímž bezpečně garantuje, že relevantní velké celky nebudou odstraněny.")
if "outliery" in GRAPH_FILES:
    doc.add_picture(GRAPH_FILES["outliery"], width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Graf 2: Identifikace extrémů na homogenizovaných datech").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

# KAP 4: POPISNÉ STATISTIKY
add_heading(doc, "4. Základní popisné statistiky (Očištěný vzorek)", 1)
add_paragraph(doc, f"Níže jsou uvedeny klíčové statistické ukazatele jednotkových cen pro čistý vzorek {len(df_clean)} transakcí před aplikací koeficientů.")
add_table_from_df(doc, pd.DataFrame([["Vážený průměr (Hrubý)", f"{stats_jc['průměr']:,.0f} Kč/m²"], ["Medián", f"{stats_jc['medián']:,.0f} Kč/m²"], ["Směrodatná odchylka", f"{stats_jc['std']:,.0f} Kč/m²"]], columns=["Statistický ukazatel", "Hodnota"]))
if "histogram_JC" in GRAPH_FILES:
    doc.add_paragraph("")
    doc.add_picture(GRAPH_FILES["histogram_JC"], width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Graf 3: Histogram jednotkových cen (Zdravý trh)").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

# KAP 5: WHITE-BOX MODELY
add_heading(doc, "5. Exaktní zdůvodnění a výpočet korekčních koeficientů (K1–K5)", 1)
add_paragraph(doc, "Tato část detailně rozkrývá matematické modely použité pro transformaci srovnávacího vzorku na úroveň oceňované nemovitosti. Limity korekcí byly s ohledem na extrémní rozdíly typů ÚP v modelu uvolněny na velkorysý metodický rozptyl, aby matematické ořezy nezkreslovaly realitu.")

# K1
slovo_zmeny = "růst" if rocni_zmena_k1 > 0 else "pokles"
add_paragraph(doc, f"\nK1 – Korekce na čas (Vývoj trhu)\nByla provedena log-lineární regrese v závislosti na čase (chronologických rocích od data ocenění, t=0). Z odhadnutého regresního sklonu (β = {slope_k1:.6f}) lze odvodit, že zkoumaný vzorek v čase vykazuje průměrný roční {slovo_zmeny} cen o {abs(rocni_zmena_k1):.2f} %. Koeficient K1 posunuje historickou cenu k datu ocenění vzorcem K1 = exp(β · t).", bold=True)
if "casovy_vyvoj" in GRAPH_FILES:
    doc.add_picture(GRAPH_FILES["casovy_vyvoj"], width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Graf 4: Časový vývoj jednotkových cen v tržním vzorku").alignment = WD_ALIGN_PARAGRAPH.CENTER

# K2, K3
add_paragraph(doc, f"\nK2 – Korekce na lokalitu\nPočítá se jako podíl mediánu JC v KÚ oceňované nemovitosti ({OCE_KU}) a mediánu JC srovnávací transakce. Medián matematicky přesně identifikuje střední obvyklou hodnotu a eliminuje vliv lokálních anomálií. (Medián JC pro {OCE_KU} = {median_oce_ku:,.0f} Kč/m²).\n\nK3 – Korekce na typ ÚP\nDo čitatele vstupuje medián ceny pozemků s funkčním využitím {OCE_UP} a do jmenovatele medián JC plochy u porovnávané transakce. Rozdíly v ÚP tvoří v tomto modelu nejdůležitější cenový skok (např. převod zeleň -> stavební parcely). (Medián JC pro ÚP={OCE_UP} = {median_oce_up:,.0f} Kč/m²).", bold=True)
if "JC_dle_UP" in GRAPH_FILES:
    doc.add_picture(GRAPH_FILES["JC_dle_UP"], width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Graf 5: Porovnání mediánů JC dle územního plánu").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

# K4
add_paragraph(doc, f"K4 – Korekce na velikost (Diskont na velikost)\nS rostoucí výměrou pozemku klesá jeho jednotková cena (klesající mezní užitek). Model využívá log-log (mocninnou) regresi s konstantní elasticitou. {'Záporná hodnota sklonu (elasticita β = ' + str(round(slope_k4, 4)) + ') tento diskont exaktně potvrzuje.' if slope_k4 < 0 else 'Skutečnost, že regresní sklon není záporný (β = ' + str(round(slope_k4, 4)) + '), naznačuje, že se sleva na velikost neprojevuje.'} Transakce jsou přepočteny na úroveň oceňované výměry {OCE_VYMERA:,.0f} m².", bold=True)
if "JC_vs_vymera" in GRAPH_FILES:
    doc.add_picture(GRAPH_FILES["JC_vs_vymera"], width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Graf 6: Závislost jednotkové ceny na výměře s proloženým log-log trendem").alignment = WD_ALIGN_PARAGRAPH.CENTER

# K5
if not pp_valid_data.empty:
    pp_med = pp_valid_data.median()
    tvar_komentar = f"nad úrovní tržního mediánu (trh = {pp_med:.2f}). Je kompaktnější, což vede k vyššímu koeficientu" if OCE_PP_INDEX > pp_med else f"pod úrovní tržního mediánu (trh = {pp_med:.2f}). Vykazuje mimořádně složitou geometrii, což vyžaduje cenovou penalizaci (K5 < 1)"
else: tvar_komentar = "nelze v tomto vzorku určit"

add_paragraph(doc, f"\nK5 – Korekce na tvar pozemku (Polsby-Popper index)\nGeometrická komplexita je objektivizována Polsby-Popper indexem (1,0 = dokonalý kruh). Oceňovaný pozemek má index {OCE_PP_INDEX}. Ve srovnání s trhem se nachází {tvar_komentar}.", bold=True)
if "K5_tvar" in GRAPH_FILES:
    doc.add_picture(GRAPH_FILES["K5_tvar"], width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Graf 7: Rozložení tvarového indexu v tržním vzorku").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

# KAP 6: PŘEHLED KOEFICIENTŮ
add_heading(doc, "6. Přehled průměrných uplatněných koeficientů", 1)
mean_k1, mean_k2, mean_k3, mean_k4, mean_k5 = df_clean['K1'].mean(), df_clean['K2'].mean(), df_clean['K3'].mean(), df_clean['K4'].mean(), df_clean['K5'].mean()
mean_k_celkem = (df_clean['K1']*df_clean['K2']*df_clean['K3']*df_clean['K4']*df_clean['K5']).mean()

add_paragraph(doc, f"Pro absolutní transparentnost uvádí tabulka průměrné hodnoty koeficientů uplatněné na čistý vzorek {len(df_clean)} transakcí. Odhaluje, co primárně srazilo nebo navýšilo konečnou jednotkovou cenu.")
add_table_from_df(doc, pd.DataFrame([
    ["K1 - Korekce na čas", f"{mean_k1:.3f}"], ["K2 - Korekce na lokalitu", f"{mean_k2:.3f}"],
    ["K3 - Korekce na typ ÚP", f"{mean_k3:.3f}"], ["K4 - Korekce na velikost", f"{mean_k4:.3f}"],
    ["K5 - Korekce na tvar", f"{mean_k5:.3f}"], ["Celkový průměrný násobitel na transakce (K_celk)", f"{mean_k_celkem:.3f}"]
], columns=["Označení koeficientu", "Průměrná aplikovaná hodnota na tržní vzorek"]))

varovani_text = ""
if mean_k4 <= 0.06 or mean_k4 >= 19.95: varovani_text += f"\nUPOZORNĚNÍ: K4 (Velikost) dosahuje extrémních limitů kvůli propastnému rozdílu mezi oceňovanou výměrou a běžnými obchody trhu."
if mean_k5 <= 0.06 or mean_k5 >= 19.95: varovani_text += f"\nUPOZORNĚNÍ: K5 (Tvar) dosahuje extrémních limitů. Pozemek má mimořádně nevýhodný tvar."
if varovani_text: add_paragraph(doc, varovani_text.strip(), bold=True)

# KAP 7: ETALON
add_heading(doc, "7. Stanovení ETALONU JC", 1)
add_paragraph(doc, f"Upravená JC pro každou transakci vznikla součinem její původní JC a všech vypočtených koeficientů. Na základě porovnávací metody byl z čistého vzorku {len(df_clean)} transakcí odvozen ETALON JC k datu {OCE_DATUM.strftime('%d. %m. %Y')} následovně:")
add_table_from_df(doc, pd.DataFrame([["ETALON JC – vážený průměr", f"{etalon_vazeny:,.0f} Kč/m²"], ["ETALON JC – medián", f"{etalon_median:,.0f} Kč/m²"], ["Směrodatná odchylka", f"{etalon_std:,.0f} Kč/m²"]], columns=["Statistický ukazatel", "Výsledná hodnota"]))
add_paragraph(doc, "\nZ výše uplatněných průměrných koeficientů je matematicky patrné, jak vlastnosti zadání zredukovaly nebo povýšily původní tržní medián na výsledný Etalon.")
doc.add_page_break()

# KAP 8: CITLIVOSTNÍ ANALÝZA
add_heading(doc, "8. Citlivostní analýza modelu", 1)
add_paragraph(doc, f"Analýza testuje robustnost ETALONU vůči izolovaným změnám hlavních statistických parametrů pro K1–K5 o ±{int(SENSITIVITY_DELTA*100)} %.")
if "tornado" in GRAPH_FILES:
    doc.add_picture(GRAPH_FILES["tornado"], width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Graf 8: Tornádový graf citlivostní analýzy").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

# KAP 9: CÍLENÁ ANALÝZA ÚP
add_heading(doc, f"9. Cílená analýza podmnožiny: ÚP {OCE_UP}", 1)
if has_sub_up:
    add_paragraph(doc, f"Z celkového čistého vzorku ({len(df_clean)} vkladů) vyhovuje této klasifikaci přesně {len(df_sub_up)} transakcí. Níže je uvedeno statistické srovnání této podmnožiny s celkovým trhem, které empiricky verifikuje úroveň korekčního koeficientu K3.")
    add_table_from_df(doc, pd.DataFrame([["Vážený průměr", f"{stats_sub_jc['průměr']:,.0f} Kč/m²", f"{stats_jc['průměr']:,.0f} Kč/m²"], ["Medián", f"{stats_sub_jc['medián']:,.0f} Kč/m²", f"{stats_jc['medián']:,.0f} Kč/m²"]], columns=["Ukazatel", f"Podmnožina ({OCE_UP})", "Celý trh"]))
    if "podmnozina_up" in GRAPH_FILES:
        doc.add_paragraph("")
        doc.add_picture(GRAPH_FILES["podmnozina_up"], width=Inches(6.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("Graf 9: Boxplot podmnožiny").alignment = WD_ALIGN_PARAGRAPH.CENTER
else: add_paragraph(doc, f"Pro typ územního plánu {OCE_UP} není ve vzorku dostatečný počet transakcí pro izolovanou statistiku.")
doc.add_page_break()

# KAP 10: HEATMAPA
if "heatmap" in GRAPH_FILES:
    add_heading(doc, "10. Prostorové rozložení cenové hladiny v území", 1)
    add_paragraph(doc, "Pro vizuální ověření lokalizačního vlivu na cenu byla vypracována teplotní mapa (Heatmapa). Zelené odstíny reprezentují území s nižší cenovou hladinou, zatímco žluté až červené oblasti signalizují prémiové ceny. Tento model demonstruje reprezentativnost stanovených průměrů pro okolí oceňované nemovitosti (modrý trojúhelník).")
    doc.add_picture(GRAPH_FILES["heatmap"], width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Obrázek 2: Cenová heatmapa").alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save(OUT_DOCX)
print(f"[14] Uložen kompletní Word dokument: {OUT_DOCX}")

print("\n" + "=" * 70)
print("ANALÝZA DOKONČENA")
print(f"  ETALON JC (vážený průměr) = {etalon_vazeny:,.0f} Kč/m²")
print(f"  Celkový průměrný koeficient aplikovaný na trh = {mean_k_celkem:.3f}")
print("=" * 70)